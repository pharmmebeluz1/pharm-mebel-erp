# -*- coding: utf-8 -*-
import csv
import io
import os
import sqlite3
import secrets
from datetime import datetime, date, timedelta
from flask import Flask, jsonify, request, render_template_string, Response, session, redirect, url_for, send_file, send_from_directory, flash
import qrcode
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import hashlib
_original_md5 = hashlib.md5
def _md5_windows7_compat(*args, **kwargs):
    kwargs.pop("usedforsecurity", None)
    return _original_md5(*args, **kwargs)
hashlib.md5 = _md5_windows7_compat
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
import re
from zoneinfo import ZoneInfo

app = Flask(__name__)


def _tashkent_now():
    """Server qayerda ishlashidan qat’i nazar Toshkent vaqtini qaytaradi."""
    return datetime.now(ZoneInfo("Asia/Tashkent"))


def _tashkent_today():
    return _tashkent_now().date().isoformat()


def _safe_time(value, default="18:00"):
    """HH:MM formatini xavfsiz normal holatga keltiradi."""
    value = str(value or "").strip()
    if re.fullmatch(r"(?:[01]\d|2[0-3]):[0-5]\d", value):
        return value
    return default


def _order_deadline(order):
    """Buyurtmaning Toshkent vaqtidagi aniq muddatini qaytaradi.

    Yangi buyurtmalarda rasmiy muddat avans olingandan yoki ishonchli
    mijoz uchun rahbar tasdiqlagandan keyingina ishlaydi. Eski buyurtmalar
    esa avvalgi taxminiy/tugash sanalari bilan ishlashda davom etadi.
    """
    if not order:
        return None
    keys = set(order.keys()) if hasattr(order, "keys") else set()
    workflow_mode = str(order["muddat_tartibi"] or "") if "muddat_tartibi" in keys else "Eski"
    started = str(order["muddat_boshlanish_vaqt"] or "").strip() if "muddat_boshlanish_vaqt" in keys else ""
    if workflow_mode == "Tasdiqdan keyin" and not started:
        return None
    date_value = ""
    for key in ("rasmiy_muddat_sana", "taxminiy_sana", "tugash_sana"):
        if key in keys and str(order[key] or "").strip():
            date_value = str(order[key]).strip()
            break
    if not date_value:
        return None
    time_source = "18:00"
    if "rasmiy_muddat_vaqt" in keys and str(order["rasmiy_muddat_vaqt"] or "").strip():
        time_source = order["rasmiy_muddat_vaqt"]
    elif "taxminiy_vaqt" in keys:
        time_source = order["taxminiy_vaqt"]
    time_value = _safe_time(time_source)
    try:
        return datetime.strptime(f"{date_value} {time_value}", "%Y-%m-%d %H:%M").replace(
            tzinfo=ZoneInfo("Asia/Tashkent")
        )
    except (TypeError, ValueError):
        return None


def _money(value):
    try:
        return f"{float(value or 0):,.0f}".replace(",", " ")
    except (TypeError, ValueError):
        return "0"


def _currency_code(value):
    """Dasturda ishlatiladigan valyutani xavfsiz normal holatga keltiradi."""
    code=str(value or "UZS").strip().upper()
    return code if code in {"UZS", "USD"} else "UZS"


def _currency_money(value, currency="UZS"):
    """Kelishuv valyutasidagi summani foydalanuvchiga tushunarli ko‘rsatadi."""
    code=_currency_code(currency)
    try:
        number=float(value or 0)
    except (TypeError, ValueError):
        number=0.0
    if code=="USD":
        decimals=2 if abs(number-round(number))>0.000001 else 0
        return "$"+f"{number:,.{decimals}f}".replace(",", " ")
    return _money(number)+" so‘m"


def _payment_conversion(order_currency, payment_currency, received_amount, usd_rate=0):
    """Qabul qilingan pulni buyurtma valyutasi va UZS ekvivalentiga aylantiradi.

    USD qatnashgan har bir to‘lovda aynan o‘sha to‘lov kunidagi kurs saqlanadi.
    Eski to‘lovlar va UZS→UZS hisobida kurs 1 bo‘lib qoladi.
    """
    order_code=_currency_code(order_currency)
    payment_code=_currency_code(payment_currency)
    received=float(received_amount or 0)
    rate=float(usd_rate or 0)
    if received<=0:
        raise ValueError("To‘lov summasi 0 dan katta bo‘lsin.")
    if "USD" in {order_code,payment_code} and rate<=0:
        raise ValueError("USD qatnashgan to‘lov uchun to‘lov kunidagi 1 USD kursini kiriting.")
    if order_code==payment_code:
        order_amount=received
    elif order_code=="USD" and payment_code=="UZS":
        order_amount=received/rate
    elif order_code=="UZS" and payment_code=="USD":
        order_amount=received*rate
    else:
        raise ValueError("Bu valyuta juftligi qo‘llab-quvvatlanmaydi.")
    if payment_code=="UZS":
        uzs_equivalent=received
    else:
        uzs_equivalent=received*rate
    return round(order_amount,2),round(uzs_equivalent,2),rate if rate>0 else 1.0


DB_NAME = os.environ.get("PHARM_ERP_DB", "pharm_mebel_erp_pro.db")
_APP_DIR = os.path.dirname(os.path.abspath(__file__))
CONSTRUCTOR_UPLOAD_DIR = os.path.join(_APP_DIR, "uploads", "konstruktor")
os.makedirs(CONSTRUCTOR_UPLOAD_DIR, exist_ok=True)
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024


def _get_or_create_secret_key():
    """PHARM_ERP_SECRET muhit o'zgaruvchisi bo'lmasa, bir martalik tasodifiy
    kalit yaratib, faylga saqlaydi (har safar ilova qayta ishga tushganda
    bir xil kalit ishlatilishi uchun, aks holda barcha sessiyalar uziladi)."""
    env_key = os.environ.get("PHARM_ERP_SECRET")
    if env_key:
        return env_key
    key_path = os.path.join(_APP_DIR, ".secret_key")
    if os.path.exists(key_path):
        with open(key_path, "r", encoding="utf-8") as f:
            saved = f.read().strip()
            if saved:
                return saved
    new_key = secrets.token_hex(32)
    try:
        with open(key_path, "w", encoding="utf-8") as f:
            f.write(new_key)
    except Exception:
        pass
    return new_key


app.secret_key = _get_or_create_secret_key()
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.environ.get("PHARM_ERP_HTTPS", "0") == "1",
)

# ---------- UCH TILLI TIZIM: O‘ZBEKCHA / RUSCHA / INGLIZCHA ----------
PHARM_I18N_WIDGET = '\n<!-- PHARM_MEBEL_3_LANGUAGE_SYSTEM_V2 -->\n<style id="pharm-i18n-style">\n#pharm-lang-box{position:fixed;top:10px;right:10px;z-index:99999;display:flex;gap:5px;padding:5px;border-radius:12px;background:rgba(255,255,255,.96);border:1px solid rgba(15,23,42,.14);box-shadow:0 8px 24px rgba(15,23,42,.18);font-family:Arial,sans-serif}\n#pharm-lang-box.pharm-lang-inline{position:static;box-shadow:none;background:rgba(255,255,255,.13);border-color:rgba(255,255,255,.28);padding:4px}\n#pharm-lang-box button{width:auto!important;margin:0!important;padding:7px 10px!important;border:0!important;border-radius:8px!important;background:#e2e8f0!important;color:#0f172a!important;font-size:12px!important;font-weight:800!important;line-height:1!important;cursor:pointer!important;white-space:nowrap!important}\n#pharm-lang-box button.active{background:#16a34a!important;color:#fff!important;box-shadow:0 3px 10px rgba(22,163,74,.28)!important}\n@media(max-width:700px){#pharm-lang-box:not(.pharm-lang-inline){top:auto;right:8px;bottom:8px}#pharm-lang-box button{padding:7px 8px!important;font-size:11px!important}}\n</style>\n<div id="pharm-lang-box" data-pharm-i18n-ignore="1" aria-label="Til tanlash">\n<button type="button" data-pharm-lang="uz">O‘zbekcha</button>\n<button type="button" data-pharm-lang="ru">Русский</button>\n<button type="button" data-pharm-lang="en">English</button>\n</div>\n<script id="pharm-i18n-script">\n(function(){\n\'use strict\';\nconst DICT={ru:{"Dastur ochilmoqda...":"Программа загружается...","Dastur ochilmoqda.":"Программа загружается.","Mebel360° dasturiga xush kelibsiz":"Добро пожаловать в программу Mebel360°","Faqat sizga biriktirilgan vazifalar va ruxsat berilgan bo‘limlar ochiladi.":"Откроются только назначенные вам задачи и разрешённые разделы.","Korxona, buyurtmalar, ishchilar va omborni yagona tizimda boshqaring.":"Управляйте предприятием, заказами, работниками и складом в единой системе.","22 yillik amaliy tajriba va yo‘l qo‘yilgan xatolardan olingan saboqlar asosida shakllangan tizim":"Система, созданная на основе 22-летнего практического опыта и уроков, извлечённых из допущенных ошибок","Dastur boshqaruv paneliga kirish uchun login va parolingizni kiriting.":"Введите логин и пароль для входа в панель управления.","Menejer, Konstruktor, Ishchi va Shofyor kabinetlari":"Кабинеты менеджера, конструктора, работника и водителя","Buyurtmalar va ishlab chiqarish nazorati":"Контроль заказов и производства","Ishchilar, buyurtmalar, ombor, ishlab chiqarish va moliya":"Работники, заказы, склад, производство и финансы","Ombor, xarajat va hisobotlar":"Склад, расходы и отчёты","Buyurtmalar va mijozlar":"Заказы и клиенты","Chizma, kroy va CNC":"Чертежи, раскрой и CNC","Chizma, kroy va CNC fayli":"Чертёж, раскрой и CNC-файл","Mebel360° boshqaruv tizimi":"Система управления Mebel360°","Mebel360° - Birinchi xavfsiz sozlash":"Mebel360° — первая безопасная настройка","🔐 Birinchi xavfsiz sozlash":"🔐 Первая безопасная настройка","Admin akkauntini yaratish":"Создать аккаунт администратора","Akkaunt yaratish yoki parolni yangilash":"Создать аккаунт или обновить пароль","Menejer va Konstruktor akkauntlari":"Аккаунты менеджера и конструктора","🔐 Menejer va Konstruktor akkauntlari":"🔐 Аккаунты менеджера и конструктора","Birinchi xavfsiz sozlash":"Первая безопасная настройка","Boshqa kabinetlar":"Другие кабинеты","Kabinetga kirish":"Войти в кабинет","Rahbar kirishi":"Вход руководителя","← Rahbar kirishi":"← Вход руководителя","Mebel360° — Rahbar kirishi":"Mebel360° — вход руководителя","Menejer kirishi":"Вход менеджера","Konstruktor kirishi":"Вход конструктора","Ishchi kirishi":"Вход работника","Shofyor kirishi":"Вход водителя","Ishchi ro‘yxatdan o‘tishi":"Регистрация работника","Shofyor ro‘yxatdan o‘tishi":"Регистрация водителя","Yangi shofyor — ro‘yxatdan o‘tish":"Новый водитель — зарегистрироваться","Akkauntim bor — kirish":"У меня есть аккаунт — войти","Ro‘yxatdan o‘tish":"Зарегистрироваться","Ro‘yxatdan o‘tdingiz":"Вы зарегистрировались","Yangi ro‘yxatdan o‘tish":"Новая регистрация","Telefonni tasdiqlash":"Подтвердить телефон","Kodni tasdiqlash":"Подтверждение кода","Kodni kiritish":"Введите код","Kod olish":"Получить код","Haqiqiy SMS xizmati ulanmaguncha shu koddan foydalaning.":"Используйте этот код до подключения настоящего SMS-сервиса.","Endi administrator akkauntingizni tasdiqlaydi. Tasdiqlangach login va parolingiz bilan kirasiz.":"Теперь администратор подтвердит ваш аккаунт. После подтверждения войдите с логином и паролем.","Menga biriktirilgan yuklar":"Назначенные мне доставки","Hozircha yuk biriktirilmagan.":"Доставки пока не назначены.","Lokatsiyani o‘zingiz xohlagan xaritada oching.":"Откройте геолокацию в удобной карте.","Sahifa har 15 soniyada avtomatik yangilanadi.":"Страница автоматически обновляется каждые 15 секунд.","Hali shofyor biriktirilmagan":"Водитель ещё не назначен","Hozircha topshiriq yo‘q. Holatingiz:":"Заданий пока нет. Ваш статус:","Hali akkaunt yaratilmagan.":"Аккаунт ещё не создан.","Hali fayl yuklanmagan.":"Файл ещё не загружен.","Buyurtmani tanlang.":"Выберите заказ.","Buyurtma yo‘q.":"Заказов нет.","Mijoz tasdiqlashi kutilmoqda":"Ожидается подтверждение клиента","Material tayyorlanmoqda":"Материал подготавливается","Chizma tayyorlanmoqda":"Чертёж подготавливается","Ishlab chiqarish bosqichlari":"Этапы производства","Buyurtma bosqichlari":"Этапы заказа","Buyurtma kuzatuvi":"Отслеживание заказа","Yetkazib berish holati":"Статус доставки","Yetkazishga tayyor":"Готово к доставке","Haydovchiga topshirildi":"Передано водителю","Yetkazib berildi":"Доставлено","Buyurtma yopildi":"Заказ закрыт","Yo‘lga chiqqan vaqt":"Время выезда","Yetib kelgan vaqt":"Время прибытия","Topshirilgan vaqt":"Время передачи","Yo‘lga chiqdim":"Я выехал","Yetib keldim":"Я прибыл","Yetkazib berdim":"Доставил","Avtomatik keldi-ketdi":"Автоматический приход-уход","Hozir keldi":"Пришёл сейчас","Hozir ketdi":"Ушёл сейчас","Qo‘lda kiritish":"Ввести вручную","Hozir ishlayotgan ishim":"Моя текущая работа","Boshlangan vaqt":"Время начала","Ishchi kabinetlarini boshqarish":"Управление кабинетами работников","Yangi topshiriq":"Новое задание","Topshiriq berish":"Назначить задание","Boshlanish rejasi":"План начала","Tugash rejasi":"План окончания","Ro‘yxatdan o‘tganlar":"Зарегистрированные","Shofyor login yaratish":"Создать логин водителя","Yuk biriktirish":"Назначить доставку","Biriktirilgan yuklar":"Назначенные доставки","Yetkazish navbati":"Очередь доставки","Qadoq soni":"Количество упаковок","Xarita tanlang":"Выберите карту","Mavjud akkauntlar":"Существующие аккаунты","Yangi akkaunt":"Новый аккаунт","Admin akkaunti":"Аккаунт администратора","Rahbar kabineti":"Кабинет руководителя","Menejer kabineti":"Кабинет менеджера","Konstruktor kabineti":"Кабинет конструктора","Ishchi kabineti":"Кабинет работника","Shofyor kabineti":"Кабинет водителя","Ishchi boshqaruvi":"Управление работниками","Shofyor boshqaruvi":"Управление водителями","PRO boshqaruv":"PRO управление","ERP bosh sahifa":"Главная ERP","Menejer / Konstruktor":"Менеджер / Конструктор","Shartnoma Word":"Договор Word","Shartnoma PDF":"Договор PDF","Toshkent vaqti":"Время Ташкента","Maxfiy hujjatlar":"Конфиденциальные документы","Bu ma’lumotlarni faqat rahbar ko‘radi.":"Эти данные видит только руководитель.","Pasport/ID seriya-raqami":"Серия и номер паспорта/ID","Favqulodda aloqa telefoni":"Экстренный телефон","Pasportni bergan tashkilot":"Орган, выдавший паспорт","Pasport berilgan sana":"Дата выдачи паспорта","Yashash manzili":"Адрес проживания","Tug‘ilgan sana":"Дата рождения","Ishga kirgan sana":"Дата приёма","Sifat balli (1-10)":"Оценка качества (1-10)","Tezlik balli (1-10)":"Оценка скорости (1-10)","Intizom balli (1-10)":"Оценка дисциплины (1-10)","Kechikish chegirmasi (%/kun)":"Скидка за задержку (%/день)","Maksimal chegirma %":"Максимальная скидка %","Taxminiy tayyor sana":"Ориентировочная дата готовности","Mas’ul xodim":"Ответственный сотрудник","Kechikish sababi":"Причина задержки","Lokatsiya havolasi":"Ссылка на геолокацию","Kafolat boshlanish":"Начало гарантии","Kafolat tugash":"Окончание гарантии","Kafolat sharti":"Условия гарантии","Kafolat muddati":"Срок гарантии","Oraliq to‘lov":"Промежуточный платёж","To‘lov usuli":"Способ оплаты","Bank orqali":"Банковский перевод","Plastik karta":"Банковская карта","Alohida haq":"Отдельная оплата","Kiritilmagan":"Не включено","Kiritilgan":"Включено","Katta mashina":"Крупная машина","Kira oladi":"Может заехать","Kira olmaydi":"Не может заехать","Ombor harakati":"Движение склада","Ombor qoldig‘i":"Остатки склада","Shofyor safari":"Поездка водителя","Masofa km":"Расстояние, км","Yoqilg‘i litr":"Топливо, л","Yo‘l xarajati":"Дорожные расходы","Moliyaviy xulosa":"Финансовый итог","Ishchi to‘lov":"Выплаты работникам","Sof foyda":"Чистая прибыль","Mijoz bahosi":"Оценка клиента","Mebel sifati (1-5)":"Качество мебели (1-5)","Muddat (1-5)":"Срок (1-5)","Muomala (1-5)":"Обслуживание (1-5)","Yetkazish (1-5)":"Доставка (1-5)","Montaj (1-5)":"Монтаж (1-5)","Bahoni saqlash":"Сохранить оценку","Qo‘shimcha ishlar":"Дополнительные работы","Qo‘shimcha summa":"Дополнительная сумма","Qo‘shimcha kun":"Дополнительные дни","Servis / kafolat murojaati":"Обращение по сервису / гарантии","Pullik xizmat":"Платная услуга","Servis sanasi":"Дата сервиса","Yetkazishni rejalash":"Планирование доставки","Fayl nomi/izohi":"Название файла / примечание","Faylni yuklash":"Загрузить файл","Holatni yangilash":"Обновить статус","Buyurtma holati":"Статус заказа","Buyurtmani saqlash":"Сохранить заказ","Yangi buyurtma":"Новый заказ","Umumiy narx":"Общая стоимость","Summa / qoldiq":"Сумма / остаток","Kod / mijoz":"Код / клиент","Taxminiy tayyor":"Ориентировочная готовность","JAMI BUYURTMA":"ВСЕГО ЗАКАЗОВ","FAOL BUYURTMALAR":"АКТИВНЫЕ ЗАКАЗЫ","JAMI SUMMA":"ОБЩАЯ СУММА","BUGUNGI SOAT":"ЧАСОВ СЕГОДНЯ","BUGUNGI ISH":"РАБОТА СЕГОДНЯ","BUGUNGI KM":"КМ СЕГОДНЯ","KAM QOLDIQ":"МАЛЫЙ ОСТАТОК","ISHCHILAR":"РАБОТНИКИ","ISHCHI TO‘LOV":"ВЫПЛАТЫ РАБОТНИКАМ","SOF FOYDA":"ЧИСТАЯ ПРИБЫЛЬ","KIRIM":"ДОХОД","XARAJAT":"РАСХОД","TUSHUM":"ПОСТУПЛЕНИЯ","BONUS":"БОНУС","Dushanba":"Понедельник","Seshanba":"Вторник","Chorshanba":"Среда","Payshanba":"Четверг","Juma":"Пятница","Shanba":"Суббота","Yakshanba":"Воскресенье","yanvar":"января","fevral":"февраля","mart":"марта","aprel":"апреля","may":"мая","iyun":"июня","iyul":"июля","avgust":"августа","sentabr":"сентября","oktabr":"октября","noyabr":"ноября","dekabr":"декабря","O‘zbekcha":"O‘zbekcha","Ruscha":"Русский","Inglizcha":"Английский","Akkaunt yaratish":"Создать аккаунт","Yangi parol":"Новый пароль","Yangi login":"Новый логин","Parolni qayta yozing":"Повторите пароль","Telefon raqamingiz":"Ваш номер телефона","Sinov kodi":"Тестовый код","Kirish sahifasi":"Страница входа","Topshiriq":"Задание","Tasdiqlangan":"Подтверждено","Tasdiqlash":"Подтвердить","Kutilmoqda":"Ожидается","Bloklash":"Заблокировать","Biriktirish":"Назначить","Qadoqlar":"Упаковки","Qadoq":"Упаковка","Orqaga":"Назад","Bosh sahifa":"Главная","Xush kelibsiz":"Добро пожаловать","Ko‘rsatish":"Показать","Kirish →":"Войти →","Kirish":"Войти","Chiqish":"Выйти","Login":"Логин","Parol":"Пароль","Saqlash":"Сохранить","Yangilash":"Обновить","O‘chirish":"Удалить","Qo‘shish":"Добавить","Yopish":"Закрыть","Hisoblash":"Рассчитать","Rejalash":"Запланировать","Qabul qilish":"Принять","Excel/CSV":"Excel/CSV","Backup":"Резервная копия","Rahbar":"Руководитель","Menejer":"Менеджер","Konstruktor":"Конструктор","Ishchi":"Работник","Shofyor":"Водитель","Haydovchi":"Водитель","Mijoz":"Клиент","Usta":"Мастер","Administrator":"Администратор","Akkaunt":"Аккаунт","Kabinet":"Кабинет","Ishchilar":"Работники","Buyurtmalar":"Заказы","Buyurtma":"Заказ","Ombor":"Склад","Ishlab chiqarish":"Производство","Korxona uchun":"Для предприятия","Korxonada":"На предприятии","Korxonaga":"На предприятие","Korxona":"Предприятие","Xarajat":"Расход","Hisobot":"Отчёт","Foyda":"Прибыль","Jarima":"Штраф","Bonus":"Бонус","Ta’til":"Отпуск","Servis":"Сервис","Kafolat":"Гарантия","Yetkazish":"Доставка","Montaj":"Монтаж","To‘lov":"Оплата","Avans":"Аванс","Qarz":"Долг","Naqd":"Наличные","Sana":"Дата","Boshlanish":"Начало","Tugash":"Окончание","Boshlangan":"Начато","Tugadi":"Завершено","Bajarildi":"Выполнено","Jarayonda":"В процессе","Jarayon":"Процесс","Tayyor":"Готово","Yangi":"Новый","Yopildi":"Закрыто","Yetkazildi":"Доставлено","Faol":"Активен","FAOL":"АКТИВНЫЕ","Holat":"Статус","Amal":"Действие","Izoh":"Примечание","Nomi":"Название","Ism":"Имя","Familiya":"Фамилия","Telefon":"Телефон","Manzil":"Адрес","Lavozim":"Должность","Kod":"Код","Kodi":"Код","Mahsulot":"Изделие","Material":"Материал","Rang":"Цвет","O‘lcham":"Размер","Soni":"Количество","Miqdor":"Количество","Summa":"Сумма","Narx":"Цена","Qoldiq":"Остаток","Birlik narxi":"Цена за единицу","Birlik":"Единица","Haq":"Оплата труда","Ish haqi":"Заработок","Oylik maosh":"Месячная зарплата","Kunlik stavka":"Дневная ставка","Oylik":"Месячная","Kunlik":"Дневная","Staj (yil)":"Стаж (лет)","Staj":"Стаж","JSHSHIR":"ПИНФЛ","Pasport/ID":"Паспорт/ID","Mas’ul":"Ответственный","Sababi":"Причина","Sabab":"Причина","Muammo":"Проблема","Turi":"Тип","Kategoriya":"Категория","Kirim":"Приход","Chiqim":"Расход","Keldi-ketdi":"Приход-уход","Keldi":"Пришёл","Ketdi":"Ушёл","Soat":"Часы","Kun":"Дни","O‘rin":"Место","Reyting":"Рейтинг","Jami/Reyting":"Итоги/Рейтинг","Jami":"Итого","Shu oy":"Этот месяц","Yo‘nalish":"Маршрут","Qayerdan":"Откуда","Qayerga":"Куда","Mashina":"Машина","Navbat":"Очередь","Mo‘ljal":"Ориентир","Qavat":"Этаж","Lift":"Лифт","Bor":"Есть","Yo‘q":"Нет","Favqulodda holat":"Чрезвычайная ситуация","Favqulodda":"Экстренное","Boshqa":"Другое","Uy uchun":"Для дома","Ofis va reklama":"Офис и реклама","Mijoz va buyurtma":"Клиент и заказ","Bonus/Ta’til":"Бонус/Отпуск","Tayyor mahsulot":"Готовая продукция","Mijoz PRO":"Клиент PRO","Servis/Kafolat":"Сервис/Гарантия","Ish natijasi":"Результат работы","Yangi ishchi":"Новый работник","Fayl":"Файл","Rasm":"Фото","Chek":"Чек","MDF":"МДФ","DSP":"ДСП","LDSP":"ЛДСП","DVP":"ДВП","HDF":"ХДФ","Akril":"Акрил","Fanera":"Фанера","Oyna":"Стекло","Shisha":"Стекло","Kromka":"Кромка","Furnitura":"Фурнитура","Bo‘yoq":"Краска","Yelim":"Клей","Silikon":"Силикон","Porolon":"Поролон","Mato":"Ткань","Profil":"Профиль","Kesish":"Резка","Bo‘yash":"Покраска","Yig‘ish":"Сборка","Teshish":"Сверление","Frezalash":"Фрезеровка","Sayqalash":"Шлифовка","Qadoqlash":"Упаковка","Razmer olish":"Замер","Chizma tayyorlash":"Подготовка чертежа","dona":"шт.","komplekt":"комплект","metr":"метр","soat":"час","kun":"день","daqiqa":"минута","safar":"поездка","loyiha":"проект","list":"лист","kg":"кг","m²":"м²"},en:{"Dastur ochilmoqda...":"Loading the system...","Dastur ochilmoqda.":"Loading the system.","Mebel360° dasturiga xush kelibsiz":"Welcome to the Mebel360° system","Faqat sizga biriktirilgan vazifalar va ruxsat berilgan bo‘limlar ochiladi.":"Only tasks assigned to you and permitted sections will be available.","Korxona, buyurtmalar, ishchilar va omborni yagona tizimda boshqaring.":"Manage the company, orders, workers and inventory in one system.","22 yillik amaliy tajriba va yo‘l qo‘yilgan xatolardan olingan saboqlar asosida shakllangan tizim":"A system shaped by 22 years of practical experience and lessons learned from mistakes","Dastur boshqaruv paneliga kirish uchun login va parolingizni kiriting.":"Enter your login and password to access the control panel.","Menejer, Konstruktor, Ishchi va Shofyor kabinetlari":"Manager, designer, worker and driver dashboards","Buyurtmalar va ishlab chiqarish nazorati":"Order and production control","Ishchilar, buyurtmalar, ombor, ishlab chiqarish va moliya":"Workers, orders, inventory, production and finance","Ombor, xarajat va hisobotlar":"Inventory, expenses and reports","Buyurtmalar va mijozlar":"Orders and customers","Chizma, kroy va CNC":"Drawings, cutting and CNC","Chizma, kroy va CNC fayli":"Drawing, cutting plan and CNC file","Mebel360° boshqaruv tizimi":"Mebel360° management system","Mebel360° - Birinchi xavfsiz sozlash":"Mebel360° — initial secure setup","🔐 Birinchi xavfsiz sozlash":"🔐 Initial secure setup","Admin akkauntini yaratish":"Create administrator account","Akkaunt yaratish yoki parolni yangilash":"Create an account or update the password","Menejer va Konstruktor akkauntlari":"Manager and designer accounts","🔐 Menejer va Konstruktor akkauntlari":"🔐 Manager and designer accounts","Birinchi xavfsiz sozlash":"Initial secure setup","Boshqa kabinetlar":"Other dashboards","Kabinetga kirish":"Open dashboard","Rahbar kirishi":"Owner sign-in","← Rahbar kirishi":"← Owner sign-in","Mebel360° — Rahbar kirishi":"Mebel360° — owner sign-in","Menejer kirishi":"Manager sign-in","Konstruktor kirishi":"Designer sign-in","Ishchi kirishi":"Worker sign-in","Shofyor kirishi":"Driver sign-in","Ishchi ro‘yxatdan o‘tishi":"Worker registration","Shofyor ro‘yxatdan o‘tishi":"Driver registration","Yangi shofyor — ro‘yxatdan o‘tish":"New driver — register","Akkauntim bor — kirish":"I already have an account — sign in","Ro‘yxatdan o‘tish":"Register","Ro‘yxatdan o‘tdingiz":"Registration completed","Yangi ro‘yxatdan o‘tish":"New registration","Telefonni tasdiqlash":"Verify phone","Kodni tasdiqlash":"Code verification","Kodni kiritish":"Enter code","Kod olish":"Get code","Haqiqiy SMS xizmati ulanmaguncha shu koddan foydalaning.":"Use this code until a real SMS service is connected.","Endi administrator akkauntingizni tasdiqlaydi. Tasdiqlangach login va parolingiz bilan kirasiz.":"The administrator will now approve your account. After approval, sign in with your login and password.","Menga biriktirilgan yuklar":"Deliveries assigned to me","Hozircha yuk biriktirilmagan.":"No deliveries have been assigned yet.","Lokatsiyani o‘zingiz xohlagan xaritada oching.":"Open the location in your preferred map.","Sahifa har 15 soniyada avtomatik yangilanadi.":"The page refreshes automatically every 15 seconds.","Hali shofyor biriktirilmagan":"No driver has been assigned yet","Hozircha topshiriq yo‘q. Holatingiz:":"No tasks yet. Your status:","Hali akkaunt yaratilmagan.":"No account has been created yet.","Hali fayl yuklanmagan.":"No file has been uploaded yet.","Buyurtmani tanlang.":"Select an order.","Buyurtma yo‘q.":"No orders.","Mijoz tasdiqlashi kutilmoqda":"Waiting for customer approval","Material tayyorlanmoqda":"Material in preparation","Chizma tayyorlanmoqda":"Drawing in preparation","Ishlab chiqarish bosqichlari":"Production stages","Buyurtma bosqichlari":"Order stages","Buyurtma kuzatuvi":"Order tracking","Yetkazib berish holati":"Delivery status","Yetkazishga tayyor":"Ready for delivery","Haydovchiga topshirildi":"Handed to the driver","Yetkazib berildi":"Delivered","Buyurtma yopildi":"Order closed","Yo‘lga chiqqan vaqt":"Departure time","Yetib kelgan vaqt":"Arrival time","Topshirilgan vaqt":"Handover time","Yo‘lga chiqdim":"I have departed","Yetib keldim":"I have arrived","Yetkazib berdim":"Delivered","Avtomatik keldi-ketdi":"Automatic attendance","Hozir keldi":"Clock in now","Hozir ketdi":"Clock out now","Qo‘lda kiritish":"Enter manually","Hozir ishlayotgan ishim":"My current task","Boshlangan vaqt":"Start time","Ishchi kabinetlarini boshqarish":"Manage worker accounts","Yangi topshiriq":"New task","Topshiriq berish":"Assign task","Boshlanish rejasi":"Planned start","Tugash rejasi":"Planned completion","Ro‘yxatdan o‘tganlar":"Registered users","Shofyor login yaratish":"Create driver login","Yuk biriktirish":"Assign delivery","Biriktirilgan yuklar":"Assigned deliveries","Yetkazish navbati":"Delivery queue","Qadoq soni":"Number of packages","Xarita tanlang":"Select a map","Mavjud akkauntlar":"Existing accounts","Yangi akkaunt":"New account","Admin akkaunti":"Administrator account","Rahbar kabineti":"Owner dashboard","Menejer kabineti":"Manager dashboard","Konstruktor kabineti":"Designer dashboard","Ishchi kabineti":"Worker dashboard","Shofyor kabineti":"Driver dashboard","Ishchi boshqaruvi":"Worker management","Shofyor boshqaruvi":"Driver management","PRO boshqaruv":"PRO management","ERP bosh sahifa":"ERP home","Menejer / Konstruktor":"Manager / Designer","Shartnoma Word":"Word contract","Shartnoma PDF":"PDF contract","Toshkent vaqti":"Tashkent time","Maxfiy hujjatlar":"Confidential documents","Bu ma’lumotlarni faqat rahbar ko‘radi.":"Only the owner can view this information.","Pasport/ID seriya-raqami":"Passport/ID series and number","Favqulodda aloqa telefoni":"Emergency contact phone","Pasportni bergan tashkilot":"Issuing authority","Pasport berilgan sana":"Passport issue date","Yashash manzili":"Home address","Tug‘ilgan sana":"Date of birth","Ishga kirgan sana":"Employment date","Sifat balli (1-10)":"Quality score (1-10)","Tezlik balli (1-10)":"Speed score (1-10)","Intizom balli (1-10)":"Discipline score (1-10)","Kechikish chegirmasi (%/kun)":"Delay discount (%/day)","Maksimal chegirma %":"Maximum discount %","Taxminiy tayyor sana":"Estimated completion date","Mas’ul xodim":"Responsible employee","Kechikish sababi":"Reason for delay","Lokatsiya havolasi":"Location link","Kafolat boshlanish":"Warranty start","Kafolat tugash":"Warranty end","Kafolat sharti":"Warranty terms","Kafolat muddati":"Warranty period","Oraliq to‘lov":"Intermediate payment","To‘lov usuli":"Payment method","Bank orqali":"Bank transfer","Plastik karta":"Bank card","Alohida haq":"Separate fee","Kiritilmagan":"Not included","Kiritilgan":"Included","Katta mashina":"Large vehicle","Kira oladi":"Can enter","Kira olmaydi":"Cannot enter","Ombor harakati":"Inventory movement","Ombor qoldig‘i":"Inventory balance","Shofyor safari":"Driver trip","Masofa km":"Distance, km","Yoqilg‘i litr":"Fuel, L","Yo‘l xarajati":"Travel expense","Moliyaviy xulosa":"Financial summary","Ishchi to‘lov":"Worker payments","Sof foyda":"Net profit","Mijoz bahosi":"Customer rating","Mebel sifati (1-5)":"Furniture quality (1-5)","Muddat (1-5)":"Timeline (1-5)","Muomala (1-5)":"Service (1-5)","Yetkazish (1-5)":"Delivery (1-5)","Montaj (1-5)":"Installation (1-5)","Bahoni saqlash":"Save rating","Qo‘shimcha ishlar":"Additional work","Qo‘shimcha summa":"Additional amount","Qo‘shimcha kun":"Additional days","Servis / kafolat murojaati":"Service / warranty request","Pullik xizmat":"Paid service","Servis sanasi":"Service date","Yetkazishni rejalash":"Plan delivery","Fayl nomi/izohi":"File name / note","Faylni yuklash":"Upload file","Holatni yangilash":"Update status","Buyurtma holati":"Order status","Buyurtmani saqlash":"Save order","Yangi buyurtma":"New order","Umumiy narx":"Total price","Summa / qoldiq":"Amount / balance","Kod / mijoz":"Code / customer","Taxminiy tayyor":"Estimated completion","JAMI BUYURTMA":"TOTAL ORDERS","FAOL BUYURTMALAR":"ACTIVE ORDERS","JAMI SUMMA":"TOTAL AMOUNT","BUGUNGI SOAT":"HOURS TODAY","BUGUNGI ISH":"WORK TODAY","BUGUNGI KM":"KM TODAY","KAM QOLDIQ":"LOW STOCK","ISHCHILAR":"WORKERS","ISHCHI TO‘LOV":"WORKER PAYMENTS","SOF FOYDA":"NET PROFIT","KIRIM":"INCOME","XARAJAT":"EXPENSE","TUSHUM":"REVENUE","BONUS":"BONUS","Dushanba":"Monday","Seshanba":"Tuesday","Chorshanba":"Wednesday","Payshanba":"Thursday","Juma":"Friday","Shanba":"Saturday","Yakshanba":"Sunday","yanvar":"January","fevral":"February","mart":"March","aprel":"April","may":"May","iyun":"June","iyul":"July","avgust":"August","sentabr":"September","oktabr":"October","noyabr":"November","dekabr":"December","O‘zbekcha":"O‘zbekcha","Ruscha":"Russian","Inglizcha":"English","Akkaunt yaratish":"Create account","Yangi parol":"New password","Yangi login":"New login","Parolni qayta yozing":"Repeat password","Telefon raqamingiz":"Your phone number","Sinov kodi":"Test code","Kirish sahifasi":"Sign-in page","Topshiriq":"Task","Tasdiqlangan":"Approved","Tasdiqlash":"Approve","Kutilmoqda":"Pending","Bloklash":"Block","Biriktirish":"Assign","Qadoqlar":"Packages","Qadoq":"Package","Orqaga":"Back","Bosh sahifa":"Home","Xush kelibsiz":"Welcome","Ko‘rsatish":"Show","Kirish →":"Sign in →","Kirish":"Sign in","Chiqish":"Sign out","Login":"Login","Parol":"Password","Saqlash":"Save","Yangilash":"Update","O‘chirish":"Delete","Qo‘shish":"Add","Yopish":"Close","Hisoblash":"Calculate","Rejalash":"Plan","Qabul qilish":"Accept","Excel/CSV":"Excel/CSV","Backup":"Backup","Rahbar":"Owner","Menejer":"Manager","Konstruktor":"Designer","Ishchi":"Worker","Shofyor":"Driver","Haydovchi":"Driver","Mijoz":"Customer","Usta":"Technician","Administrator":"Administrator","Akkaunt":"Account","Kabinet":"Dashboard","Ishchilar":"Workers","Buyurtmalar":"Orders","Buyurtma":"Order","Ombor":"Inventory","Ishlab chiqarish":"Production","Korxona uchun":"For the company","Korxonada":"At the company","Korxonaga":"To the company","Korxona":"Company","Xarajat":"Expense","Hisobot":"Report","Foyda":"Profit","Jarima":"Penalty","Bonus":"Bonus","Ta’til":"Leave","Servis":"Service","Kafolat":"Warranty","Yetkazish":"Delivery","Montaj":"Installation","To‘lov":"Payment","Avans":"Advance","Qarz":"Debt","Naqd":"Cash","Sana":"Date","Boshlanish":"Start","Tugash":"End","Boshlangan":"Started","Tugadi":"Completed","Bajarildi":"Completed","Jarayonda":"In progress","Jarayon":"Progress","Tayyor":"Ready","Yangi":"New","Yopildi":"Closed","Yetkazildi":"Delivered","Faol":"Active","FAOL":"ACTIVE","Holat":"Status","Amal":"Action","Izoh":"Note","Nomi":"Name","Ism":"First name","Familiya":"Last name","Telefon":"Phone","Manzil":"Address","Lavozim":"Position","Kod":"Code","Kodi":"Code","Mahsulot":"Product","Material":"Material","Rang":"Color","O‘lcham":"Size","Soni":"Quantity","Miqdor":"Quantity","Summa":"Amount","Narx":"Price","Qoldiq":"Balance","Birlik narxi":"Unit price","Birlik":"Unit","Haq":"Pay","Ish haqi":"Earnings","Oylik maosh":"Monthly salary","Kunlik stavka":"Daily rate","Oylik":"Monthly","Kunlik":"Daily","Staj (yil)":"Experience (years)","Staj":"Experience","JSHSHIR":"PINFL","Pasport/ID":"Passport/ID","Mas’ul":"Responsible","Sababi":"Reason","Sabab":"Reason","Muammo":"Problem","Turi":"Type","Kategoriya":"Category","Kirim":"Incoming","Chiqim":"Outgoing","Keldi-ketdi":"Attendance","Keldi":"Clock-in","Ketdi":"Clock-out","Soat":"Hours","Kun":"Days","O‘rin":"Rank","Reyting":"Rating","Jami/Reyting":"Totals/Rating","Jami":"Total","Shu oy":"This month","Yo‘nalish":"Route","Qayerdan":"From","Qayerga":"To","Mashina":"Vehicle","Navbat":"Queue","Mo‘ljal":"Landmark","Qavat":"Floor","Lift":"Elevator","Bor":"Available","Yo‘q":"No","Favqulodda holat":"Emergency","Favqulodda":"Emergency","Boshqa":"Other","Uy uchun":"For home","Ofis va reklama":"Office and advertising","Mijoz va buyurtma":"Customer and order","Bonus/Ta’til":"Bonus/Leave","Tayyor mahsulot":"Finished products","Mijoz PRO":"Customer PRO","Servis/Kafolat":"Service/Warranty","Ish natijasi":"Work results","Yangi ishchi":"New worker","Fayl":"File","Rasm":"Image","Chek":"Receipt","MDF":"MDF","DSP":"Chipboard","LDSP":"Laminated chipboard","DVP":"Fiberboard","HDF":"HDF","Akril":"Acrylic","Fanera":"Plywood","Oyna":"Glass","Shisha":"Glass","Kromka":"Edge band","Furnitura":"Hardware","Bo‘yoq":"Paint","Yelim":"Glue","Silikon":"Silicone","Porolon":"Foam","Mato":"Fabric","Profil":"Profile","Kesish":"Cutting","Bo‘yash":"Painting","Yig‘ish":"Assembly","Teshish":"Drilling","Frezalash":"Milling","Sayqalash":"Sanding","Qadoqlash":"Packaging","Razmer olish":"Measuring","Chizma tayyorlash":"Drawing preparation","dona":"pcs","komplekt":"set","metr":"meter","soat":"hour","kun":"day","daqiqa":"minute","safar":"trip","loyiha":"project","list":"sheet","kg":"kg","m²":"m²"}};\nconst STORE=\'pharm_language_v2\';\nconst ALLOWED=[\'uz\',\'ru\',\'en\'];\nlet current=localStorage.getItem(STORE)||\'uz\';\nif(!ALLOWED.includes(current))current=\'uz\';\nconst originals=new WeakMap();\nconst attrOriginals=new WeakMap();\nconst ordered={\n ru:Object.keys(DICT.ru).sort((a,b)=>b.length-a.length),\n en:Object.keys(DICT.en).sort((a,b)=>b.length-a.length)\n};\nlet busy=false;\nfunction ignored(node){\n const e=node&&node.nodeType===1?node:node&&node.parentElement;\n return !!(e&&e.closest&&e.closest(\'[data-pharm-i18n-ignore],script,style,code,pre,textarea\'));\n}\nfunction translateString(value,lang){\n if(lang===\'uz\'||value===null||value===undefined)return String(value??\'\');\n const raw=String(value);\n const lead=(raw.match(/^\\s*/)||[\'\'])[0];\n const tail=(raw.match(/\\s*$/)||[\'\'])[0];\n const core=raw.slice(lead.length,raw.length-tail.length);\n if(!core)return raw;\n const exact=DICT[lang][core];\n if(exact!==undefined)return lead+exact+tail;\n let out=core;\n const escapeRe=s=>s.replace(/[.*+?^${}()|[\\]\\\\]/g,\'\\\\$&\');\n for(const key of ordered[lang]){\n   if(key.length<2||!out.includes(key))continue;\n   const value=DICT[lang][key];\n   if(!key.includes(\' \')&&/^[\\p{L}\\p{N}ʻ‘’ʼ\'°²/%.-]+$/u.test(key)){\n     const re=new RegExp(\'(^|[^\\\\p{L}\\\\p{N}_])\'+escapeRe(key)+\'(?=$|[^\\\\p{L}\\\\p{N}_])\',\'gu\');\n     out=out.replace(re,(match,prefix)=>prefix+value);\n   }else{\n     out=out.split(key).join(value);\n   }\n }\n return lead+out+tail;\n}\nfunction textNode(node){\n if(!node||node.nodeType!==3||ignored(node)||!node.nodeValue.trim())return;\n const rec=originals.get(node);\n const source=(!rec||node.nodeValue!==rec.last)?node.nodeValue:rec.source;\n const next=translateString(source,current);\n originals.set(node,{source:source,last:next});\n if(node.nodeValue!==next){busy=true;node.nodeValue=next;busy=false;}\n}\nfunction attributes(el){\n if(!el||el.nodeType!==1||ignored(el))return;\n const rec=attrOriginals.get(el)||{};\n const attrs=[\'placeholder\',\'title\',\'aria-label\'];\n if(el.matches(\'input[type="button"],input[type="submit"],input[type="reset"]\'))attrs.push(\'value\');\n for(const name of attrs){\n   if(!el.hasAttribute(name))continue;\n   const currentValue=el.getAttribute(name)||\'\';\n   const old=rec[name];\n   const source=(!old||currentValue!==old.last)?currentValue:old.source;\n   const next=translateString(source,current);\n   rec[name]={source:source,last:next};\n   if(currentValue!==next)el.setAttribute(name,next);\n }\n attrOriginals.set(el,rec);\n}\nfunction walk(root){\n if(!root)return;\n if(root.nodeType===3){textNode(root);return;}\n if(root.nodeType!==1&&root.nodeType!==9&&root.nodeType!==11)return;\n if(root.nodeType===1){if(ignored(root))return;attributes(root);}\n const walker=document.createTreeWalker(root,NodeFilter.SHOW_ELEMENT|NodeFilter.SHOW_TEXT);\n let n;\n while((n=walker.nextNode())){if(n.nodeType===3)textNode(n);else attributes(n);}\n}\nconst originalTitle=document.title;\nfunction apply(lang){\n if(!ALLOWED.includes(lang))lang=\'uz\';\n current=lang;\n localStorage.setItem(STORE,lang);\n document.documentElement.lang=lang;\n walk(document.body);\n document.title=lang===\'uz\'?originalTitle:translateString(originalTitle,lang);\n document.querySelectorAll(\'[data-pharm-lang]\').forEach(b=>b.classList.toggle(\'active\',b.dataset.pharmLang===lang));\n window.dispatchEvent(new CustomEvent(\'pharm-language-change\',{detail:{language:lang}}));\n}\nfunction start(){\n const box=document.getElementById(\'pharm-lang-box\');\n const actions=document.querySelector(\'.header-actions\');\n if(box&&actions){box.classList.add(\'pharm-lang-inline\');actions.prepend(box);}\n document.querySelectorAll(\'[data-pharm-lang]\').forEach(b=>b.addEventListener(\'click\',()=>apply(b.dataset.pharmLang)));\n const nativeAlert=window.alert.bind(window);\n const nativeConfirm=window.confirm.bind(window);\n const nativePrompt=window.prompt.bind(window);\n window.alert=(m)=>nativeAlert(translateString(m,current));\n window.confirm=(m)=>nativeConfirm(translateString(m,current));\n window.prompt=(m,d)=>nativePrompt(translateString(m,current),d);\n apply(current);\n const observer=new MutationObserver(records=>{\n   if(busy)return;\n   for(const r of records){\n     if(r.type===\'characterData\')textNode(r.target);\n     else for(const n of r.addedNodes)walk(n);\n   }\n });\n observer.observe(document.body,{subtree:true,childList:true,characterData:true});\n}\nif(document.readyState===\'loading\')document.addEventListener(\'DOMContentLoaded\',start);else start();\n})();\n</script>\n'


# ---------- MEBEL360° PWA / TELEFON ILOVASI ----------
MEBEL360_PWA_HEAD = r"""
<!-- MEBEL360_PWA_V5 -->
<link rel="manifest" href="/manifest.webmanifest?v=6">
<meta name="theme-color" content="#0757a6">
<meta name="application-name" content="Mebel360°">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="default">
<meta name="apple-mobile-web-app-title" content="Mebel360°">
<link rel="apple-touch-icon" href="/static/icons/icon-192.png?v=5">
<link rel="icon" type="image/png" sizes="192x192" href="/static/icons/icon-192.png?v=5">
"""

MEBEL360_PWA_BODY = r"""
<!-- MEBEL360_PWA_INSTALL_V5 -->
<style>
#mebel360-install-btn{display:none;position:fixed;left:12px;bottom:12px;z-index:99998;border:0;border-radius:13px;padding:11px 15px;background:linear-gradient(135deg,#0757a6,#10a952);color:#fff;font-weight:900;box-shadow:0 9px 26px rgba(2,34,79,.30);cursor:pointer}
@media(max-width:700px){#mebel360-install-btn{font-size:13px;padding:10px 12px}}
</style>
<button id="mebel360-install-btn" type="button">📲 Telefonga o‘rnatish</button>
<script>
(function(){
  'use strict';
  if('serviceWorker' in navigator){
    window.addEventListener('load', function(){
      navigator.serviceWorker.register('/service-worker.js?v=6', {scope:'/'}).catch(function(){});
    });
  }
  var deferredPrompt = null;
  var btn = document.getElementById('mebel360-install-btn');
  window.addEventListener('beforeinstallprompt', function(e){
    e.preventDefault();
    deferredPrompt = e;
    if(btn) btn.style.display = 'block';
  });
  if(btn){
    btn.addEventListener('click', async function(){
      if(!deferredPrompt) return;
      deferredPrompt.prompt();
      try{ await deferredPrompt.userChoice; }catch(_e){}
      deferredPrompt = null;
      btn.style.display = 'none';
    });
  }
  window.addEventListener('appinstalled', function(){ if(btn) btn.style.display='none'; });
})();
</script>
"""

@app.route('/manifest.webmanifest')
def mebel360_manifest():
    payload = {
        'id': '/',
        'name': 'Mebel360°',
        'short_name': 'Mebel360',
        'description': 'Mebel korxonasini 360° boshqarish: buyurtmalar, ishchilar, keldi-ketdi, ombor va moliya.',
        'lang': 'uz',
        'start_url': '/?source=pwa',
        'scope': '/',
        'display': 'standalone',
        'display_override': ['window-controls-overlay', 'standalone', 'minimal-ui'],
        'orientation': 'any',
        'background_color': '#eef3f8',
        'theme_color': '#0757a6',
        'categories': ['business', 'productivity'],
        'icons': [
            {'src':'/static/icons/icon-192.png?v=5','sizes':'192x192','type':'image/png','purpose':'any'},
            {'src':'/static/icons/icon-512.png?v=5','sizes':'512x512','type':'image/png','purpose':'any'},
            {'src':'/static/icons/icon-maskable-512.png?v=5','sizes':'512x512','type':'image/png','purpose':'maskable'}
        ]
    }
    import json
    return Response(json.dumps(payload, ensure_ascii=False), mimetype='application/manifest+json', headers={'Cache-Control':'no-cache, no-store, must-revalidate'})

@app.route('/service-worker.js')
def mebel360_service_worker():
    js = r"""
const CACHE='mebel360-static-v6';
const STATIC=['/offline.html','/static/mebel360-logo.png?v=20260722','/static/icons/icon-192.png?v=5','/static/icons/icon-512.png?v=5'];
self.addEventListener('install',event=>{event.waitUntil(caches.open(CACHE).then(c=>c.addAll(STATIC)).then(()=>self.skipWaiting()))});
self.addEventListener('activate',event=>{event.waitUntil(caches.keys().then(keys=>Promise.all(keys.filter(k=>k!==CACHE).map(k=>caches.delete(k)))).then(()=>self.clients.claim()))});
self.addEventListener('fetch',event=>{
  const req=event.request;
  if(req.method!=='GET') return;
  const url=new URL(req.url);
  if(url.origin!==location.origin) return;
  if(url.pathname.startsWith('/api/') || url.pathname==='/service-worker.js' || url.pathname==='/manifest.webmanifest') return;
  if(req.mode==='navigate'){
    event.respondWith(fetch(req).catch(()=>caches.match('/offline.html')));
    return;
  }
  if(url.pathname.startsWith('/static/')){
    event.respondWith(caches.match(req).then(hit=>hit||fetch(req).then(res=>{const copy=res.clone();caches.open(CACHE).then(c=>c.put(req,copy));return res;})));
  }
});
"""
    return Response(js, mimetype='application/javascript', headers={'Cache-Control':'no-cache, no-store, must-revalidate','Service-Worker-Allowed':'/'})

@app.route('/offline.html')
def mebel360_offline():
    return Response("""<!doctype html><html lang='uz'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Mebel360° — Internet yo‘q</title><style>body{margin:0;min-height:100vh;display:grid;place-items:center;font-family:Arial;background:#eef3f8;color:#0f1b33}.box{max-width:520px;margin:18px;padding:28px;text-align:center;background:#fff;border-radius:22px;box-shadow:0 18px 50px #0f172a22}.box img{max-width:260px;width:75%}button{border:0;border-radius:12px;padding:12px 17px;background:#0757a6;color:#fff;font-weight:800}</style></head><body><div class='box'><img src='/static/mebel360-logo.png'><h1>Mebel360°</h1><p>Internet aloqasi yo‘q. Internet qaytgach sahifani yangilang.</p><button onclick='location.reload()'>Qayta urinish</button></div></body></html>""", mimetype='text/html')

@app.route('/api/health')
def mebel360_health():
    return jsonify({'status':'ok','app':'Mebel360°','time':_tashkent_now().isoformat()})


@app.after_request
def _pharm_inject_language_selector(response):
    """Barcha HTML sahifalarga til tanlash va Mebel360° PWA qismlarini qo‘shadi."""
    try:
        content_type = response.headers.get('Content-Type', '')
        if response.status_code == 200 and 'text/html' in content_type.lower() and not response.direct_passthrough:
            page = response.get_data(as_text=True)
            low = page.lower()
            if '</head>' in low and 'MEBEL360_PWA_V5' not in page:
                pos = low.rfind('</head>')
                page = page[:pos] + MEBEL360_PWA_HEAD + page[pos:]
                low = page.lower()
            additions = ''
            if 'PHARM_MEBEL_3_LANGUAGE_SYSTEM_V2' not in page:
                additions += PHARM_I18N_WIDGET
            if 'MEBEL360_PWA_INSTALL_V5' not in page:
                additions += MEBEL360_PWA_BODY
            if additions and '</body>' in low:
                pos = low.rfind('</body>')
                page = page[:pos] + additions + page[pos:]
            response.set_data(page)
    except Exception:
        pass
    return response

def get_db():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db():
    conn = get_db()
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS ishchilar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ism TEXT NOT NULL,
        familiya TEXT DEFAULT '',
        telefon TEXT DEFAULT '',
        lavozim TEXT DEFAULT '',
        ishga_kirgan_sana TEXT DEFAULT '',
        staj_yil REAL DEFAULT 0,
        kunlik_stavka REAL DEFAULT 0,
        oylik_maosh REAL DEFAULT 0,
        sifat_ball REAL DEFAULT 5,
        tezlik_ball REAL DEFAULT 5,
        intizom_ball REAL DEFAULT 5,
        izoh TEXT DEFAULT '',
        faol INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS ish_turlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nomi TEXT UNIQUE NOT NULL,
        kategoriya TEXT DEFAULT 'Ish',
        birlik TEXT DEFAULT 'dona',
        standart_narx REAL DEFAULT 0,
        faol INTEGER DEFAULT 1
    );

    CREATE TABLE IF NOT EXISTS keldi_ketdi (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        keldi_vaqti TEXT NOT NULL,
        ketdi_vaqti TEXT NOT NULL,
        ish_soatlari REAL DEFAULT 0,
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS ish_natijalari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        ish_turi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        birlik_narxi REAL DEFAULT 0,
        jami_haq REAL DEFAULT 0,
        buyurtma_kodi TEXT DEFAULT '',
        izoh TEXT DEFAULT '',
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE,
        FOREIGN KEY (ish_turi_id) REFERENCES ish_turlari(id)
    );

    CREATE TABLE IF NOT EXISTS tolovlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        turi TEXT DEFAULT 'Avans',
        tavsifi TEXT DEFAULT '',
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS jarimalar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        sababi TEXT DEFAULT '',
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtmalar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        kod TEXT UNIQUE NOT NULL,
        mijoz TEXT NOT NULL,
        telefon TEXT DEFAULT '',
        manzil TEXT DEFAULT '',
        mahsulot TEXT DEFAULT '',
        umumiy_narx REAL DEFAULT 0,
        oldindan_tolov REAL DEFAULT 0,
        tugash_sana TEXT DEFAULT '',
        holat TEXT DEFAULT 'Yangi',
        izoh TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS shartnoma_versiyalari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        versiya INTEGER DEFAULT 1,
        docx_fayl TEXT DEFAULT '',
        pdf_fayl TEXT DEFAULT '',
        yaratilgan_vaqt TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtma_bosqichlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        bosqich TEXT NOT NULL,
        bajarildi INTEGER DEFAULT 0,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE,
        UNIQUE(buyurtma_id, bosqich)
    );

    CREATE TABLE IF NOT EXISTS ombor (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nomi TEXT UNIQUE NOT NULL,
        kategoriya TEXT DEFAULT '',
        birlik TEXT DEFAULT 'dona',
        qoldiq REAL DEFAULT 0,
        min_qoldiq REAL DEFAULT 0,
        narx REAL DEFAULT 0
    );

    CREATE TABLE IF NOT EXISTS ombor_harakat (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        material_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        turi TEXT NOT NULL,
        miqdor REAL NOT NULL,
        buyurtma_kodi TEXT DEFAULT '',
        izoh TEXT DEFAULT '',
        FOREIGN KEY (material_id) REFERENCES ombor(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS safarlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        mashina TEXT DEFAULT '',
        qayerdan TEXT DEFAULT '',
        qayerga TEXT DEFAULT '',
        masofa_km REAL DEFAULT 0,
        sabab TEXT DEFAULT '',
        yonilgi REAL DEFAULT 0,
        xarajat REAL DEFAULT 0,
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS shofyor_akkauntlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER UNIQUE NOT NULL,
        login TEXT UNIQUE NOT NULL,
        parol_hash TEXT NOT NULL,
        faol INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );


    CREATE TABLE IF NOT EXISTS xarajatlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        sana TEXT NOT NULL,
        kategoriya TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        tavsifi TEXT DEFAULT '',
        buyurtma_kodi TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS bonuslar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        sababi TEXT DEFAULT '',
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS ishchi_holatlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        turi TEXT NOT NULL,
        izoh TEXT DEFAULT '',
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtma_tolovlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        sana TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        turi TEXT DEFAULT 'To‘lov',
        izoh TEXT DEFAULT '',
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS tayyor_mahsulot (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nomi TEXT NOT NULL,
        kodi TEXT DEFAULT '',
        rang TEXT DEFAULT '',
        miqdor REAL DEFAULT 0,
        birlik TEXT DEFAULT 'dona',
        narx REAL DEFAULT 0,
        izoh TEXT DEFAULT ''
    );


    CREATE TABLE IF NOT EXISTS mijozlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fio TEXT NOT NULL,
        telefon TEXT DEFAULT '',
        pasport_id TEXT DEFAULT '',
        manzil TEXT DEFAULT '',
        rozilik_reklama INTEGER DEFAULT 0,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS mijoz_akkauntlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mijoz_id INTEGER,
        telefon TEXT NOT NULL,
        parol_hash TEXT DEFAULT '',
        sms_tasdiq INTEGER DEFAULT 0,
        faol INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (mijoz_id) REFERENCES mijozlar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtma_bosqich_hodisalari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        bosqich TEXT NOT NULL,
        boshlandi TEXT DEFAULT '',
        tugadi TEXT DEFAULT '',
        ishchi_id INTEGER,
        izoh TEXT DEFAULT '',
        media_havola TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE,
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE SET NULL
    );

    CREATE TABLE IF NOT EXISTS buyurtma_media (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        turi TEXT DEFAULT 'Rasm',
        havola TEXT NOT NULL,
        izoh TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtma_tasdiqlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        turi TEXT NOT NULL,
        holat TEXT DEFAULT 'Kutilmoqda',
        izoh TEXT DEFAULT '',
        tasdiqlagan TEXT DEFAULT '',
        tasdiqlangan_vaqt TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtma_muddat_tarixi (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        eski_sana TEXT DEFAULT '',
        eski_vaqt TEXT DEFAULT '',
        yangi_sana TEXT DEFAULT '',
        yangi_vaqt TEXT DEFAULT '',
        sabab TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS keshbek_harakatlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mijoz_telefon TEXT NOT NULL,
        buyurtma_id INTEGER,
        turi TEXT NOT NULL,
        miqdor REAL DEFAULT 0,
        amal_muddati TEXT DEFAULT '',
        izoh TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE SET NULL
    );

    CREATE TABLE IF NOT EXISTS baholar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        sifat INTEGER DEFAULT 5,
        muddat INTEGER DEFAULT 5,
        muomala INTEGER DEFAULT 5,
        yetkazish INTEGER DEFAULT 5,
        montaj INTEGER DEFAULT 5,
        umumiy INTEGER DEFAULT 5,
        izoh TEXT DEFAULT '',
        rasm_havola TEXT DEFAULT '',
        reklama_ruxsat INTEGER DEFAULT 0,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS servis_murojaatlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        turi TEXT DEFAULT 'Servis',
        muammo TEXT NOT NULL,
        media_havola TEXT DEFAULT '',
        holat TEXT DEFAULT 'Qabul qilindi',
        usta_id INTEGER,
        servis_sana TEXT DEFAULT '',
        pullik INTEGER DEFAULT 0,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE,
        FOREIGN KEY (usta_id) REFERENCES ishchilar(id) ON DELETE SET NULL
    );

    CREATE TABLE IF NOT EXISTS kafolatlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER UNIQUE NOT NULL,
        boshlanish TEXT DEFAULT '',
        tugash TEXT DEFAULT '',
        shartlar TEXT DEFAULT '',
        qr_token TEXT DEFAULT '',
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS qoshimcha_ishlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        nomi TEXT NOT NULL,
        summa REAL DEFAULT 0,
        qoshimcha_kun INTEGER DEFAULT 0,
        mijoz_tasdiq INTEGER DEFAULT 0,
        kim_kiritdi TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS ombor_rezervlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        material_id INTEGER NOT NULL,
        miqdor REAL DEFAULT 0,
        holat TEXT DEFAULT 'Rezerv',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE,
        FOREIGN KEY (material_id) REFERENCES ombor(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS buyurtma_tannarxi (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER UNIQUE NOT NULL,
        material REAL DEFAULT 0,
        ishchi REAL DEFAULT 0,
        boyoq REAL DEFAULT 0,
        oyna REAL DEFAULT 0,
        furnitura REAL DEFAULT 0,
        transport REAL DEFAULT 0,
        elektr REAL DEFAULT 0,
        tashqi_xizmat REAL DEFAULT 0,
        boshqa REAL DEFAULT 0,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS tizim_xabarlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        foydalanuvchi_turi TEXT DEFAULT '',
        foydalanuvchi_id INTEGER,
        buyurtma_id INTEGER,
        mavzu TEXT NOT NULL,
        matn TEXT NOT NULL,
        kanal TEXT DEFAULT 'Sayt',
        oqildi INTEGER DEFAULT 0,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS foydalanuvchi_rollari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nomi TEXT UNIQUE NOT NULL,
        huquqlar TEXT DEFAULT ''
    );


    CREATE TABLE IF NOT EXISTS admin_akkauntlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        login TEXT UNIQUE NOT NULL,
        parol_hash TEXT NOT NULL,
        faol INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS audit_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        sana_vaqt TEXT DEFAULT CURRENT_TIMESTAMP,
        amal TEXT NOT NULL,
        tafsilot TEXT DEFAULT ''
    );

    CREATE TABLE IF NOT EXISTS xodim_akkauntlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ism TEXT NOT NULL,
        login TEXT UNIQUE NOT NULL,
        parol_hash TEXT NOT NULL,
        rol TEXT NOT NULL,
        faol INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS buyurtma_hujjatlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        nomi TEXT NOT NULL,
        fayl_nomi TEXT NOT NULL,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS ishchi_akkauntlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER UNIQUE,
        telefon TEXT UNIQUE NOT NULL,
        login TEXT UNIQUE,
        parol_hash TEXT DEFAULT '',
        tasdiqlangan INTEGER DEFAULT 0,
        admin_tasdiq INTEGER DEFAULT 0,
        faol INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE SET NULL
    );

    CREATE TABLE IF NOT EXISTS ishchi_otp (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        telefon TEXT NOT NULL,
        kod_hash TEXT NOT NULL,
        muddati TEXT NOT NULL,
        ishlatildi INTEGER DEFAULT 0,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS ishchi_topshiriqlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ishchi_id INTEGER NOT NULL,
        buyurtma_kodi TEXT DEFAULT '',
        ish_turi TEXT NOT NULL,
        tavsif TEXT DEFAULT '',
        holat TEXT DEFAULT 'Yangi',
        progress INTEGER DEFAULT 0,
        sana TEXT NOT NULL,
        tugash_sana TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (ishchi_id) REFERENCES ishchilar(id) ON DELETE CASCADE
    );

    CREATE TABLE IF NOT EXISTS mijoz_baholari (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        sifat INTEGER DEFAULT 5, muddat INTEGER DEFAULT 5, muomala INTEGER DEFAULT 5,
        yetkazish INTEGER DEFAULT 5, montaj INTEGER DEFAULT 5, izoh TEXT DEFAULT '',
        reklama_ruxsat INTEGER DEFAULT 0, created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS servis_murojaatlari (
        id INTEGER PRIMARY KEY AUTOINCREMENT, buyurtma_id INTEGER NOT NULL,
        turi TEXT DEFAULT 'Servis', muammo TEXT NOT NULL, holat TEXT DEFAULT 'Qabul qilindi',
        servis_sana TEXT DEFAULT '', usta TEXT DEFAULT '', izoh TEXT DEFAULT '',
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS qoshimcha_ishlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT, buyurtma_id INTEGER NOT NULL,
        nomi TEXT NOT NULL, summa REAL DEFAULT 0, qoshimcha_kun INTEGER DEFAULT 0,
        tasdiq INTEGER DEFAULT 0, izoh TEXT DEFAULT '', created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE
    );
    CREATE TABLE IF NOT EXISTS yetkazishlar (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        buyurtma_id INTEGER NOT NULL,
        haydovchi_id INTEGER,
        sana TEXT DEFAULT '',
        navbat INTEGER DEFAULT 1,
        mashina TEXT DEFAULT '',
        holat TEXT DEFAULT 'Rejalashtirilgan',
        yolga_chiqdi TEXT DEFAULT '',
        yetib_keldi TEXT DEFAULT '',
        topshirildi TEXT DEFAULT '',
        yetkazildi TEXT DEFAULT '',
        benzin REAL DEFAULT 0,
        yol_xarajati REAL DEFAULT 0,
        izoh TEXT DEFAULT '',
        qadoq_soni INTEGER DEFAULT 1,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE,
        FOREIGN KEY (haydovchi_id) REFERENCES ishchilar(id) ON DELETE SET NULL
    );
    """)

    roles=[
        ("Rahbar","all"),("Administrator","orders,customers,workers"),
        ("Menejer","orders,customers,payments"),("Konstruktor","drawings,stages,files"),
        ("Hisobchi","finance,payments,expenses"),("Omborchi","stock"),
        ("Ishchi","own_tasks"),("Bo‘lim boshlig‘i","department"),
        ("Shofyor","own_deliveries"),("Mijoz","own_orders")
    ]
    conn.executemany("INSERT OR IGNORE INTO foydalanuvchi_rollari(nomi,huquqlar) VALUES(?,?)",roles)

    ishlar = [
        ("Kesish","Ishlab chiqarish","dona"),("Bo‘yash","Ishlab chiqarish","dona"),
        ("Yig‘ish","Ishlab chiqarish","dona"),("Roverda ishlash","Stanok","soat"),
        ("Razmer olish","Xizmat","buyurtma"),("Kromka urish","Ishlab chiqarish","metr"),
        ("Sayqalash","Ishlab chiqarish","dona"),("Teshish","Ishlab chiqarish","dona"),
        ("Lazer","Stanok","soat"),("Korxonada yig‘ish","Ishlab chiqarish","dona"),
        ("Korxona ishlari","Umumiy","soat"),("Yangi loyihalar","Loyiha","loyiha"),
        ("Frezalash","Stanok","soat"),("Yumshoq mebel","Ishlab chiqarish","dona"),
        ("Shofyorlik","Transport","safar"),("Ish vaqtida hordiq","Tanaffus","daqiqa"),
        ("Bozorga tushish","Xizmat","safar"),("Qadoqlash","Ishlab chiqarish","dona"),
        ("Oyna o‘rnatish","Ishlab chiqarish","dona"),("Oyna teshish","Ishlab chiqarish","dona"),
        ("Tumba yig‘ish","Mebel turi","dona"),("Oyoq kiyim shkafi yig‘ish","Mebel turi","dona"),
        ("Kravot yig‘ish","Mebel turi","dona"),("Shkaf yig‘ish","Mebel turi","dona"),
        ("Kuxniy yig‘ish","Mebel turi","dona"),("Kamod yig‘ish","Mebel turi","dona"),
        ("Buyurtmani raspildan olib kelish","Transport","safar"),
        ("Abed","Tanaffus","daqiqa"),("Tanaffus","Tanaffus","daqiqa"),
        ("Namoz vaqti","Tanaffus","daqiqa")
    ]
    conn.executemany(
        "INSERT OR IGNORE INTO ish_turlari(nomi,kategoriya,birlik) VALUES(?,?,?)", ishlar
    )


    # PHARM_MEBEL_KORXONA_MIGRATION_V3_SAFE
    # Eski bazadagi "Sex" va "Seh" yozuvlarini ma'lumot yo'qotmasdan
    # "Korxona" yozuviga o'tkazadi. Takroriy ish turlari xavfsiz birlashtiriladi.
    def _korxona_nomi(value):
        result = str(value or "")
        pairs = (
            ("Sex uchun", "Korxona uchun"), ("sex uchun", "korxona uchun"), ("SEX UCHUN", "KORXONA UCHUN"),
            ("Seh uchun", "Korxona uchun"), ("seh uchun", "korxona uchun"), ("SEH UCHUN", "KORXONA UCHUN"),
            ("Sexda", "Korxonada"), ("sexda", "korxonada"), ("SEXDA", "KORXONADA"),
            ("Sehda", "Korxonada"), ("sehda", "korxonada"), ("SEHDA", "KORXONADA"),
            ("Sexga", "Korxonaga"), ("sexga", "korxonaga"), ("SEXGA", "KORXONAGA"),
            ("Sehga", "Korxonaga"), ("sehga", "korxonaga"), ("SEHGA", "KORXONAGA"),
            ("Sexdan", "Korxonadan"), ("sexdan", "korxonadan"), ("SEXDAN", "KORXONADAN"),
            ("Sehdan", "Korxonadan"), ("sehdan", "korxonadan"), ("SEHDAN", "KORXONADAN"),
            ("Sexning", "Korxonaning"), ("sexning", "korxonaning"), ("SEXNING", "KORXONANING"),
            ("Sehning", "Korxonaning"), ("sehning", "korxonaning"), ("SEHNING", "KORXONANING"),
            ("Sexni", "Korxonani"), ("sexni", "korxonani"), ("SEXNI", "KORXONANI"),
            ("Sehni", "Korxonani"), ("sehni", "korxonani"), ("SEHNI", "KORXONANI"),
            ("Sexlar", "Korxonalar"), ("sexlar", "korxonalar"), ("SEXLAR", "KORXONALAR"),
            ("Sehlar", "Korxonalar"), ("sehlar", "korxonalar"), ("SEHLAR", "KORXONALAR"),
            ("Sex", "Korxona"), ("sex", "korxona"), ("SEX", "KORXONA"),
            ("Seh", "Korxona"), ("seh", "korxona"), ("SEH", "KORXONA"),
        )
        for old, new in pairs:
            result = result.replace(old, new)
        return result

    old_work_types = conn.execute("""
        SELECT id, nomi FROM ish_turlari
        WHERE nomi LIKE '%Sex%' OR nomi LIKE '%sex%' OR nomi LIKE '%SEX%'
           OR nomi LIKE '%Seh%' OR nomi LIKE '%seh%' OR nomi LIKE '%SEH%'
    """).fetchall()
    for old_work in old_work_types:
        old_id = int(old_work["id"])
        new_name = _korxona_nomi(old_work["nomi"])
        existing_work = conn.execute(
            "SELECT id FROM ish_turlari WHERE nomi=? AND id<>?",
            (new_name, old_id),
        ).fetchone()
        if existing_work:
            new_id = int(existing_work["id"])
            conn.execute(
                "UPDATE ish_natijalari SET ish_turi_id=? WHERE ish_turi_id=?",
                (new_id, old_id),
            )
            conn.execute("DELETE FROM ish_turlari WHERE id=?", (old_id,))
        else:
            conn.execute("UPDATE ish_turlari SET nomi=? WHERE id=?", (new_name, old_id))

    old_stages = conn.execute("""
        SELECT id, buyurtma_id, bosqich, bajarildi
        FROM buyurtma_bosqichlari
        WHERE bosqich LIKE '%Sex%' OR bosqich LIKE '%sex%' OR bosqich LIKE '%SEX%'
           OR bosqich LIKE '%Seh%' OR bosqich LIKE '%seh%' OR bosqich LIKE '%SEH%'
    """).fetchall()
    for old_stage in old_stages:
        old_stage_id = int(old_stage["id"])
        order_id = int(old_stage["buyurtma_id"])
        new_stage_name = _korxona_nomi(old_stage["bosqich"])
        existing_stage = conn.execute("""
            SELECT id, bajarildi FROM buyurtma_bosqichlari
            WHERE buyurtma_id=? AND bosqich=? AND id<>?
        """, (order_id, new_stage_name, old_stage_id)).fetchone()
        if existing_stage:
            merged_done = max(
                int(existing_stage["bajarildi"] or 0),
                int(old_stage["bajarildi"] or 0),
            )
            conn.execute(
                "UPDATE buyurtma_bosqichlari SET bajarildi=? WHERE id=?",
                (merged_done, int(existing_stage["id"])),
            )
            conn.execute("DELETE FROM buyurtma_bosqichlari WHERE id=?", (old_stage_id,))
        else:
            conn.execute(
                "UPDATE buyurtma_bosqichlari SET bosqich=? WHERE id=?",
                (new_stage_name, old_stage_id),
            )

    # Matn ko'rinishida saqlangan eski yozuvlar.
    for table_name, column_name in (
        ("ishchi_topshiriqlari", "ish_turi"),
        ("buyurtma_bosqich_hodisalari", "bosqich"),
        ("xarajatlar", "kategoriya"),
        ("xarajatlar", "xarajat_nomi"),
        ("buyurtmalar", "kechikish_turi"),
    ):
        try:
            rows = conn.execute(
                f'SELECT id, "{column_name}" FROM "{table_name}" '
                f'WHERE "{column_name}" IS NOT NULL'
            ).fetchall()
            for row in rows:
                current_value = row[column_name]
                new_value = _korxona_nomi(current_value)
                if new_value != current_value:
                    conn.execute(
                        f'UPDATE "{table_name}" SET "{column_name}"=? WHERE id=?',
                        (new_value, int(row["id"])),
                    )
        except sqlite3.OperationalError:
            # Eski bazada ustun hali yaratilmagan bo'lsa, keyingi ishga tushishda o'tadi.
            pass

    materials = [
        ("MDF 18 mm","Plita","list",0,5,0),
        ("MDF 16 mm","Plita","list",0,5,0),
        ("Kromka","Kromka","metr",0,100,0),
        ("Furnitura","Furnitura","dona",0,50,0),
        ("Bo‘yoq","Bo‘yoq","kg",0,10,0),
        ("Oyna","Oyna","m²",0,5,0),
        ("Akril","Plita","list",0,3,0)
    ]
    conn.executemany(
        "INSERT OR IGNORE INTO ombor(nomi,kategoriya,birlik,qoldiq,min_qoldiq,narx) VALUES(?,?,?,?,?,?)",
        materials
    )
    # V5 migratsiya: ishchining maxfiy pasport ma'lumotlari
    worker_cols=[r[1] for r in conn.execute("PRAGMA table_info(ishchilar)").fetchall()]
    worker_new_cols={
        "pasport":"TEXT DEFAULT ''",
        "jshshir":"TEXT DEFAULT ''",
        "tugilgan_sana":"TEXT DEFAULT ''",
        "yashash_manzil":"TEXT DEFAULT ''",
        "pasport_berilgan_sana":"TEXT DEFAULT ''",
        "pasport_bergan":"TEXT DEFAULT ''",
        "favqulodda_telefon":"TEXT DEFAULT ''"
    }
    for col,ctype in worker_new_cols.items():
        if col not in worker_cols:
            conn.execute(f"ALTER TABLE ishchilar ADD COLUMN {col} {ctype}")

    # V5 migratsiya: xarajat tafsilotlari
    expense_cols=[r[1] for r in conn.execute("PRAGMA table_info(xarajatlar)").fetchall()]
    expense_new_cols={
        "xarajat_nomi":"TEXT DEFAULT ''",
        "kimga_berildi":"TEXT DEFAULT ''",
        "chek_havola":"TEXT DEFAULT ''"
    }
    for col,ctype in expense_new_cols.items():
        if col not in expense_cols:
            conn.execute(f"ALTER TABLE xarajatlar ADD COLUMN {col} {ctype}")

    # V4.5 migratsiya: shofyor ro'yxatdan o'tishi va admin tasdig'i
    driver_cols=[r[1] for r in conn.execute("PRAGMA table_info(shofyor_akkauntlari)").fetchall()]
    if "telefon" not in driver_cols:
        conn.execute("ALTER TABLE shofyor_akkauntlari ADD COLUMN telefon TEXT DEFAULT ''")
    if "admin_tasdiq" not in driver_cols:
        conn.execute("ALTER TABLE shofyor_akkauntlari ADD COLUMN admin_tasdiq INTEGER DEFAULT 0")

    # V4.4.2 migratsiya: yetkazishlar jadvalini eski va yangi tizimga moslash
    y_cols=[r[1] for r in conn.execute("PRAGMA table_info(yetkazishlar)").fetchall()]
    required_y_cols={"haydovchi_id","sana","navbat","mashina","holat","yolga_chiqdi",
                     "yetib_keldi","topshirildi","yetkazildi","benzin","yol_xarajati",
                     "izoh","qadoq_soni","created_at"}
    if not required_y_cols.issubset(set(y_cols)):
        conn.execute("PRAGMA foreign_keys = OFF")
        conn.execute("""
            CREATE TABLE yetkazishlar_yangi (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                buyurtma_id INTEGER NOT NULL,
                haydovchi_id INTEGER,
                sana TEXT DEFAULT '',
                navbat INTEGER DEFAULT 1,
                mashina TEXT DEFAULT '',
                holat TEXT DEFAULT 'Rejalashtirilgan',
                yolga_chiqdi TEXT DEFAULT '',
                yetib_keldi TEXT DEFAULT '',
                topshirildi TEXT DEFAULT '',
                yetkazildi TEXT DEFAULT '',
                benzin REAL DEFAULT 0,
                yol_xarajati REAL DEFAULT 0,
                izoh TEXT DEFAULT '',
                qadoq_soni INTEGER DEFAULT 1,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (buyurtma_id) REFERENCES buyurtmalar(id) ON DELETE CASCADE,
                FOREIGN KEY (haydovchi_id) REFERENCES ishchilar(id) ON DELETE SET NULL
            )
        """)
        old_cols=set(y_cols)
        rows=conn.execute("SELECT * FROM yetkazishlar").fetchall()
        for r in rows:
            def rv(name, default=None):
                return r[name] if name in old_cols else default
            driver_id=rv("haydovchi_id")
            if driver_id is None:
                driver_id=rv("shofyor_id")
            conn.execute("""
                INSERT INTO yetkazishlar_yangi
                (id,buyurtma_id,haydovchi_id,sana,navbat,mashina,holat,yolga_chiqdi,
                 yetib_keldi,topshirildi,yetkazildi,benzin,yol_xarajati,izoh,qadoq_soni,created_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,(
                rv("id"),rv("buyurtma_id"),driver_id,
                rv("sana",_tashkent_today()) or _tashkent_today(),
                rv("navbat",1) or 1,rv("mashina","") or "",
                rv("holat","Rejalashtirilgan") or "Rejalashtirilgan",
                rv("yolga_chiqdi","") or "",rv("yetib_keldi","") or "",
                rv("topshirildi","") or "",rv("yetkazildi","") or "",
                rv("benzin",0) or 0,rv("yol_xarajati",0) or 0,
                rv("izoh","") or "",rv("qadoq_soni",1) or 1,
                rv("created_at",_tashkent_now().strftime("%Y-%m-%d %H:%M:%S"))
            ))
        conn.execute("DROP TABLE yetkazishlar")
        conn.execute("ALTER TABLE yetkazishlar_yangi RENAME TO yetkazishlar")
        conn.execute("PRAGMA foreign_keys = ON")

    # V4.4 migratsiya: ishchi topshirig'ining haqiqiy boshlanish va tugash vaqti
    task_cols=[r[1] for r in conn.execute("PRAGMA table_info(ishchi_topshiriqlari)").fetchall()]
    if "boshlandi_vaqt" not in task_cols:
        conn.execute("ALTER TABLE ishchi_topshiriqlari ADD COLUMN boshlandi_vaqt TEXT DEFAULT ''")
    if "tugadi_vaqt" not in task_cols:
        conn.execute("ALTER TABLE ishchi_topshiriqlari ADD COLUMN tugadi_vaqt TEXT DEFAULT ''")

    # V4.2 migratsiya: xarajat to'lov usuli
    xarajat_cols=[r[1] for r in conn.execute("PRAGMA table_info(xarajatlar)").fetchall()]
    if "tolov_usuli" not in xarajat_cols:
        conn.execute("ALTER TABLE xarajatlar ADD COLUMN tolov_usuli TEXT DEFAULT 'Naqd'")

    # V5.0.6 migratsiya: shartnoma uchun qo'shimcha buyurtma ma'lumotlari
    order_contract_cols=[r[1] for r in conn.execute("PRAGMA table_info(buyurtmalar)").fetchall()]
    order_contract_new={
        "pasport_id":"TEXT DEFAULT ''",
        "olcham":"TEXT DEFAULT ''",
        "soni":"INTEGER DEFAULT 1",
        "material":"TEXT DEFAULT ''",
        "rang":"TEXT DEFAULT ''",
        "tolov_usuli":"TEXT DEFAULT 'Naqd'",
        "oraliq_tolov":"REAL DEFAULT 0",
        "montaj":"TEXT DEFAULT 'Kiritilgan'",
        "yetkazish":"TEXT DEFAULT 'Kiritilgan'",
        "kafolat_muddati":"TEXT DEFAULT '12 oy'"
    }
    for col,ctype in order_contract_new.items():
        if col not in order_contract_cols:
            conn.execute(f"ALTER TABLE buyurtmalar ADD COLUMN {col} {ctype}")

    # V3 migratsiya: mijoz kuzatuv tokeni
    cols=[r[1] for r in conn.execute("PRAGMA table_info(buyurtmalar)").fetchall()]
    if "tracking_token" not in cols:
        conn.execute("ALTER TABLE buyurtmalar ADD COLUMN tracking_token TEXT DEFAULT ''")
    for row in conn.execute("SELECT id FROM buyurtmalar WHERE tracking_token IS NULL OR tracking_token='' ").fetchall():
        conn.execute("UPDATE buyurtmalar SET tracking_token=? WHERE id=?",(secrets.token_urlsafe(8),row[0]))
    # V4 migratsiya: mijoz, chegirma, keshbek, kafolat va lokatsiya ustunlari
    order_cols={r[1] for r in conn.execute("PRAGMA table_info(buyurtmalar)").fetchall()}
    migrations={
      "boshlanish_sana":"TEXT DEFAULT ''", "taxminiy_sana":"TEXT DEFAULT ''",
      "kechikish_foiz":"REAL DEFAULT 0", "maks_chegirma_foiz":"REAL DEFAULT 20",
      "kechikish_turi":"TEXT DEFAULT 'Korxona'", "keshbek_foiz":"REAL DEFAULT 0",
      "keshbek_summa":"REAL DEFAULT 0", "keshbek_amal_sana":"TEXT DEFAULT ''",
      "kafolat_boshlanish":"TEXT DEFAULT ''", "kafolat_tugash":"TEXT DEFAULT ''",
      "kafolat_sharti":"TEXT DEFAULT ''", "lokatsiya":"TEXT DEFAULT ''",
      "moljal":"TEXT DEFAULT ''", "qavat":"TEXT DEFAULT ''", "lift":"TEXT DEFAULT ''",
      "katta_mashina":"TEXT DEFAULT ''", "masul_xodim":"TEXT DEFAULT ''",
      "taxminiy_vaqt":"TEXT DEFAULT '18:00'", "rang_kodi":"TEXT DEFAULT ''",
      "yaltiroqlik":"TEXT DEFAULT ''", "mijozga_izoh":"TEXT DEFAULT ''",
      "kechikish_sababi":"TEXT DEFAULT ''", "aloqa_telefon":"TEXT DEFAULT ''",
      "bloklangan":"INTEGER DEFAULT 0",
      "buyurtma_turi":"TEXT DEFAULT 'To‘liq mebel'",
      "tolov_sharti":"TEXT DEFAULT 'Avans majburiy'",
      "avans_talab":"REAL DEFAULT 0",
      "avans_muddat_sana":"TEXT DEFAULT ''",
      "taklif_amal_sana":"TEXT DEFAULT ''",
      "chizma_versiya":"INTEGER DEFAULT 1",
      "bepul_ozgarish_limit":"INTEGER DEFAULT 2",
      "ozgarish_soni":"INTEGER DEFAULT 0",
      "rahbar_tasdiq":"INTEGER DEFAULT 0",
      "rahbar_tasdiq_vaqt":"TEXT DEFAULT ''",
      "muddat_tartibi":"TEXT DEFAULT 'Eski'",
      "muddat_boshlanish_vaqt":"TEXT DEFAULT ''",
      "ishlab_chiqarish_kun":"INTEGER DEFAULT 0",
      "rasmiy_muddat_sana":"TEXT DEFAULT ''",
      "rasmiy_muddat_vaqt":"TEXT DEFAULT '18:00'"
    }
    for col,typ in migrations.items():
        if col not in order_cols:
            conn.execute(f"ALTER TABLE buyurtmalar ADD COLUMN {col} {typ}")

    # V5.2 migratsiya: UZS / USD va to‘lov kunidagi kurs bo‘yicha hisob
    order_currency_cols={r[1] for r in conn.execute("PRAGMA table_info(buyurtmalar)").fetchall()}
    for col,typ in {
        "valyuta":"TEXT DEFAULT 'UZS'",
        "kurs_tartibi":"TEXT DEFAULT 'To‘lov kunidagi kurs'",
        "oxirgi_kurs":"REAL DEFAULT 0"
    }.items():
        if col not in order_currency_cols:
            conn.execute(f"ALTER TABLE buyurtmalar ADD COLUMN {col} {typ}")
    conn.execute("UPDATE buyurtmalar SET valyuta='UZS' WHERE valyuta IS NULL OR TRIM(valyuta)='' ")
    conn.execute("UPDATE buyurtmalar SET kurs_tartibi='To‘lov kunidagi kurs' WHERE kurs_tartibi IS NULL OR TRIM(kurs_tartibi)='' ")

    payment_currency_cols={r[1] for r in conn.execute("PRAGMA table_info(buyurtma_tolovlari)").fetchall()}
    for col,typ in {
        "tolov_valyutasi":"TEXT DEFAULT 'UZS'",
        "kurs":"REAL DEFAULT 1",
        "qabul_qilingan_summa":"REAL DEFAULT 0",
        "buyurtma_summa":"REAL DEFAULT 0",
        "uzs_ekvivalent":"REAL DEFAULT 0"
    }.items():
        if col not in payment_currency_cols:
            conn.execute(f"ALTER TABLE buyurtma_tolovlari ADD COLUMN {col} {typ}")
    conn.execute("UPDATE buyurtma_tolovlari SET tolov_valyutasi='UZS' WHERE tolov_valyutasi IS NULL OR TRIM(tolov_valyutasi)='' ")
    conn.execute("UPDATE buyurtma_tolovlari SET kurs=1 WHERE kurs IS NULL OR kurs<=0")
    conn.execute("UPDATE buyurtma_tolovlari SET qabul_qilingan_summa=miqdor WHERE COALESCE(qabul_qilingan_summa,0)=0 AND COALESCE(miqdor,0)<>0")
    conn.execute("UPDATE buyurtma_tolovlari SET buyurtma_summa=miqdor WHERE COALESCE(buyurtma_summa,0)=0 AND COALESCE(miqdor,0)<>0")
    conn.execute("UPDATE buyurtma_tolovlari SET uzs_ekvivalent=miqdor WHERE COALESCE(uzs_ekvivalent,0)=0 AND COALESCE(miqdor,0)<>0")

    stage_cols={r[1] for r in conn.execute("PRAGMA table_info(buyurtma_bosqichlari)").fetchall()}
    for col,typ in {"boshlanish_vaqti":"TEXT DEFAULT ''","tugash_vaqti":"TEXT DEFAULT ''","ishchi":"TEXT DEFAULT ''","izoh":"TEXT DEFAULT ''","media_url":"TEXT DEFAULT ''"}.items():
        if col not in stage_cols: conn.execute(f"ALTER TABLE buyurtma_bosqichlari ADD COLUMN {col} {typ}")
    # ADMIN_XAVFSIZLIK_V1
    # Eski ochiq admin parolini bir marta shifrlab bazaga ko'chiradi.
    # Keyingi ishga tushirishlarda parol hech qachon TXT yoki CMDda ko'rsatilmaydi.
    admin_login = (os.environ.get("PHARM_ERP_USER", "admin") or "admin").strip()
    admin_row = conn.execute(
        "SELECT id FROM admin_akkauntlari WHERE login=? LIMIT 1",
        (admin_login,),
    ).fetchone()
    old_password_path = os.path.join(_APP_DIR, "admin_parol.txt")
    if not admin_row:
        bootstrap_password = (os.environ.get("PHARM_ERP_PASSWORD") or "").strip()
        migrated_from_txt = False
        if not bootstrap_password and os.path.exists(old_password_path):
            try:
                with open(old_password_path, "r", encoding="utf-8") as f:
                    bootstrap_password = f.read().strip()
                migrated_from_txt = bool(bootstrap_password)
            except OSError:
                bootstrap_password = ""
        if bootstrap_password:
            conn.execute(
                "INSERT INTO admin_akkauntlari(login,parol_hash,faol) VALUES(?,?,1)",
                (admin_login, generate_password_hash(bootstrap_password)),
            )
            if migrated_from_txt:
                try:
                    os.remove(old_password_path)
                except OSError:
                    pass
    elif os.path.exists(old_password_path):
        # Bazada xavfsiz akkaunt mavjud bo'lsa, eski ochiq parol fayli kerak emas.
        try:
            os.remove(old_password_path)
        except OSError:
            pass

    conn.commit()
    conn.close()



def _safe_filename(value):
    value=str(value or '').strip()
    value=re.sub(r'[^A-Za-z0-9_-]+','_',value)
    return value.strip('_') or 'buyurtma'


def _contract_dir():
    path=os.path.join(os.path.dirname(os.path.abspath(__file__)),'shartnomalar')
    os.makedirs(path,exist_ok=True)
    return path


def _fmt_money(value):
    try:
        return f"{float(value or 0):,.0f}".replace(","," ")
    except Exception:
        return "0"


def _add_docx_field(document,label,value):
    p=document.add_paragraph()
    p.paragraph_format.space_after=Pt(3)
    r=p.add_run(label+": ")
    r.bold=True
    p.add_run(str(value or "________________"))


def generate_order_contract(oid):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    if not order:
        c.close()
        raise ValueError("Buyurtma topilmadi")
    last=c.execute("SELECT COALESCE(MAX(versiya),0) FROM shartnoma_versiyalari WHERE buyurtma_id=?",(oid,)).fetchone()[0]
    version=int(last or 0)+1
    base=f"{_safe_filename(order['kod'])}_shartnoma_v{version}"
    docx_path=os.path.join(_contract_dir(),base+".docx")
    pdf_path=os.path.join(_contract_dir(),base+".pdf")
    order_currency=_currency_code(order["valyuta"] if "valyuta" in order.keys() else "UZS")
    currency_name="AQSh dollari (USD)" if order_currency=="USD" else "O‘zbekiston so‘mi (UZS)"
    currency_rule=("So‘mda amalga oshirilgan har bir to‘lov to‘lov kunida kiritilgan 1 USD kursi bo‘yicha "
                   "AQSh dollaridagi qarzdorlikdan ayriladi." if order_currency=="USD" else
                   "Hisob-kitob O‘zbekiston so‘mida amalga oshiriladi.")

    # DOCX
    doc=Document()
    section=doc.sections[0]
    section.top_margin=Cm(1.5); section.bottom_margin=Cm(1.5)
    section.left_margin=Cm(1.7); section.right_margin=Cm(1.7)
    styles=doc.styles
    styles['Normal'].font.name='Arial'
    styles['Normal'].font.size=Pt(10)

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("MEBEL360°\nMEBEL ISHLAB CHIQARISH (BUYURTMA) SHARTNOMASI")
    r.bold=True; r.font.size=Pt(14)
    p2=doc.add_paragraph()
    p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Shartnoma № {order['kod']}  |  Sana: {str(order['created_at'])[:10]}").bold=True

    doc.add_heading("1. TOMONLAR VA BUYURTMA", level=1)
    _add_docx_field(doc,"Buyurtmachi",order['mijoz'])
    _add_docx_field(doc,"Pasport/ID",order['pasport_id'])
    _add_docx_field(doc,"Telefon",order['telefon'])
    _add_docx_field(doc,"Manzil",order['manzil'])

    doc.add_heading("2. BUYURTMA TARKIBI", level=1)
    table=doc.add_table(rows=2,cols=7)
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style='Table Grid'
    headers=["Kod","Mahsulot","O‘lcham","Soni","Material","Rang","Jami narx"]
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]
        cell.text=h
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    vals=[order['kod'],order['mahsulot'],order['olcham'],order['soni'],
          order['material'],order['rang'],_currency_money(order['umumiy_narx'],order_currency)]
    for i,v in enumerate(vals):
        table.rows[1].cells[i].text=str(v or "")

    doc.add_heading("3. TO‘LOV TARTIBI", level=1)
    _add_docx_field(doc,"Kelishuv valyutasi",currency_name)
    _add_docx_field(doc,"Umumiy summa",_currency_money(order['umumiy_narx'],order_currency))
    _add_docx_field(doc,"Avans",_currency_money(order['oldindan_tolov'],order_currency))
    _add_docx_field(doc,"Oraliq to‘lov",_currency_money(order['oraliq_tolov'],order_currency))
    qoldiq=float(order['umumiy_narx'] or 0)-float(order['oldindan_tolov'] or 0)-float(order['oraliq_tolov'] or 0)
    _add_docx_field(doc,"Qoldiq",_currency_money(max(0,qoldiq),order_currency))
    _add_docx_field(doc,"Kurs tartibi",order['kurs_tartibi'] if 'kurs_tartibi' in order.keys() else 'To‘lov kunidagi kurs')
    _add_docx_field(doc,"To‘lov usuli",order['tolov_usuli'])
    doc.add_paragraph(currency_rule)

    doc.add_heading("4. MUDDAT, YETKAZISH VA MONTAJ", level=1)
    _add_docx_field(doc,"Boshlanish sanasi",order['boshlanish_sana'])
    _add_docx_field(doc,"Rejalashtirilgan tugash sanasi",order['tugash_sana'] or order['taxminiy_sana'])
    _add_docx_field(doc,"Yetkazish",order['yetkazish'])
    _add_docx_field(doc,"Montaj",order['montaj'])
    _add_docx_field(doc,"Yetkazish manzili",order['manzil'])
    _add_docx_field(doc,"Mo‘ljal",order['moljal'])
    _add_docx_field(doc,"Qavat",order['qavat'])
    _add_docx_field(doc,"Lift",order['lift'])
    _add_docx_field(doc,"Katta mashina kirishi",order['katta_mashina'])

    doc.add_heading("5. SIFAT, TASDIQLASH VA KAFOLAT", level=1)
    paragraphs=[
        "Mebel tasdiqlangan chizma, o‘lcham, material va rang asosida ishlab chiqariladi.",
        "Buyurtmachi chizma yoki materialni tasdiqlagandan keyingi o‘zgarishlar qo‘shimcha narx va muddat bilan bajariladi.",
        f"Kafolat muddati: {order['kafolat_muddati'] or '12 oy'}. "
        f"Kafolat sharti: {order['kafolat_sharti'] or 'Ishlab chiqarish va montaj nuqsonlariga amal qiladi.'}",
        "Noto‘g‘ri foydalanish, namlik, mexanik shikast va uchinchi shaxs ta’miri kafolatga kirmaydi.",
        "Buyurtmachi sababli tasdiq, to‘lov yoki obyektga kirish kechiksa, ish muddati tegishlicha uzayadi."
    ]
    for s in paragraphs:
        doc.add_paragraph(s,style=None)

    doc.add_heading("6. QABUL QILISH VA YAKUNIY SHARTLAR", level=1)
    for s in [
        "Mebel topshirilganda qabul-topshirish dalolatnomasi imzolanadi.",
        "Yakuniy to‘lov to‘liq amalga oshirilmaguncha Pudratchi mahsulotni topshirmaslikka haqli.",
        "Buyurtmachi buyurtmani bekor qilsa, bajarilgan ishlar, sotib olingan materiallar va tasdiqlangan xarajatlar hisobdan ushlab qolinadi.",
        "Tomonlarning elektron xabar, SMS yoki dasturdagi tasdig‘i yozma kelishuv sifatida qabul qilinadi."
    ]:
        doc.add_paragraph(s)

    if order['izoh']:
        doc.add_heading("Qo‘shimcha izoh",level=2)
        doc.add_paragraph(order['izoh'])

    doc.add_paragraph("\n")
    sig=doc.add_table(rows=3,cols=2)
    sig.style='Table Grid'
    sig.cell(0,0).text="PUDRATCHI: Mebel360°"
    sig.cell(0,1).text=f"BUYURTMACHI: {order['mijoz']}"
    sig.cell(1,0).text="Imzo: __________________"
    sig.cell(1,1).text="Imzo: __________________"
    sig.cell(2,0).text="Sana: __________________"
    sig.cell(2,1).text="Sana: __________________"
    doc.save(docx_path)

    # PDF
    out=canvas.Canvas(pdf_path,pagesize=A4)
    w,h=A4
    y=h-45
    def draw_wrapped(txt,bold=False,size=10,indent=45,space=15):
        nonlocal y
        out.setFont('Helvetica-Bold' if bold else 'Helvetica',size)
        ascii_txt=(str(txt).replace("‘","'").replace("’","'").replace("–","-")
                   .replace("—","-").replace("o‘","o'").replace("g‘","g'"))
        lines=simpleSplit(ascii_txt,'Helvetica-Bold' if bold else 'Helvetica',size,w-indent-45)
        for line in lines:
            if y<55:
                out.showPage(); y=h-45
            out.drawString(indent,y,line); y-=space

    draw_wrapped("MEBEL360°",True,16,45,20)
    draw_wrapped("MEBEL ISHLAB CHIQARISH (BUYURTMA) SHARTNOMASI",True,13,45,20)
    draw_wrapped(f"Shartnoma № {order['kod']} | Sana: {str(order['created_at'])[:10]}",True,10)
    y-=5
    pdf_items=[
        ("Buyurtmachi",order['mijoz']),("Pasport/ID",order['pasport_id']),
        ("Telefon",order['telefon']),("Manzil",order['manzil']),
        ("Mahsulot",order['mahsulot']),("O'lcham",order['olcham']),
        ("Soni",order['soni']),("Material",order['material']),("Rang",order['rang']),
        ("Kelishuv valyutasi",currency_name),
        ("Umumiy summa",_currency_money(order['umumiy_narx'],order_currency)),
        ("Avans",_currency_money(order['oldindan_tolov'],order_currency)),
        ("Oraliq to'lov",_currency_money(order['oraliq_tolov'],order_currency)),
        ("Qoldiq",_currency_money(max(0,qoldiq),order_currency)),
        ("Kurs tartibi",order['kurs_tartibi'] if 'kurs_tartibi' in order.keys() else "To'lov kunidagi kurs"),
        ("To'lov usuli",order['tolov_usuli']),
        ("Boshlanish",order['boshlanish_sana']),("Tugash",order['tugash_sana'] or order['taxminiy_sana']),
        ("Yetkazish",order['yetkazish']),("Montaj",order['montaj']),
        ("Mo'ljal",order['moljal']),("Qavat",order['qavat']),("Lift",order['lift']),
        ("Katta mashina",order['katta_mashina']),("Kafolat",order['kafolat_muddati'])
    ]
    for label,val in pdf_items:
        draw_wrapped(f"{label}: {val or '-'}",False,10)
    y-=7
    for s in paragraphs:
        draw_wrapped("- "+s,False,9,50,13)
    if order['izoh']:
        draw_wrapped("Izoh: "+order['izoh'],False,9)
    y-=20
    draw_wrapped("Pudratchi imzosi: ____________________",False,10)
    draw_wrapped("Buyurtmachi imzosi: ____________________",False,10)
    out.save()

    c.execute("""INSERT INTO shartnoma_versiyalari
                 (buyurtma_id,versiya,docx_fayl,pdf_fayl)
                 VALUES(?,?,?,?)""",(oid,version,docx_path,pdf_path))
    c.commit(); c.close()
    return {"version":version,"docx":docx_path,"pdf":pdf_path}


def latest_order_contract(oid):
    c=get_db()
    row=c.execute("""SELECT * FROM shartnoma_versiyalari
                     WHERE buyurtma_id=? ORDER BY versiya DESC,id DESC LIMIT 1""",(oid,)).fetchone()
    c.close()
    return row


def jdata():
    d = request.get_json(silent=True)
    return d if isinstance(d, dict) else {}


def calc_hours(start_text, end_text):
    s = datetime.strptime(start_text, "%H:%M")
    e = datetime.strptime(end_text, "%H:%M")
    sec = (e - s).total_seconds()
    if sec < 0:
        sec += 86400
    return round(sec / 3600, 2)


def log_action(amal, tafsilot=''):
    try:
        c=get_db(); c.execute("INSERT INTO audit_log(sana_vaqt,amal,tafsilot) VALUES(?,?,?)",(_tashkent_now().strftime("%Y-%m-%d %H:%M:%S"),amal,tafsilot)); c.commit(); c.close()
    except Exception:
        pass


# ---------- CSRF HIMOYASI ----------
def get_csrf_token():
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(16)
    return session['csrf_token']

app.jinja_env.globals['csrf_token'] = get_csrf_token


def _csrf_valid():
    token = session.get('csrf_token')
    sent = request.form.get('csrf_token', '')
    return bool(token) and secrets.compare_digest(token, sent)


# ---------- LOGIN URINISHLARINI CHEKLASH (BRUTE-FORCE HIMOYASI) ----------
_login_attempts = {}
LOGIN_MAX_ATTEMPTS = 5
LOGIN_LOCKOUT_SECONDS = 900

def _login_key(username):
    return f"{request.remote_addr}:{username}"

def _is_login_locked(username):
    now = datetime.now().timestamp()
    rec = _login_attempts.get(_login_key(username))
    if rec and rec[1] and now < rec[1]:
        return True, int(rec[1]-now)
    return False, 0

def _register_failed_login(username):
    now = datetime.now().timestamp()
    key = _login_key(username)
    count, locked_until = _login_attempts.get(key, (0, 0))
    count += 1
    locked_until = now + LOGIN_LOCKOUT_SECONDS if count >= LOGIN_MAX_ATTEMPTS else 0
    _login_attempts[key] = (count, locked_until)

def _clear_login_attempts(username):
    _login_attempts.pop(_login_key(username), None)


# ---------- PAROL MUSTAHKAMLIGI ----------
def _weak_password(pw):
    """True qaytaradi agar parol zaif bo'lsa: kamida 8 belgi, kamida bitta harf va bitta raqam kerak."""
    if not pw or len(pw) < 8:
        return True
    has_letter = any(ch.isalpha() for ch in pw)
    has_digit = any(ch.isdigit() for ch in pw)
    return not (has_letter and has_digit)


# ---------- AVTOMATIK KUNLIK ZAXIRA ----------
def _auto_backup_if_needed():
    try:
        backup_dir = os.path.join(_APP_DIR, 'backups')
        os.makedirs(backup_dir, exist_ok=True)
        today = _tashkent_today()
        marker = os.path.join(backup_dir, f".last_backup_{today}")
        if os.path.exists(marker):
            return
        if not os.path.exists(DB_NAME):
            return
        import shutil
        dest = os.path.join(backup_dir, f"pharm_mebel_backup_{today}.db")
        shutil.copyfile(DB_NAME, dest)
        with open(marker, 'w') as f:
            f.write('1')
        # Oxirgi 30 kunlik zaxiradan ortig'ini o'chirib, joy tejash
        old_backups = sorted(
            [f for f in os.listdir(backup_dir) if f.startswith('pharm_mebel_backup_') and f.endswith('.db')]
        )
        for old in old_backups[:-30]:
            try:
                os.remove(os.path.join(backup_dir, old))
            except Exception:
                pass
    except Exception:
        pass




# PHARM_MEBEL_SPLASH_LOGO_CLOCK_V1
SPLASH_HTML = '<!doctype html>\n<html lang="uz">\n<head>\n<meta charset="utf-8">\n<meta name="viewport" content="width=device-width,initial-scale=1">\n<title>Mebel360°</title>\n<style>\n*{box-sizing:border-box}\nhtml,body{margin:0;min-height:100%;font-family:Arial,sans-serif}\nbody{\n  min-height:100vh;display:grid;place-items:center;overflow:hidden;color:#14532d;\n  background:\n    radial-gradient(circle at 15% 18%,rgba(134,239,172,.50),transparent 34%),\n    radial-gradient(circle at 85% 82%,rgba(187,247,208,.78),transparent 38%),\n    linear-gradient(145deg,#fbfffc,#dcfce7);\n}\n.splash{\n  width:min(830px,94vw);min-height:min(680px,94vh);padding:24px 30px;\n  border:1px solid rgba(34,197,94,.25);border-radius:30px;\n  background:rgba(255,255,255,.84);box-shadow:0 28px 80px rgba(20,83,45,.18);\n  backdrop-filter:blur(13px);display:flex;flex-direction:column;align-items:center;\n  justify-content:center;text-align:center;animation:appear .65s ease both;\n}\n.logo{\n  width:min(680px,96%);max-height:385px;object-fit:contain;border-radius:22px;\n  filter:drop-shadow(0 14px 20px rgba(20,83,45,.13));animation:logoIn .9s ease both;\n}\n.clock{\n  margin-top:5px;font-size:clamp(48px,9vw,86px);line-height:1;font-weight:900;\n  letter-spacing:4px;color:#166534;text-shadow:0 4px 18px rgba(22,101,52,.13);\n}\n.date{margin-top:11px;font-size:clamp(16px,3vw,23px);font-weight:700;color:#3f7a52}\n.welcome{margin-top:17px;font-size:clamp(18px,3vw,23px);font-weight:900}\n.status{margin-top:7px;font-size:14px;color:#568467}\n.progress{\n  width:min(500px,84%);height:8px;margin-top:18px;overflow:hidden;\n  border-radius:999px;background:#d1fae5;\n}\n.bar{\n  width:0;height:100%;border-radius:999px;\n  background:linear-gradient(90deg,#4ade80,#15803d);animation:loading 5s linear forwards;\n}\n.enter{\n  margin-top:16px;padding:12px 28px;border:0;border-radius:13px;background:#16a34a;\n  color:#fff;font-size:16px;font-weight:800;cursor:pointer;\n  box-shadow:0 9px 24px rgba(22,163,74,.25);\n}\n.enter:hover{background:#15803d}\n@keyframes appear{from{opacity:0;transform:scale(.97)}to{opacity:1;transform:scale(1)}}\n@keyframes logoIn{from{opacity:0;transform:translateY(-15px)}to{opacity:1;transform:translateY(0)}}\n@keyframes loading{to{width:100%}}\n@media(max-width:600px){\n  .splash{min-height:94vh;padding:17px 12px;border-radius:22px}\n  .logo{width:100%;max-height:300px}\n  .clock{letter-spacing:2px}\n}\n</style>\n</head>\n<body>\n<section class="splash">\n  <img class="logo" src="/static/mebel360-logo.png?v=20260722" alt="Mebel360° logosi">\n  <div id="clock" class="clock">00:00:00</div>\n  <div id="date" class="date"></div>\n  <div class="welcome">Mebel360° dasturiga xush kelibsiz</div>\n  <div class="status">Dastur ochilmoqda...</div>\n  <div class="progress"><div class="bar"></div></div>\n  <button class="enter" type="button" onclick="goNext()">Kirish</button>\n</section>\n<script>\nconst NEXT_URL = {{ next_url|tojson }};\nconst DAYS = ["Yakshanba","Dushanba","Seshanba","Chorshanba","Payshanba","Juma","Shanba"];\nconst MONTHS = ["yanvar","fevral","mart","aprel","may","iyun",\n                "iyul","avgust","sentabr","oktabr","noyabr","dekabr"];\nfunction two(n){ return String(n).padStart(2,"0"); }\nfunction updateClock(){\n  // UTC vaqtiga aniq 5 soat qo‘shiladi. Bu eski Windows 7 brauzerlarida ham ishlaydi.\n  const now = new Date(Date.now() + 5*60*60*1000);\n  document.getElementById("clock").textContent =\n    two(now.getUTCHours())+":"+two(now.getUTCMinutes())+":"+two(now.getUTCSeconds());\n  document.getElementById("date").textContent =\n    DAYS[now.getUTCDay()]+", "+now.getUTCDate()+" "+MONTHS[now.getUTCMonth()]+" "+now.getUTCFullYear()+" · Toshkent";\n}\nlet leaving = false;\nfunction goNext(){\n  if(leaving) return;\n  leaving = true;\n  document.body.style.transition = "opacity .35s ease";\n  document.body.style.opacity = "0";\n  setTimeout(() => window.location.replace(NEXT_URL), 330);\n}\nupdateClock();\nsetInterval(updateClock,1000);\nsetTimeout(goNext,5000);\n</script>\n</body>\n</html>'

@app.before_request
def require_login():
    _auto_backup_if_needed()
    public_endpoints = {
        "splash", "mebel360_manifest", "mebel360_service_worker", "mebel360_offline", "mebel360_health",
        "login", "admin_setup", "static", "public_track", "order_qr",
        "worker_register", "worker_verify", "worker_login", "worker_logout", "driver_login", "driver_logout", "driver_register",
        "manager_login", "manager_logout", "constructor_login", "constructor_logout"
    }
    # CSRF tekshiruvi: barcha oddiy HTML-forma orqali yuboriladigan POST so'rovlar uchun.
    # /api/ yo'nalishlari brauzerdagi JS orqali (fetch, o'sha domendan) chaqiriladi, shuning
    # uchun bu yerda tekshirilmaydi.
    if request.method == 'POST' and not request.path.startswith('/api/'):
        if not _csrf_valid():
            return "Xato: sessiya muddati tugagan yoki noto'g'ri so'rov (CSRF). Sahifani qayta yuklab, qayta urinib ko'ring.", 400
    if request.endpoint in public_endpoints or request.path.startswith('/ishchi/public/'):
        return None
    if request.path.startswith('/ishchi/'):
        if not session.get("worker_account_id"):
            return redirect(url_for("worker_login"))
        return None
    if request.path.startswith('/shofyor/'):
        if not session.get("driver_account_id"):
            return redirect(url_for("driver_login"))
        return None
    if request.path.startswith('/menejer/'):
        if session.get("staff_role") != "menejer" and not session.get("logged_in"):
            return redirect(url_for("manager_login"))
        return None
    if request.path.startswith('/konstruktor/'):
        if session.get("staff_role") != "konstruktor" and not session.get("logged_in"):
            return redirect(url_for("constructor_login"))
        return None
    if not session.get("logged_in"):
        return redirect(url_for("login"))


def _admin_account(login):
    c = get_db()
    row = c.execute(
        "SELECT id,login,parol_hash,faol FROM admin_akkauntlari WHERE login=? LIMIT 1",
        ((login or "").strip(),),
    ).fetchone()
    c.close()
    return row


def _admin_is_configured():
    c = get_db()
    row = c.execute("SELECT id FROM admin_akkauntlari WHERE faol=1 LIMIT 1").fetchone()
    c.close()
    return bool(row)


@app.route("/admin/setup", methods=["GET", "POST"])
def admin_setup():
    if _admin_is_configured():
        return redirect(url_for("login"))
    error = ""
    if request.method == "POST":
        user = (request.form.get("user") or "admin").strip()
        password = request.form.get("password") or ""
        confirm = request.form.get("confirm") or ""
        if len(user) < 3:
            error = "Login kamida 3 ta belgidan iborat bo‘lsin."
        elif _weak_password(password):
            error = "Parol kamida 8 belgi, kamida bitta harf va bitta raqamdan iborat bo‘lsin."
        elif password != confirm:
            error = "Ikki parol bir xil emas."
        else:
            c = get_db()
            try:
                c.execute(
                    "INSERT INTO admin_akkauntlari(login,parol_hash,faol) VALUES(?,?,1)",
                    (user, generate_password_hash(password)),
                )
                c.commit()
            except sqlite3.IntegrityError:
                c.rollback()
                error = "Bu login avval ishlatilgan."
            finally:
                c.close()
            if not error:
                session.clear()
                session["logged_in"] = True
                session["user"] = user
                log_action("admin_first_setup", f"user={user}")
                return redirect(url_for("home"))
    return render_template_string(ADMIN_SETUP_HTML, error=error)


@app.route("/login", methods=["GET","POST"])
def login():
    if not _admin_is_configured():
        return redirect(url_for("admin_setup"))
    error=''
    if request.method=='POST':
        user=(request.form.get('user') or '').strip()
        password=request.form.get('password') or ''
        locked, wait_sec = _is_login_locked(user)
        account = _admin_account(user)
        if locked:
            minutes = max(1, (wait_sec + 59) // 60)
            error=f'Juda ko‘p xato urinish. Taxminan {minutes} daqiqadan so‘ng qayta urinib ko‘ring.'
        elif account and int(account['faol'] or 0) == 1 and check_password_hash(account['parol_hash'], password):
            _clear_login_attempts(user)
            session.clear()
            session['logged_in']=True
            session['user']=account['login']
            log_action('admin_login', f"user={account['login']}")
            return redirect(url_for('home'))
        else:
            _register_failed_login(user)
            log_action('admin_login_failed', f'user={user}')
            error='Login yoki parol xato'
    return render_template_string(LOGIN_HTML,error=error)


@app.route("/logout")
def logout():
    user = session.get("user", "")
    log_action("admin_logout", f"user={user}")
    session.clear()
    return redirect(url_for('login'))


@app.route("/")
def splash():
    if session.get("logged_in"):
        next_url = url_for("home")
    elif session.get("staff_role") == "menejer":
        next_url = url_for("manager_dashboard")
    elif session.get("staff_role") == "konstruktor":
        next_url = url_for("constructor_dashboard")
    else:
        next_url = url_for("login")
    return render_template_string(SPLASH_HTML, next_url=next_url)


@app.route("/dashboard")
def home():
    return render_template_string(HTML)


# ---------- ISHCHILAR ----------
@app.route("/api/ishchilar", methods=["GET"])
def workers_get():
    c = get_db()
    rows = c.execute("SELECT * FROM ishchilar WHERE faol=1 ORDER BY ism,familiya").fetchall()
    c.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/ishchilar", methods=["POST"])
def workers_add():
    d = jdata()
    if not str(d.get("ism","")).strip():
        return jsonify({"message":"Ism kiritilmagan"}),400
    c = get_db()
    c.execute("""INSERT INTO ishchilar
    (ism,familiya,telefon,lavozim,ishga_kirgan_sana,staj_yil,kunlik_stavka,oylik_maosh,
     sifat_ball,tezlik_ball,intizom_ball,izoh,pasport,jshshir,tugilgan_sana,yashash_manzil,
     pasport_berilgan_sana,pasport_bergan,favqulodda_telefon)
    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(
        d.get("ism","").strip(),d.get("familiya","").strip(),d.get("telefon","").strip(),
        d.get("lavozim","").strip(),d.get("ishga_kirgan_sana",""),
        float(d.get("staj_yil") or 0),float(d.get("kunlik_stavka") or 0),
        float(d.get("oylik_maosh") or 0),float(d.get("sifat_ball") or 5),
        float(d.get("tezlik_ball") or 5),float(d.get("intizom_ball") or 5),
        d.get("izoh","").strip(),d.get("pasport","").strip(),d.get("jshshir","").strip(),
        d.get("tugilgan_sana",""),d.get("yashash_manzil","").strip(),
        d.get("pasport_berilgan_sana",""),d.get("pasport_bergan","").strip(),
        d.get("favqulodda_telefon","").strip()
    ))
    c.commit(); c.close()
    return jsonify({"status":"ok"})


@app.route("/api/ishchilar/<int:i>", methods=["DELETE"])
def workers_delete(i):
    c=get_db(); c.execute("UPDATE ishchilar SET faol=0 WHERE id=?",(i,)); c.commit(); c.close()
    return jsonify({"status":"ok"})


# ---------- KELDI KETDI ----------
@app.route("/api/keldi-ketdi", methods=["GET"])
def attendance_get():
    c=get_db()
    rows=c.execute("""SELECT k.*,i.ism,i.familiya FROM keldi_ketdi k
    JOIN ishchilar i ON i.id=k.ishchi_id ORDER BY sana DESC,id DESC LIMIT 200""").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/keldi-ketdi", methods=["POST"])
def attendance_add():
    d=jdata()
    try:
        hours=calc_hours(d["keldi_vaqti"],d["ketdi_vaqti"])
        c=get_db()
        c.execute("""INSERT INTO keldi_ketdi(ishchi_id,sana,keldi_vaqti,ketdi_vaqti,ish_soatlari)
        VALUES(?,?,?,?,?)""",(int(d["ishchi_id"]),d["sana"],d["keldi_vaqti"],d["ketdi_vaqti"],hours))
        c.commit(); c.close()
        return jsonify({"status":"ok"})
    except Exception as e:
        return jsonify({"message":str(e)}),400



@app.route("/api/davomat/keldi", methods=["POST"])
def attendance_clock_in():
    d=jdata()
    try:
        worker_id=int(d["ishchi_id"])
        now=_tashkent_now()
        sana=now.strftime("%Y-%m-%d")
        vaqt=now.strftime("%H:%M")
        c=get_db()
        row=c.execute("""SELECT id FROM keldi_ketdi
                         WHERE ishchi_id=? AND sana=? AND (ketdi_vaqti='' OR ketdi_vaqti IS NULL)
                         ORDER BY id DESC LIMIT 1""",(worker_id,sana)).fetchone()
        if row:
            c.close()
            return jsonify({"message":"Bu ishchi bugun allaqachon kelgan deb belgilangan."}),400
        c.execute("""INSERT INTO keldi_ketdi(ishchi_id,sana,keldi_vaqti,ketdi_vaqti,ish_soatlari)
                     VALUES(?,?,?,?,0)""",(worker_id,sana,vaqt,""))
        c.commit(); c.close()
        return jsonify({"status":"ok","sana":sana,"vaqt":vaqt})
    except Exception as e:
        return jsonify({"message":str(e)}),400


@app.route("/api/davomat/ketdi", methods=["POST"])
def attendance_clock_out():
    d=jdata()
    try:
        worker_id=int(d["ishchi_id"])
        now=_tashkent_now()
        sana=now.strftime("%Y-%m-%d")
        vaqt=now.strftime("%H:%M")
        c=get_db()
        row=c.execute("""SELECT * FROM keldi_ketdi
                         WHERE ishchi_id=? AND sana=? AND (ketdi_vaqti='' OR ketdi_vaqti IS NULL)
                         ORDER BY id DESC LIMIT 1""",(worker_id,sana)).fetchone()
        if not row:
            c.close()
            return jsonify({"message":"Avval “Hozir keldi” tugmasini bosing."}),400
        hours=calc_hours(row["keldi_vaqti"],vaqt)
        c.execute("UPDATE keldi_ketdi SET ketdi_vaqti=?, ish_soatlari=? WHERE id=?",
                  (vaqt,hours,row["id"]))
        c.commit(); c.close()
        return jsonify({"status":"ok","sana":sana,"vaqt":vaqt,"ish_soatlari":hours})
    except Exception as e:
        return jsonify({"message":str(e)}),400


# ---------- ISH NATIJALARI ----------
@app.route("/api/ish-turlari")
def work_types_get():
    c=get_db(); rows=c.execute("SELECT * FROM ish_turlari WHERE faol=1 ORDER BY kategoriya,nomi").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/natijalar", methods=["GET"])
def results_get():
    c=get_db()
    rows=c.execute("""SELECT n.*,i.ism,i.familiya,t.nomi ish_turi,t.birlik
    FROM ish_natijalari n JOIN ishchilar i ON i.id=n.ishchi_id
    JOIN ish_turlari t ON t.id=n.ish_turi_id
    ORDER BY sana DESC,n.id DESC LIMIT 250""").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/natijalar", methods=["POST"])
def results_add():
    d=jdata()
    try:
        qty=float(d.get("miqdor") or 0); price=float(d.get("birlik_narxi") or 0)
        c=get_db()
        c.execute("""INSERT INTO ish_natijalari
        (ishchi_id,ish_turi_id,sana,miqdor,birlik_narxi,jami_haq,buyurtma_kodi,izoh)
        VALUES(?,?,?,?,?,?,?,?)""",(int(d["ishchi_id"]),int(d["ish_turi_id"]),d["sana"],
        qty,price,round(qty*price,2),d.get("buyurtma_kodi",""),d.get("izoh","")))
        c.commit(); c.close()
        return jsonify({"status":"ok"})
    except Exception as e:
        return jsonify({"message":str(e)}),400


# ---------- TOLOV VA JARIMA ----------
@app.route("/api/tolovlar", methods=["GET","POST"])
def payments():
    if request.method=="POST":
        d=jdata(); c=get_db()
        c.execute("INSERT INTO tolovlar(ishchi_id,sana,miqdor,turi,tavsifi) VALUES(?,?,?,?,?)",
                  (int(d["ishchi_id"]),d["sana"],float(d.get("miqdor") or 0),d.get("turi","Avans"),d.get("tavsifi","")))
        c.commit(); c.close(); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("""SELECT t.*,i.ism,i.familiya FROM tolovlar t
    JOIN ishchilar i ON i.id=t.ishchi_id ORDER BY sana DESC,t.id DESC LIMIT 200""").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/jarimalar", methods=["GET","POST"])
def penalties():
    if request.method=="POST":
        d=jdata(); c=get_db()
        c.execute("INSERT INTO jarimalar(ishchi_id,sana,miqdor,sababi) VALUES(?,?,?,?)",
                  (int(d["ishchi_id"]),d["sana"],float(d.get("miqdor") or 0),d.get("sababi","")))
        c.commit(); c.close(); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("""SELECT j.*,i.ism,i.familiya FROM jarimalar j
    JOIN ishchilar i ON i.id=j.ishchi_id ORDER BY sana DESC,j.id DESC LIMIT 200""").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


# ---------- BUYURTMALAR ----------
ORDER_STAGE_TEMPLATES = {
    "To‘liq mebel": [
        "O‘lchov olindi", "Chizma tayyorlandi", "Hisob-kitob jarayonda",
        "Narx kelishildi", "Mijoz tasdiqladi", "Material tayyorlandi",
        "Kesish", "Kromka", "Teshish", "Frezalash", "Bo‘yash",
        "Sayqalash", "Yig‘ish", "Sifat nazorati", "Qadoqlash",
        "Yetkazish/o‘rnatish", "Yakunlandi",
    ],
    "Faqat oyna": [
        "O‘lchov olindi", "Chizma tayyorlandi", "Hisob-kitob jarayonda",
        "Mijoz tasdiqladi", "Oynaga buyurtma berildi", "Kesildi",
        "Qirralari ishlov berildi", "Teshik ochildi", "Sifat nazorati",
        "Yetkazish/o‘rnatish", "Yakunlandi",
    ],
    "MDF fasad": [
        "O‘lchov olindi", "Chizma tayyorlandi", "Hisob-kitob jarayonda",
        "Mijoz tasdiqladi", "MDF tayyorlandi", "Kesish", "Frezalash",
        "Sayqalash", "Bo‘yash", "Sifat nazorati", "Qadoqlash",
        "Yetkazish/o‘rnatish", "Yakunlandi",
    ],
    "Ta’mirlash": [
        "O‘lchov olindi", "Muammo aniqlandi", "Hisob-kitob jarayonda",
        "Mijoz tasdiqladi", "Material tayyorlandi", "Ta’mirlash boshlandi",
        "Sifat nazorati", "Yetkazish/o‘rnatish", "Yakunlandi",
    ],
    "Aralash buyurtma": [
        "O‘lchov olindi", "Chizma tayyorlandi", "Hisob-kitob jarayonda",
        "Mijoz tasdiqladi", "Material tayyorlandi", "Kesish", "Kromka",
        "Oynaga buyurtma berildi", "Qirralari ishlov berildi", "Teshish",
        "Frezalash", "Bo‘yash", "Yig‘ish", "Sifat nazorati", "Qadoqlash",
        "Yetkazish/o‘rnatish", "Yakunlandi",
    ],
}

PAYMENT_STAGE_TEMPLATES = {
    "Avans majburiy": ["Avans olindi"],
    "Qisman avans": ["Qisman avans olindi"],
    "Avans talab qilinmaydi — ishonchli mijoz": ["Avans talab qilinmaydi", "Rahbar tasdiqladi"],
    "Muddatli to‘lov": ["Muddatli to‘lov tasdiqlandi", "Rahbar tasdiqladi"],
    "Shartnoma asosida to‘lov": ["Shartnoma tasdiqlandi", "Rahbar tasdiqladi"],
}

PRE_START_STAGES = {
    "O‘lchov olindi", "Chizma tayyorlandi", "Hisob-kitob jarayonda",
    "Narx kelishildi", "Mijoz tasdiqladi", "Muammo aniqlandi",
    "Avans olindi", "Qisman avans olindi", "Avans talab qilinmaydi",
    "Muddatli to‘lov tasdiqlandi", "Shartnoma tasdiqlandi", "Rahbar tasdiqladi",
}


def _custom_stage_list(raw):
    parts = re.split(r"[\n,;→]+", str(raw or ""))
    clean=[]
    for item in parts:
        item=re.sub(r"\s+", " ", item).strip(" -—>\t")
        if item and item not in clean:
            clean.append(item)
    return clean[:60]


def _workflow_stages(order_type, payment_condition, custom_stages=""):
    order_type=str(order_type or "To‘liq mebel").strip()
    payment_condition=str(payment_condition or "Avans majburiy").strip()
    base=_custom_stage_list(custom_stages) if order_type=="Maxsus" else list(
        ORDER_STAGE_TEMPLATES.get(order_type, ORDER_STAGE_TEMPLATES["To‘liq mebel"])
    )
    if not base:
        base=list(ORDER_STAGE_TEMPLATES["To‘liq mebel"])
    payment=list(PAYMENT_STAGE_TEMPLATES.get(payment_condition, ["Avans olindi"]))
    insert_at=next((i+1 for i,s in enumerate(base) if s=="Mijoz tasdiqladi"), min(5,len(base)))
    return base[:insert_at]+payment+base[insert_at:]


def _complete_order_stage(conn, order_id, stage_names, note=""):
    if isinstance(stage_names, str):
        stage_names=[stage_names]
    now_text=_tashkent_now().strftime("%Y-%m-%d %H:%M:%S")
    changed=0
    for name in stage_names:
        cur=conn.execute("""UPDATE buyurtma_bosqichlari SET bajarildi=1,
            boshlanish_vaqti=CASE WHEN COALESCE(boshlanish_vaqti,'')='' THEN ? ELSE boshlanish_vaqti END,
            tugash_vaqti=?, izoh=CASE WHEN ?<>'' THEN ? ELSE izoh END
            WHERE buyurtma_id=? AND bosqich=?""",
            (now_text,now_text,note,note,order_id,name))
        changed+=cur.rowcount
    return changed


def _payment_ready(order):
    keys=set(order.keys()) if hasattr(order,"keys") else set()
    condition=str(order["tolov_sharti"] or "Avans majburiy") if "tolov_sharti" in keys else "Avans majburiy"
    if condition in ("Avans talab qilinmaydi — ishonchli mijoz", "Muddatli to‘lov", "Shartnoma asosida to‘lov"):
        return bool(int(order["rahbar_tasdiq"] or 0)) if "rahbar_tasdiq" in keys else False
    paid=float(order["oldindan_tolov"] or 0)
    required=float(order["avans_talab"] or 0) if "avans_talab" in keys else 0
    return paid >= required if required>0 else paid>0


def _start_official_deadline(conn, order_id, trigger=""):
    order=conn.execute("SELECT * FROM buyurtmalar WHERE id=?",(order_id,)).fetchone()
    if not order:
        return {"started":False,"reason":"Buyurtma topilmadi"}
    keys=set(order.keys())
    if str(order["muddat_boshlanish_vaqt"] or "").strip():
        return {"started":False,"already":True,"date":str(order["rasmiy_muddat_sana"] or "")}
    now=_tashkent_now()
    days=int(order["ishlab_chiqarish_kun"] or 0) if "ishlab_chiqarish_kun" in keys else 0
    target_date=""
    if days>0:
        target_date=(now+timedelta(days=days)).date().isoformat()
    else:
        for key in ("rasmiy_muddat_sana","taxminiy_sana","tugash_sana"):
            if key in keys and str(order[key] or "").strip():
                target_date=str(order[key]).strip(); break
    target_time=_safe_time(
        order["rasmiy_muddat_vaqt"] if "rasmiy_muddat_vaqt" in keys and str(order["rasmiy_muddat_vaqt"] or "").strip()
        else (order["taxminiy_vaqt"] if "taxminiy_vaqt" in keys else "18:00")
    )
    now_text=now.strftime("%Y-%m-%d %H:%M:%S")
    conn.execute("""UPDATE buyurtmalar SET muddat_boshlanish_vaqt=?, rasmiy_muddat_sana=?,
        rasmiy_muddat_vaqt=?, boshlanish_sana=?,
        taxminiy_sana=CASE WHEN ?<>'' THEN ? ELSE taxminiy_sana END,
        taxminiy_vaqt=?, tugash_sana=CASE WHEN ?<>'' THEN ? ELSE tugash_sana END,
        holat='Jarayonda' WHERE id=?""",
        (now_text,target_date,target_time,now.date().isoformat(),target_date,target_date,
         target_time,target_date,target_date,order_id))
    if target_date:
        conn.execute("""INSERT INTO buyurtma_muddat_tarixi
          (buyurtma_id,eski_sana,eski_vaqt,yangi_sana,yangi_vaqt,sabab)
          VALUES(?,?,?,?,?,?)""",
          (order_id,str(order["taxminiy_sana"] or ""),_safe_time(order["taxminiy_vaqt"] if "taxminiy_vaqt" in keys else "18:00"),
           target_date,target_time,trigger or "Rasmiy muddat boshlandi"))
    return {"started":True,"date":target_date,"time":target_time,"start":now_text}


@app.route("/api/buyurtmalar", methods=["GET","POST"])
def orders():
    if request.method=="POST":
        d=jdata(); c=get_db()
        try:
            order_type=d.get("buyurtma_turi","To‘liq mebel") or "To‘liq mebel"
            payment_condition=d.get("tolov_sharti","Avans majburiy") or "Avans majburiy"
            order_currency=_currency_code(d.get("valyuta","UZS"))
            initial_received=float(d.get("oldindan_tolov") or 0)
            initial_payment_currency=_currency_code(d.get("boshlangich_tolov_valyutasi") or order_currency)
            initial_rate=float(d.get("boshlangich_kurs") or 0)
            initial_order_amount=0.0
            initial_uzs=0.0
            if initial_received>0:
                initial_order_amount,initial_uzs,initial_rate=_payment_conversion(
                    order_currency,initial_payment_currency,initial_received,initial_rate
                )
            columns=[
                "kod","mijoz","telefon","manzil","mahsulot","umumiy_narx","oldindan_tolov",
                "boshlanish_sana","tugash_sana","taxminiy_sana","taxminiy_vaqt","holat","izoh","tracking_token",
                "kechikish_foiz","maks_chegirma_foiz","keshbek_foiz","keshbek_summa",
                "kafolat_boshlanish","kafolat_tugash","lokatsiya","moljal","qavat","lift",
                "katta_mashina","masul_xodim","pasport_id","olcham","soni","material","rang","rang_kodi",
                "yaltiroqlik","mijozga_izoh","kechikish_sababi","aloqa_telefon",
                "tolov_usuli","oraliq_tolov","montaj","yetkazish","kafolat_muddati",
                "buyurtma_turi","tolov_sharti","avans_talab","avans_muddat_sana","taklif_amal_sana",
                "chizma_versiya","bepul_ozgarish_limit","ozgarish_soni","muddat_tartibi",
                "ishlab_chiqarish_kun","rasmiy_muddat_sana","rasmiy_muddat_vaqt",
                "valyuta","kurs_tartibi","oxirgi_kurs"
            ]
            values=[
                d["kod"],d["mijoz"],d.get("telefon",""),d.get("manzil",""),d.get("mahsulot",""),
                float(d.get("umumiy_narx") or 0),initial_order_amount,
                d.get("boshlanish_sana",""),d.get("tugash_sana",""),d.get("taxminiy_sana",""),
                _safe_time(d.get("taxminiy_vaqt")),d.get("holat","Yangi"),d.get("izoh",""),secrets.token_urlsafe(8),
                float(d.get("kechikish_foiz") or 0),float(d.get("maks_chegirma_foiz") or 20),
                float(d.get("keshbek_foiz") or 0),float(d.get("keshbek_summa") or 0),
                d.get("kafolat_boshlanish",""),d.get("kafolat_tugash",""),
                d.get("lokatsiya",""),d.get("moljal",""),d.get("qavat",""),d.get("lift",""),
                d.get("katta_mashina",""),d.get("masul_xodim",""),d.get("pasport_id",""),
                d.get("olcham",""),int(d.get("soni") or 1),d.get("material",""),d.get("rang",""),
                d.get("rang_kodi",""),d.get("yaltiroqlik",""),d.get("mijozga_izoh",""),
                d.get("kechikish_sababi",""),d.get("aloqa_telefon",""),d.get("tolov_usuli","Naqd"),
                float(d.get("oraliq_tolov") or 0),d.get("montaj","Kiritilgan"),
                d.get("yetkazish","Kiritilgan"),d.get("kafolat_muddati","12 oy"),
                order_type,payment_condition,float(d.get("avans_talab") or 0),d.get("avans_muddat_sana",""),
                d.get("taklif_amal_sana",""),int(d.get("chizma_versiya") or 1),
                int(d.get("bepul_ozgarish_limit") or 2),int(d.get("ozgarish_soni") or 0),
                "Tasdiqdan keyin",int(d.get("ishlab_chiqarish_kun") or 0),"",_safe_time(d.get("taxminiy_vaqt")),
                order_currency,"To‘lov kunidagi kurs",initial_rate if initial_rate>1 else 0
            ]
            sql=f"INSERT INTO buyurtmalar({','.join(columns)}) VALUES({','.join('?' for _ in columns)})"
            cur=c.execute(sql,values)
            oid=cur.lastrowid
            workflow=_workflow_stages(order_type,payment_condition,d.get("maxsus_bosqichlar",""))
            c.executemany("INSERT INTO buyurtma_bosqichlari(buyurtma_id,bosqich) VALUES(?,?)",
                          [(oid,s) for s in workflow])
            if initial_received>0:
                c.execute("""INSERT INTO buyurtma_tolovlari
                    (buyurtma_id,sana,miqdor,turi,izoh,tolov_valyutasi,kurs,
                     qabul_qilingan_summa,buyurtma_summa,uzs_ekvivalent)
                    VALUES(?,?,?,?,?,?,?,?,?,?)""",
                    (oid,d.get("boshlangich_tolov_sana") or _tashkent_today(),initial_order_amount,
                     "Avans",d.get("boshlangich_tolov_izoh","") or "Buyurtma yaratilganda kiritildi",
                     initial_payment_currency,initial_rate,initial_received,initial_order_amount,initial_uzs))
            if payment_condition=="Avans talab qilinmaydi — ishonchli mijoz":
                _complete_order_stage(c,oid,"Avans talab qilinmaydi","Ishonchli mijoz uchun avans talab qilinmaydi")
                c.execute("UPDATE buyurtmalar SET holat='Rahbar tasdiqlashi kutilmoqda' WHERE id=?",(oid,))
            initial_order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
            if _payment_ready(initial_order):
                pay_stage="Qisman avans olindi" if payment_condition=="Qisman avans" else "Avans olindi"
                _complete_order_stage(c,oid,pay_stage,"Buyurtma yaratilganda to‘lov kiritildi")
                _start_official_deadline(c,oid,"Avans qabul qilindi — rasmiy muddat boshlandi")
            if d.get("tasdiq_turi"):
                c.execute("INSERT INTO buyurtma_tasdiqlari(buyurtma_id,turi,holat,izoh) VALUES(?,?,?,?)",
                          (oid,d.get("tasdiq_turi"),"Kutilmoqda",d.get("tasdiq_izoh","")))
            if d.get("media_havola"):
                c.execute("INSERT INTO buyurtma_media(buyurtma_id,turi,havola,izoh) VALUES(?,?,?,?)",
                          (oid,d.get("media_turi","Rasm"),d.get("media_havola"),d.get("media_izoh","")))
            c.commit(); c.close()
            try:
                contract=generate_order_contract(oid)
                return jsonify({"status":"ok","buyurtma_id":oid,"shartnoma_versiya":contract["version"]})
            except Exception as contract_error:
                return jsonify({"status":"ok","buyurtma_id":oid,
                                "warning":"Buyurtma saqlandi, shartnoma yaratishda xato: "+str(contract_error)})
        except Exception as e:
            c.close(); return jsonify({"message":str(e)}),400
    c=get_db()
    rows=c.execute("""SELECT *,
    MAX(0,CAST(julianday('now')-julianday(tugash_sana) AS INTEGER)) kechikkan_kun,
    ROUND(MIN(CASE WHEN kechikish_turi='Korxona' THEN MAX(0,julianday('now')-julianday(tugash_sana))*kechikish_foiz ELSE 0 END,maks_chegirma_foiz),2) chegirma_foiz,
    ROUND(MAX(0,(umumiy_narx-oldindan_tolov)*(1-MIN(CASE WHEN kechikish_turi='Korxona' THEN MAX(0,julianday('now')-julianday(tugash_sana))*kechikish_foiz ELSE 0 END,maks_chegirma_foiz)/100.0)),2) qoldiq
    FROM buyurtmalar ORDER BY id DESC""").fetchall()
    result=[]
    for row in rows:
        item=dict(row)
        deadline=_order_deadline(row)
        item["muddat_iso"]=deadline.isoformat() if deadline else ""
        if deadline:
            item["muddat_matn"]=deadline.strftime("%d.%m.%Y %H:%M")
        elif str(item.get("muddat_tartibi") or "")=="Tasdiqdan keyin":
            condition=str(item.get("tolov_sharti") or "Avans majburiy")
            if condition in ("Avans talab qilinmaydi — ishonchli mijoz","Muddatli to‘lov","Shartnoma asosida to‘lov"):
                item["muddat_matn"]="Rahbar tasdig‘idan keyin"
            else:
                item["muddat_matn"]="Avansdan keyin"
        else:
            item["muddat_matn"]="Belgilanmagan"
        currency=_currency_code(item.get("valyuta","UZS"))
        rate=float(item.get("oxirgi_kurs") or 0)
        item["valyuta"]=currency
        item["umumiy_uzs"]=round(float(item.get("umumiy_narx") or 0)*(rate if currency=="USD" else 1),2) if (currency=="UZS" or rate>0) else 0
        item["qoldiq_uzs"]=round(float(item.get("qoldiq") or 0)*(rate if currency=="USD" else 1),2) if (currency=="UZS" or rate>0) else 0
        stage_stats=c.execute("SELECT COUNT(*) jami,COALESCE(SUM(bajarildi),0) bajarildi FROM buyurtma_bosqichlari WHERE buyurtma_id=?",(item["id"],)).fetchone()
        item["bosqich_jami"]=int(stage_stats["jami"] or 0)
        item["bosqich_bajarildi"]=int(stage_stats["bajarildi"] or 0)
        item["progress"]=round(item["bosqich_bajarildi"]*100/item["bosqich_jami"],1) if item["bosqich_jami"] else 0
        current_stage=c.execute("SELECT bosqich FROM buyurtma_bosqichlari WHERE buyurtma_id=? AND bajarildi=0 ORDER BY id LIMIT 1",(item["id"],)).fetchone()
        item["joriy_bosqich"]=current_stage["bosqich"] if current_stage else ("Yakunlandi" if item["bosqich_jami"] else "Bosqich yo‘q")
        paid_uzs=c.execute("SELECT COALESCE(SUM(uzs_ekvivalent),0) jami FROM buyurtma_tolovlari WHERE buyurtma_id=?",(item["id"],)).fetchone()
        item["tolangan_uzs"]=round(float(paid_uzs["jami"] or 0),2)
        result.append(item)
    c.close(); return jsonify(result)


@app.route("/api/buyurtma/<int:oid>/bosqichlar")
def order_stages(oid):
    c=get_db(); rows=c.execute("SELECT * FROM buyurtma_bosqichlari WHERE buyurtma_id=? ORDER BY id",(oid,)).fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/buyurtma-bosqich/<int:sid>", methods=["POST"])
def stage_toggle(sid):
    d=jdata(); c=get_db()
    stage_row=c.execute("""SELECT b.id,b.buyurtma_id,b.bosqich,o.muddat_tartibi,o.muddat_boshlanish_vaqt,
        o.tolov_sharti FROM buyurtma_bosqichlari b JOIN buyurtmalar o ON o.id=b.buyurtma_id
        WHERE b.id=?""",(sid,)).fetchone()
    if not stage_row:
        c.close(); return jsonify({"message":"Bosqich topilmadi"}),404
    oid=int(stage_row["buyurtma_id"]); stage_name=str(stage_row["bosqich"] or "")
    done=1 if d.get("bajarildi") else 0
    if done and str(stage_row["muddat_tartibi"] or "")=="Tasdiqdan keyin" and not str(stage_row["muddat_boshlanish_vaqt"] or "").strip() and stage_name not in PRE_START_STAGES:
        condition=str(stage_row["tolov_sharti"] or "Avans majburiy")
        c.close()
        wait="rahbar tasdig‘i" if condition in ("Avans talab qilinmaydi — ishonchli mijoz","Muddatli to‘lov","Shartnoma asosida to‘lov") else "avans"
        return jsonify({"message":f"Bu bosqichni boshlashdan oldin {wait} kerak."}),400
    now_text=_tashkent_now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("""UPDATE buyurtma_bosqichlari SET bajarildi=?,
      boshlanish_vaqti=CASE WHEN ?=1 AND COALESCE(boshlanish_vaqti,'')='' THEN ? ELSE boshlanish_vaqti END,
      tugash_vaqti=CASE WHEN ?=1 THEN ? ELSE '' END,
      ishchi=COALESCE(NULLIF(?,''),ishchi),izoh=COALESCE(NULLIF(?,''),izoh),media_url=COALESCE(NULLIF(?,''),media_url)
      WHERE id=?""",(done,done,now_text,done,now_text,d.get("ishchi",''),d.get("izoh",''),d.get("media_url",''),sid))
    deadline_result={}
    if done and stage_name in ("Avans olindi","Qisman avans olindi"):
        deadline_result=_start_official_deadline(c,oid,"Avans qabul qilindi — rasmiy muddat boshlandi")
    elif done and stage_name=="Rahbar tasdiqladi":
        c.execute("UPDATE buyurtmalar SET rahbar_tasdiq=1,rahbar_tasdiq_vaqt=? WHERE id=?",(now_text,oid))
        deadline_result=_start_official_deadline(c,oid,"Rahbar tasdiqladi — rasmiy muddat boshlandi")
    elif done and stage_name=="Mijoz tasdiqladi":
        condition=str(stage_row["tolov_sharti"] or "Avans majburiy")
        next_status="Rahbar tasdiqlashi kutilmoqda" if condition in ("Avans talab qilinmaydi — ishonchli mijoz","Muddatli to‘lov","Shartnoma asosida to‘lov") else "Avans kutilmoqda"
        c.execute("UPDATE buyurtmalar SET holat=? WHERE id=?",(next_status,oid))
    stats=c.execute("SELECT COUNT(*) jami,COALESCE(SUM(bajarildi),0) bajarildi FROM buyurtma_bosqichlari WHERE buyurtma_id=?",(oid,)).fetchone()
    last_done=c.execute("SELECT bosqich FROM buyurtma_bosqichlari WHERE buyurtma_id=? AND bajarildi=1 ORDER BY id DESC LIMIT 1",(oid,)).fetchone()
    jami=int(stats["jami"] or 0); bajarildi=int(stats["bajarildi"] or 0)
    last_name=(last_done["bosqich"] if last_done else '')
    current=c.execute("SELECT holat FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    holat=str(current["holat"] or "Yangi")
    if bajarildi == 0:
        holat='Yangi'
    elif jami and bajarildi >= jami:
        holat='Yetkazildi'
    elif last_name in ('Yakunlandi','Yetkazib berildi','Buyurtma yopildi'):
        holat='Yetkazildi'
    elif last_name in ('Qadoqlash','Yetkazish/o‘rnatish','Yetkazishga tayyor','Haydovchiga topshirildi'):
        holat='Tayyor'
    elif str(c.execute("SELECT muddat_boshlanish_vaqt FROM buyurtmalar WHERE id=?",(oid,)).fetchone()[0] or '').strip():
        holat='Jarayonda'
    c.execute("UPDATE buyurtmalar SET holat=? WHERE id=?",(holat,oid))
    c.commit(); c.close()
    foiz=round(bajarildi*100/jami,1) if jami else 0
    return jsonify({"status":"ok","buyurtma_id":oid,"foiz":foiz,"holat":holat,"jami":jami,"bajarildi":bajarildi,"deadline":deadline_result})


@app.route("/api/buyurtma/<int:oid>/workflow")
def order_workflow(oid):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    if not order:
        c.close(); return jsonify({"message":"Buyurtma topilmadi"}),404
    stages=c.execute("SELECT * FROM buyurtma_bosqichlari WHERE buyurtma_id=? ORDER BY id",(oid,)).fetchall()
    deadline=_order_deadline(order)
    result=dict(order)
    result["deadline_text"]=deadline.strftime("%d.%m.%Y %H:%M") if deadline else "Hali boshlanmagan"
    result["payment_ready"]=_payment_ready(order)
    c.close()
    return jsonify({"order":result,"stages":[dict(x) for x in stages]})


@app.route("/api/buyurtma/<int:oid>/rahbar-tasdiq", methods=["POST"])
def order_manager_approve(oid):
    c=get_db(); order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    if not order:
        c.close(); return jsonify({"message":"Buyurtma topilmadi"}),404
    condition=str(order["tolov_sharti"] or "")
    if condition not in ("Avans talab qilinmaydi — ishonchli mijoz","Muddatli to‘lov","Shartnoma asosida to‘lov"):
        c.close(); return jsonify({"message":"Bu buyurtmada rahbar tasdig‘i talab qilinmaydi."}),400
    now_text=_tashkent_now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("UPDATE buyurtmalar SET rahbar_tasdiq=1,rahbar_tasdiq_vaqt=? WHERE id=?",(now_text,oid))
    _complete_order_stage(c,oid,"Avans talab qilinmaydi","Ishonchli mijoz uchun avans talab qilinmaydi")
    _complete_order_stage(c,oid,"Muddatli to‘lov tasdiqlandi","Rahbar tomonidan tasdiqlandi")
    _complete_order_stage(c,oid,"Shartnoma tasdiqlandi","Rahbar tomonidan tasdiqlandi")
    _complete_order_stage(c,oid,"Rahbar tasdiqladi","Rahbar tomonidan tasdiqlandi")
    deadline=_start_official_deadline(c,oid,"Rahbar tasdiqladi — rasmiy muddat boshlandi")
    c.commit(); c.close()
    return jsonify({"status":"ok","deadline":deadline})



# ---------- V5.1 PRO MODULLAR ----------
@app.route("/api/pro/bosqich-hodisa", methods=["POST"])
def pro_stage_event_add():
    d=jdata(); c=get_db()
    c.execute("""INSERT INTO buyurtma_bosqich_hodisalari
        (buyurtma_id,bosqich,boshlandi,tugadi,ishchi_id,izoh,media_havola)
        VALUES(?,?,?,?,?,?,?)""",
        (int(d["buyurtma_id"]),d.get("bosqich",""),d.get("boshlandi",""),
         d.get("tugadi",""),int(d["ishchi_id"]) if d.get("ishchi_id") else None,
         d.get("izoh",""),d.get("media_havola","")))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/media", methods=["POST"])
def pro_media_add():
    d=jdata(); c=get_db()
    c.execute("INSERT INTO buyurtma_media(buyurtma_id,turi,havola,izoh) VALUES(?,?,?,?)",
              (int(d["buyurtma_id"]),d.get("turi","Rasm"),d["havola"],d.get("izoh","")))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/tasdiq", methods=["POST"])
def pro_approval_add():
    d=jdata(); c=get_db()
    c.execute("""INSERT INTO buyurtma_tasdiqlari
        (buyurtma_id,turi,holat,izoh,tasdiqlagan,tasdiqlangan_vaqt)
        VALUES(?,?,?,?,?,?)""",
        (int(d["buyurtma_id"]),d.get("turi","Chizma"),d.get("holat","Tasdiqlandi"),
         d.get("izoh",""),d.get("tasdiqlagan",""),
         _tashkent_now().strftime("%Y-%m-%d %H:%M:%S")))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/baho", methods=["POST"])
def pro_rating_add():
    d=jdata(); c=get_db()
    vals=[max(1,min(5,int(d.get(k,5)))) for k in ["sifat","muddat","muomala","yetkazish","montaj"]]
    umumiy=round(sum(vals)/len(vals))
    c.execute("""INSERT INTO baholar
        (buyurtma_id,sifat,muddat,muomala,yetkazish,montaj,umumiy,izoh,rasm_havola,reklama_ruxsat)
        VALUES(?,?,?,?,?,?,?,?,?,?)""",
        (int(d["buyurtma_id"]),*vals,umumiy,d.get("izoh",""),d.get("rasm_havola",""),
         1 if d.get("reklama_ruxsat") else 0))
    c.commit(); c.close(); return jsonify({"status":"ok","umumiy":umumiy})

@app.route("/api/pro/servis", methods=["POST"])
def pro_service_add():
    d=jdata(); c=get_db()
    c.execute("""INSERT INTO servis_murojaatlari
        (buyurtma_id,turi,muammo,media_havola,holat,servis_sana,pullik)
        VALUES(?,?,?,?,?,?,?)""",
        (int(d["buyurtma_id"]),d.get("turi","Servis"),d["muammo"],
         d.get("media_havola",""),d.get("holat","Qabul qilindi"),
         d.get("servis_sana",""),1 if d.get("pullik") else 0))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/qoshimcha-ish", methods=["POST"])
def pro_extra_work_add():
    d=jdata(); c=get_db()
    c.execute("""INSERT INTO qoshimcha_ishlar
        (buyurtma_id,nomi,summa,qoshimcha_kun,mijoz_tasdiq,kim_kiritdi)
        VALUES(?,?,?,?,?,?)""",
        (int(d["buyurtma_id"]),d["nomi"],float(d.get("summa") or 0),
         int(d.get("qoshimcha_kun") or 0),1 if d.get("mijoz_tasdiq") else 0,
         d.get("kim_kiritdi","Rahbar")))
    c.execute("UPDATE buyurtmalar SET umumiy_narx=umumiy_narx+? WHERE id=?",
              (float(d.get("summa") or 0),int(d["buyurtma_id"])))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/rezerv", methods=["POST"])
def pro_stock_reserve():
    d=jdata(); c=get_db()
    c.execute("""INSERT INTO ombor_rezervlari(buyurtma_id,material_id,miqdor)
                 VALUES(?,?,?)""",
              (int(d["buyurtma_id"]),int(d["material_id"]),float(d.get("miqdor") or 0)))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/tannarx", methods=["POST"])
def pro_cost_save():
    d=jdata(); c=get_db()
    vals=[float(d.get(k) or 0) for k in ["material","ishchi","boyoq","oyna","furnitura","transport","elektr","tashqi_xizmat","boshqa"]]
    c.execute("""INSERT INTO buyurtma_tannarxi
        (buyurtma_id,material,ishchi,boyoq,oyna,furnitura,transport,elektr,tashqi_xizmat,boshqa)
        VALUES(?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(buyurtma_id) DO UPDATE SET
        material=excluded.material,ishchi=excluded.ishchi,boyoq=excluded.boyoq,
        oyna=excluded.oyna,furnitura=excluded.furnitura,transport=excluded.transport,
        elektr=excluded.elektr,tashqi_xizmat=excluded.tashqi_xizmat,boshqa=excluded.boshqa""",
        (int(d["buyurtma_id"]),*vals))
    c.commit(); c.close(); return jsonify({"status":"ok","jami":sum(vals)})

@app.route("/api/pro/xabar", methods=["POST"])
def pro_message_add():
    d=jdata(); c=get_db()
    c.execute("""INSERT INTO tizim_xabarlari
        (foydalanuvchi_turi,foydalanuvchi_id,buyurtma_id,mavzu,matn,kanal)
        VALUES(?,?,?,?,?,?)""",
        (d.get("foydalanuvchi_turi","Mijoz"),d.get("foydalanuvchi_id"),
         d.get("buyurtma_id"),d["mavzu"],d["matn"],d.get("kanal","Sayt")))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/pro/buyurtma/<int:oid>")
def pro_order_bundle(oid):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    stages=c.execute("SELECT * FROM buyurtma_bosqich_hodisalari WHERE buyurtma_id=? ORDER BY id",(oid,)).fetchall()
    media=c.execute("SELECT * FROM buyurtma_media WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    approvals=c.execute("SELECT * FROM buyurtma_tasdiqlari WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    payments=c.execute("SELECT * FROM buyurtma_tolovlari WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    services=c.execute("SELECT * FROM servis_murojaatlari WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    extras=c.execute("SELECT * FROM qoshimcha_ishlar WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    rating=c.execute("SELECT * FROM baholar WHERE buyurtma_id=? ORDER BY id DESC LIMIT 1",(oid,)).fetchone()
    cost=c.execute("SELECT * FROM buyurtma_tannarxi WHERE buyurtma_id=?",(oid,)).fetchone()
    c.close()
    return jsonify({
        "order":dict(order) if order else None,
        "stages":[dict(x) for x in stages],
        "media":[dict(x) for x in media],
        "approvals":[dict(x) for x in approvals],
        "payments":[dict(x) for x in payments],
        "services":[dict(x) for x in services],
        "extras":[dict(x) for x in extras],
        "rating":dict(rating) if rating else None,
        "cost":dict(cost) if cost else None
    })

# ---------- OMBOR ----------
@app.route("/api/ombor", methods=["GET"])
def stock_get():
    c=get_db(); rows=c.execute("SELECT * FROM ombor ORDER BY kategoriya,nomi").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/ombor-harakat", methods=["POST"])
def stock_move():
    d=jdata()
    try:
        qty=float(d.get("miqdor") or 0); typ=d.get("turi","Kirim")
        delta=qty if typ=="Kirim" else -qty
        c=get_db()
        c.execute("UPDATE ombor SET qoldiq=qoldiq+? WHERE id=?",(delta,int(d["material_id"])))
        c.execute("""INSERT INTO ombor_harakat(material_id,sana,turi,miqdor,buyurtma_kodi,izoh)
        VALUES(?,?,?,?,?,?)""",(int(d["material_id"]),d["sana"],typ,qty,d.get("buyurtma_kodi",""),d.get("izoh","")))
        c.commit(); c.close(); return jsonify({"status":"ok"})
    except Exception as e:
        return jsonify({"message":str(e)}),400


# ---------- SAFAR ----------
@app.route("/api/safarlar", methods=["GET","POST"])
def trips():
    if request.method=="POST":
        d=jdata(); c=get_db()
        c.execute("""INSERT INTO safarlar(ishchi_id,sana,mashina,qayerdan,qayerga,masofa_km,sabab,yonilgi,xarajat)
        VALUES(?,?,?,?,?,?,?,?,?)""",(int(d["ishchi_id"]),d["sana"],d.get("mashina",""),d.get("qayerdan",""),
        d.get("qayerga",""),float(d.get("masofa_km") or 0),d.get("sabab",""),
        float(d.get("yonilgi") or 0),float(d.get("xarajat") or 0)))
        c.commit(); c.close(); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("""SELECT s.*,i.ism,i.familiya FROM safarlar s
    JOIN ishchilar i ON i.id=s.ishchi_id ORDER BY sana DESC,s.id DESC LIMIT 200""").fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


# ---------- QO‘SHIMCHA MODULLAR ----------
@app.route("/api/xarajatlar", methods=["GET","POST"])
def expenses():
    if request.method=="POST":
        d=jdata(); c=get_db()
        c.execute("INSERT INTO xarajatlar(sana,kategoriya,miqdor,tavsifi,buyurtma_kodi,tolov_usuli,xarajat_nomi,kimga_berildi,chek_havola) VALUES(?,?,?,?,?,?,?,?,?)",
                  (d["sana"],d.get("kategoriya","Boshqa"),float(d.get("miqdor") or 0),
                   d.get("tavsifi",""),d.get("buyurtma_kodi",""),d.get("tolov_usuli","Naqd"),
                   d.get("xarajat_nomi",""),d.get("kimga_berildi",""),d.get("chek_havola","")))
        c.commit(); c.close(); log_action("Xarajat qo‘shildi", d.get("kategoriya","")); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("SELECT * FROM xarajatlar ORDER BY sana DESC,id DESC LIMIT 300").fetchall(); c.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/bonuslar", methods=["GET","POST"])
def bonuses():
    if request.method=="POST":
        d=jdata(); c=get_db(); c.execute("INSERT INTO bonuslar(ishchi_id,sana,miqdor,sababi) VALUES(?,?,?,?)",
            (int(d["ishchi_id"]),d["sana"],float(d.get("miqdor") or 0),d.get("sababi","")))
        c.commit(); c.close(); log_action("Bonus qo‘shildi", str(d.get("ishchi_id"))); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("""SELECT b.*,i.ism,i.familiya FROM bonuslar b JOIN ishchilar i ON i.id=b.ishchi_id
        ORDER BY sana DESC,b.id DESC LIMIT 200""").fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/ishchi-holatlari", methods=["GET","POST"])
def worker_statuses():
    if request.method=="POST":
        d=jdata(); c=get_db(); c.execute("INSERT INTO ishchi_holatlari(ishchi_id,sana,turi,izoh) VALUES(?,?,?,?)",
            (int(d["ishchi_id"]),d["sana"],d.get("turi","Dam olish"),d.get("izoh","")))
        c.commit(); c.close(); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("""SELECT h.*,i.ism,i.familiya FROM ishchi_holatlari h JOIN ishchilar i ON i.id=h.ishchi_id
        ORDER BY sana DESC,h.id DESC LIMIT 200""").fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/buyurtma/<int:oid>/tolovlar", methods=["GET","POST"])
def order_payments(oid):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    if not order:
        c.close(); return jsonify({"message":"Buyurtma topilmadi"}),404
    if request.method=="POST":
        try:
            d=jdata()
            received=float(d.get("qabul_qilingan_summa",d.get("miqdor")) or 0)
            payment_currency=_currency_code(d.get("tolov_valyutasi") or order["valyuta"] or "UZS")
            order_currency=_currency_code(order["valyuta"] or "UZS")
            order_amount,uzs_equivalent,rate=_payment_conversion(
                order_currency,payment_currency,received,float(d.get("kurs") or 0)
            )
            c.execute("""INSERT INTO buyurtma_tolovlari
                (buyurtma_id,sana,miqdor,turi,izoh,tolov_valyutasi,kurs,
                 qabul_qilingan_summa,buyurtma_summa,uzs_ekvivalent)
                VALUES(?,?,?,?,?,?,?,?,?,?)""",
                (oid,d.get("sana") or _tashkent_today(),order_amount,d.get("turi","To‘lov"),
                 d.get("izoh",""),payment_currency,rate,received,order_amount,uzs_equivalent))
            c.execute("""UPDATE buyurtmalar SET oldindan_tolov=oldindan_tolov+?,
                oxirgi_kurs=CASE WHEN ?>1 THEN ? ELSE oxirgi_kurs END WHERE id=?""",
                (order_amount,rate,rate,oid))
            updated=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
            deadline_result={}
            if order_amount>0 and _payment_ready(updated):
                condition=str(updated["tolov_sharti"] or "Avans majburiy")
                pay_stage="Qisman avans olindi" if condition=="Qisman avans" else "Avans olindi"
                _complete_order_stage(c,oid,pay_stage,"Avans summasi talabga yetdi")
                deadline_result=_start_official_deadline(c,oid,"Avans qabul qilindi — rasmiy muddat boshlandi")
            c.commit(); c.close()
            return jsonify({"status":"ok","deadline":deadline_result,
                            "buyurtma_summa":order_amount,"uzs_ekvivalent":uzs_equivalent,
                            "buyurtma_valyutasi":order_currency})
        except Exception as e:
            c.rollback(); c.close(); return jsonify({"message":str(e)}),400
    rows=c.execute("SELECT * FROM buyurtma_tolovlari WHERE buyurtma_id=? ORDER BY sana DESC,id DESC",(oid,)).fetchall(); c.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/tayyor-mahsulot", methods=["GET","POST"])
def finished_goods():
    if request.method=="POST":
        d=jdata(); c=get_db(); c.execute("INSERT INTO tayyor_mahsulot(nomi,kodi,rang,miqdor,birlik,narx,izoh) VALUES(?,?,?,?,?,?,?)",
            (d["nomi"],d.get("kodi",""),d.get("rang",""),float(d.get("miqdor") or 0),d.get("birlik","dona"),float(d.get("narx") or 0),d.get("izoh","")))
        c.commit(); c.close(); return jsonify({"status":"ok"})
    c=get_db(); rows=c.execute("SELECT * FROM tayyor_mahsulot ORDER BY id DESC").fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/moliyaviy-xulosa")
def finance_summary():
    start=request.args.get("start") or "1900-01-01"; end=request.args.get("end") or "2999-12-31"
    c=get_db()
    income=c.execute("SELECT COALESCE(SUM(uzs_ekvivalent),0) FROM buyurtma_tolovlari WHERE sana BETWEEN ? AND ?",(start,end)).fetchone()[0]
    expense=c.execute("SELECT COALESCE(SUM(miqdor),0) FROM xarajatlar WHERE sana BETWEEN ? AND ?",(start,end)).fetchone()[0]
    salary=c.execute("SELECT COALESCE(SUM(miqdor),0) FROM tolovlar WHERE sana BETWEEN ? AND ?",(start,end)).fetchone()[0]
    bonus=c.execute("SELECT COALESCE(SUM(miqdor),0) FROM bonuslar WHERE sana BETWEEN ? AND ?",(start,end)).fetchone()[0]
    c.close(); return jsonify({"kirim":income,"xarajat":expense,"ishchi_tolov":salary,"bonus":bonus,"sof_foyda":income-expense-salary-bonus})

@app.route("/api/buyurtma-progress/<int:oid>")
def order_progress(oid):
    c=get_db(); row=c.execute("SELECT COUNT(*) jami,SUM(bajarildi) bajarildi FROM buyurtma_bosqichlari WHERE buyurtma_id=?",(oid,)).fetchone(); c.close()
    jami=row['jami'] or 0; done=row['bajarildi'] or 0; return jsonify({"foiz": round(done*100/jami,1) if jami else 0})

def _order_bundle(oid):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    stages=c.execute("SELECT * FROM buyurtma_bosqichlari WHERE buyurtma_id=? ORDER BY id",(oid,)).fetchall()
    pays=c.execute("SELECT * FROM buyurtma_tolovlari WHERE buyurtma_id=? ORDER BY sana,id",(oid,)).fetchall()
    c.close()
    return order,stages,pays

@app.route("/api/buyurtma/<int:oid>/mijoz-kartasi", methods=["GET","POST"])
def order_customer_card(oid):
    """Rahbar uchun mijozga ko‘rinadigan buyurtma kartasini boshqarish."""
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    if not order:
        c.close(); return jsonify({"message":"Buyurtma topilmadi"}),404
    if request.method=="POST":
        d=jdata()
        old_date=str(order["taxminiy_sana"] or "")
        old_time=_safe_time(order["taxminiy_vaqt"] if "taxminiy_vaqt" in order.keys() else "18:00")
        new_date=str(d.get("taxminiy_sana",old_date) or "")
        new_time=_safe_time(d.get("taxminiy_vaqt",old_time))
        fields={
            "taxminiy_sana":new_date,
            "taxminiy_vaqt":new_time,
            "rang":d.get("rang",order["rang"] or ""),
            "rang_kodi":d.get("rang_kodi",order["rang_kodi"] or ""),
            "material":d.get("material",order["material"] or ""),
            "yaltiroqlik":d.get("yaltiroqlik",order["yaltiroqlik"] or ""),
            "mijozga_izoh":d.get("mijozga_izoh",order["mijozga_izoh"] or ""),
            "kechikish_sababi":d.get("kechikish_sababi",order["kechikish_sababi"] or ""),
            "aloqa_telefon":d.get("aloqa_telefon",order["aloqa_telefon"] or ""),
            "masul_xodim":d.get("masul_xodim",order["masul_xodim"] or ""),
            "avans_talab":float(d.get("avans_talab",order["avans_talab"] or 0) or 0),
            "avans_muddat_sana":d.get("avans_muddat_sana",order["avans_muddat_sana"] or ""),
            "taklif_amal_sana":d.get("taklif_amal_sana",order["taklif_amal_sana"] or ""),
            "ishlab_chiqarish_kun":int(d.get("ishlab_chiqarish_kun",order["ishlab_chiqarish_kun"] or 0) or 0),
            "chizma_versiya":int(d.get("chizma_versiya",order["chizma_versiya"] or 1) or 1),
            "bepul_ozgarish_limit":int(d.get("bepul_ozgarish_limit",order["bepul_ozgarish_limit"] or 2) or 2),
        }
        c.execute("""UPDATE buyurtmalar SET taxminiy_sana=?,taxminiy_vaqt=?,rang=?,rang_kodi=?,
          material=?,yaltiroqlik=?,mijozga_izoh=?,kechikish_sababi=?,aloqa_telefon=?,masul_xodim=?,
          avans_talab=?,avans_muddat_sana=?,taklif_amal_sana=?,ishlab_chiqarish_kun=?,chizma_versiya=?,bepul_ozgarish_limit=?
          WHERE id=?""",(*fields.values(),oid))
        if (old_date,old_time)!=(new_date,new_time):
            c.execute("""INSERT INTO buyurtma_muddat_tarixi
              (buyurtma_id,eski_sana,eski_vaqt,yangi_sana,yangi_vaqt,sabab)
              VALUES(?,?,?,?,?,?)""",(oid,old_date,old_time,new_date,new_time,d.get("muddat_ozgarish_sababi","") or fields["kechikish_sababi"]))
        if str(d.get("media_havola","")).strip():
            c.execute("INSERT INTO buyurtma_media(buyurtma_id,turi,havola,izoh) VALUES(?,?,?,?)",
                      (oid,d.get("media_turi","Rasm"),str(d.get("media_havola")).strip(),d.get("media_izoh","")))
        if str(d.get("tasdiq_turi","")).strip():
            c.execute("INSERT INTO buyurtma_tasdiqlari(buyurtma_id,turi,holat,izoh) VALUES(?,?,?,?)",
                      (oid,str(d.get("tasdiq_turi")).strip(),"Kutilmoqda",d.get("tasdiq_izoh","")))
        c.commit()
        log_action("mijoz_kartasi_yangilandi",f"buyurtma_id={oid}; kod={order['kod']}")
        order=c.execute("SELECT * FROM buyurtmalar WHERE id=?",(oid,)).fetchone()
    media=c.execute("SELECT * FROM buyurtma_media WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    approvals=c.execute("SELECT * FROM buyurtma_tasdiqlari WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    history=c.execute("SELECT * FROM buyurtma_muddat_tarixi WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    c.close()
    return jsonify({"order":dict(order),"media":[dict(x) for x in media],
                    "approvals":[dict(x) for x in approvals],"deadline_history":[dict(x) for x in history]})


@app.route("/kuzatuv/<token>/tasdiq", methods=["POST"])
def public_track_approval(token):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE tracking_token=?",(token,)).fetchone()
    if not order:
        c.close(); return "Buyurtma topilmadi",404
    try:
        approval_id=int(request.form.get("approval_id") or 0)
    except ValueError:
        approval_id=0
    action=request.form.get("action","Tasdiqlandi")
    holat="Savolim bor" if action=="Savolim bor" else "Tasdiqlandi"
    row=c.execute("SELECT id,turi FROM buyurtma_tasdiqlari WHERE id=? AND buyurtma_id=?",(approval_id,order["id"])).fetchone()
    if row:
        now_text=_tashkent_now().strftime("%Y-%m-%d %H:%M:%S")
        c.execute("""UPDATE buyurtma_tasdiqlari SET holat=?,tasdiqlagan=?,tasdiqlangan_vaqt=?
          WHERE id=?""",(holat,order["mijoz"],now_text,approval_id))
        if holat=="Tasdiqlandi" and str(row["turi"] or "") in ("Chizma tasdig‘i","Narx tasdig‘i","Hammasini tasdiqlash"):
            _complete_order_stage(c,int(order["id"]),"Mijoz tasdiqladi",f"Mijoz {row['turi'].lower()}ni tasdiqladi")
            condition=str(order["tolov_sharti"] or "Avans majburiy") if "tolov_sharti" in order.keys() else "Avans majburiy"
            next_status="Rahbar tasdiqlashi kutilmoqda" if condition in ("Avans talab qilinmaydi — ishonchli mijoz","Muddatli to‘lov","Shartnoma asosida to‘lov") else "Avans kutilmoqda"
            c.execute("UPDATE buyurtmalar SET holat=? WHERE id=?",(next_status,int(order["id"])))
        c.commit()
    c.close()
    return redirect(url_for("public_track",token=token,xabar="Tasdiqingiz saqlandi" if holat=="Tasdiqlandi" else "Savolingiz menejerga yuborildi"))


@app.route("/kuzatuv/<token>")
def public_track(token):
    c=get_db()
    order=c.execute("SELECT * FROM buyurtmalar WHERE tracking_token=?",(token,)).fetchone()
    if not order:
        c.close(); return "Buyurtma topilmadi",404
    oid=int(order["id"])
    stages=c.execute("SELECT * FROM buyurtma_bosqichlari WHERE buyurtma_id=? ORDER BY id",(oid,)).fetchall()
    delivery=c.execute(
        "SELECT y.*, i.ism AS haydovchi_ism, i.familiya AS haydovchi_familiya, i.telefon AS haydovchi_telefon "
        "FROM yetkazishlar y LEFT JOIN ishchilar i ON i.id=y.haydovchi_id "
        "WHERE y.buyurtma_id=? ORDER BY y.id DESC LIMIT 1",(oid,)
    ).fetchone()
    media_rows=c.execute("SELECT * FROM buyurtma_media WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    approvals=c.execute("SELECT * FROM buyurtma_tasdiqlari WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall()
    pays=c.execute("SELECT * FROM buyurtma_tolovlari WHERE buyurtma_id=? ORDER BY sana DESC,id DESC LIMIT 12",(oid,)).fetchall()
    deadline_history=c.execute("SELECT * FROM buyurtma_muddat_tarixi WHERE buyurtma_id=? ORDER BY id DESC LIMIT 10",(oid,)).fetchall()
    c.close()

    done=sum(int(x["bajarildi"] or 0) for x in stages)
    pct=round(done*100/len(stages),1) if stages else 0
    unfinished=next((x for x in stages if not int(x["bajarildi"] or 0)),None)
    current_stage=(unfinished["bosqich"] if unfinished else ("Buyurtma tayyor" if pct>=100 else order["holat"]))
    ready=str(order["holat"] or "") in {"Tayyor","Yetkazishga tayyor","Yetkazildi","Yopildi"} or pct>=100
    deadline=_order_deadline(order)
    deadline_ms=int(deadline.timestamp()*1000) if deadline else 0
    workflow_started=bool(str(order["muddat_boshlanish_vaqt"] or "").strip()) if "muddat_boshlanish_vaqt" in order.keys() else True
    payment_condition=str(order["tolov_sharti"] or "Avans majburiy") if "tolov_sharti" in order.keys() else "Avans majburiy"
    if deadline:
        deadline_text=deadline.strftime("%d.%m.%Y, %H:%M")
        countdown_wait=""
    elif str(order["muddat_tartibi"] or "")=="Tasdiqdan keyin" if "muddat_tartibi" in order.keys() else False:
        countdown_wait="Rahbar tasdig‘idan keyin boshlanadi" if payment_condition in ("Avans talab qilinmaydi — ishonchli mijoz","Muddatli to‘lov","Shartnoma asosida to‘lov") else "Avans olingandan keyin boshlanadi"
        deadline_text=countdown_wait
    else:
        countdown_wait="Muddat belgilanmagan"
        deadline_text="Belgilanmagan"
    total=float(order["umumiy_narx"] or 0)
    paid=float(order["oldindan_tolov"] or 0)
    remaining=max(0,total-paid)
    currency=_currency_code(order["valyuta"] if "valyuta" in order.keys() else "UZS")
    contact=str(order["aloqa_telefon"] or os.environ.get("MEBEL360_PUBLIC_PHONE","")).strip()
    contact_digits="".join(ch for ch in contact if ch.isdigit())
    if contact_digits.startswith("0") and len(contact_digits)==10:
        contact_digits="998"+contact_digits[1:]
    whatsapp=(f"https://wa.me/{contact_digits}?text=Assalomu%20alaykum,%20{order['kod']}%20buyurtmam%20bo‘yicha%20savolim%20bor.") if contact_digits else ""
    rang_code=str(order["rang_kodi"] or "").strip()
    rang_css=rang_code if re.fullmatch(r"#[0-9a-fA-F]{6}",rang_code) else "#dbeafe"
    media=[]
    for row in media_rows:
        item=dict(row); url=str(item.get("havola") or "")
        low=url.lower().split("?")[0]
        if low.endswith((".png",".jpg",".jpeg",".webp",".gif")):
            item["kind"]="image"
        elif low.endswith((".mp4",".webm",".mov")):
            item["kind"]="video"
        else:
            item["kind"]="link"
        media.append(item)
    timeline=[]
    if order["created_at"]:
        timeline.append({"time":str(order["created_at"]),"title":"Buyurtma yaratildi","detail":str(order["mahsulot"] or "")})
    for s in stages:
        if int(s["bajarildi"] or 0):
            timeline.append({"time":str(s["tugash_vaqti"] or s["boshlanish_vaqti"] or ""),"title":str(s["bosqich"]),"detail":str(s["izoh"] or "Bosqich yakunlandi")})
    for pay in pays:
        pay_currency=_currency_code(pay["tolov_valyutasi"] if "tolov_valyutasi" in pay.keys() else currency)
        received=pay["qabul_qilingan_summa"] if "qabul_qilingan_summa" in pay.keys() and pay["qabul_qilingan_summa"] else pay["miqdor"]
        detail=_currency_money(received,pay_currency)
        if currency=="USD" and pay_currency=="UZS" and float(pay["kurs"] or 0)>0:
            detail+=f" · kurs {_money(pay['kurs'])} · {_currency_money(pay['miqdor'],'USD')} hisoblandi"
        timeline.append({"time":str(pay["sana"] or ""),"title":"To‘lov qabul qilindi","detail":detail})
    for h in deadline_history:
        timeline.append({"time":str(h["created_at"] or ""),"title":"Muddat yangilandi","detail":f"{h['yangi_sana']} {h['yangi_vaqt']} · {h['sabab'] or 'Sabab ko‘rsatilmagan'}"})
    timeline.sort(key=lambda x:x["time"],reverse=True)

    html=r"""<!doctype html><html lang="uz"><head>
    <meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
    <meta http-equiv="refresh" content="45"><title>{{o['kod']}} — Buyurtma kuzatuvi</title>
    <style>
    :root{--blue:#0757a6;--green:#10a952;--ink:#102038;--muted:#64748b;--bg:#eef4f8;--card:#fff}
    *{box-sizing:border-box}body{margin:0;font-family:Arial,sans-serif;background:linear-gradient(180deg,#e8f4ff 0,#f4f7fa 260px);color:var(--ink)}
    .top{background:linear-gradient(135deg,#063f7d,#0f6dbb 58%,#10a952);color:#fff;padding:20px 16px 72px}.topin{max-width:980px;margin:auto;display:flex;align-items:center;gap:14px}.logo{width:90px;height:68px;object-fit:contain;background:#fff;border-radius:15px;padding:5px;box-shadow:0 10px 28px #001b3d55}.title h1{margin:0;font-size:27px}.title p{margin:5px 0 0;opacity:.9}.wrap{max-width:980px;margin:-52px auto 40px;padding:0 14px}.hero,.card{background:var(--card);border-radius:20px;box-shadow:0 13px 35px #10203818;border:1px solid #dbe6ef}.hero{padding:20px}.headrow{display:flex;justify-content:space-between;gap:14px;align-items:flex-start}.code{font-size:26px;font-weight:900}.product{color:var(--muted);margin-top:5px}.status-pill{background:#e8f5ff;color:#0757a6;padding:9px 12px;border-radius:999px;font-weight:900;white-space:nowrap}.metrics{display:grid;grid-template-columns:repeat(3,1fr);gap:10px;margin-top:18px}.metric{border-radius:16px;padding:16px;background:#f5f9fc;text-align:center}.metric small{display:block;color:var(--muted);font-weight:800;margin-bottom:8px}.metric strong{font-size:22px}.countdown{font-variant-numeric:tabular-nums}.countdown.green{color:#15803d}.countdown.yellow{color:#a16207}.countdown.orange{color:#c2410c}.countdown.red{color:#dc2626}.bar{height:14px;background:#dfe7ee;border-radius:999px;overflow:hidden;margin-top:14px}.fill{height:100%;background:linear-gradient(90deg,#0757a6,#16a34a);border-radius:999px}.grid{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-top:14px}.card{padding:18px}.card h3{margin:0 0 14px;font-size:18px}.facts{display:grid;grid-template-columns:1fr 1fr;gap:10px}.fact{background:#f7fafc;border-radius:13px;padding:12px}.fact small{display:block;color:var(--muted);margin-bottom:5px}.fact b{font-size:15px}.swatch{display:inline-block;width:18px;height:18px;border-radius:50%;vertical-align:middle;border:2px solid #fff;box-shadow:0 0 0 1px #94a3b8;margin-right:6px}.money-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:8px}.money-box{padding:13px;border-radius:13px;background:#f8fafc}.money-box small{display:block;color:var(--muted);margin-bottom:5px}.money-box b{font-size:17px}.balance{color:#dc2626}.notice{margin-top:12px;padding:13px;border-radius:13px;background:#fff7ed;border-left:5px solid #f97316}.customer-note{background:#ecfdf5;border-left-color:#16a34a}.stage-list{position:relative}.stage{display:grid;grid-template-columns:32px 1fr;gap:10px;padding:10px 0}.dot{width:28px;height:28px;border-radius:50%;display:grid;place-items:center;background:#e2e8f0;color:#64748b;font-weight:900}.stage.done .dot{background:#dcfce7;color:#15803d}.stage.current .dot{background:#dbeafe;color:#1d4ed8;box-shadow:0 0 0 5px #dbeafe66}.stage b{display:block}.stage small{color:var(--muted)}.delivery{background:#eff6ff;border:1px solid #bfdbfe}.delivery-status{font-size:20px;font-weight:900;color:#1d4ed8}.media-grid{display:grid;grid-template-columns:repeat(2,1fr);gap:10px}.media-item{border:1px solid #dbe5ee;border-radius:14px;overflow:hidden;background:#f8fafc}.media-item img,.media-item video{display:block;width:100%;height:190px;object-fit:cover}.media-caption{padding:10px}.approval{border:1px solid #dbe5ee;border-radius:14px;padding:13px;margin-top:9px}.approval.pending{border-color:#facc15;background:#fffbeb}.approval.ok{border-color:#86efac;background:#f0fdf4}.approval.ask{border-color:#fdba74;background:#fff7ed}.actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px}.btn{display:inline-flex;align-items:center;justify-content:center;border:0;border-radius:11px;padding:10px 14px;text-decoration:none;background:#0757a6;color:#fff;font-weight:900;cursor:pointer}.btn.green{background:#16a34a}.btn.orange{background:#ea580c}.timeline{border-left:3px solid #dbe5ee;margin-left:8px;padding-left:17px}.event{position:relative;padding:0 0 16px}.event:before{content:"";position:absolute;width:11px;height:11px;border-radius:50%;background:#0f6dbb;left:-24px;top:4px}.event b{display:block}.event small{color:var(--muted)}.empty{color:var(--muted);padding:8px 0}.alert{margin-bottom:13px;background:#dcfce7;border:1px solid #86efac;color:#166534;padding:12px;border-radius:13px;font-weight:800}.footer{text-align:center;color:#64748b;padding:26px 0 0;font-size:13px}
    @media(max-width:760px){.top{padding-bottom:64px}.topin{align-items:flex-start}.logo{width:75px;height:58px}.title h1{font-size:22px}.wrap{margin-top:-48px}.headrow{display:block}.status-pill{display:inline-block;margin-top:10px}.metrics,.grid{grid-template-columns:1fr}.metric{text-align:left}.facts{grid-template-columns:1fr 1fr}.money-grid{grid-template-columns:1fr}.media-grid{grid-template-columns:1fr}.media-item img,.media-item video{height:220px}}
    @media(max-width:430px){.facts{grid-template-columns:1fr}.metric strong{font-size:19px}.code{font-size:22px}}
    </style></head><body>
    <div class="top"><div class="topin"><img class="logo" src="/static/mebel360-logo.png?v=20260722" alt="Mebel360°"><div class="title"><h1>Mebel360°</h1><p>Buyurtmangizni 360° kuzating</p></div></div></div>
    <main class="wrap">
      {% if message %}<div class="alert">✅ {{message}}</div>{% endif %}
      <section class="hero">
        <div class="headrow"><div><div class="code">{{o['kod']}}</div><div class="product">{{o['mahsulot'] or 'Mebel buyurtmasi'}} · {{o['mijoz']}}</div></div><div class="status-pill">{{o['holat']}}</div></div>
        <div class="metrics">
          <div class="metric"><small>BAJARILDI</small><strong>{{pct}}%</strong></div>
          <div class="metric"><small>QOLGAN VAQT</small><strong id="countdown" class="countdown">Hisoblanmoqda…</strong></div>
          <div class="metric"><small>QOLDIQ TO‘LOV</small><strong class="balance">{{format_money(remaining,currency)}}</strong></div>
        </div>
        <div class="bar"><div class="fill" style="width:{{pct}}%"></div></div>
      </section>
      <div class="grid">
        <section class="card"><h3>📋 Buyurtma ma’lumotlari</h3><div class="facts">
          <div class="fact"><small>Hozirgi bosqich</small><b>{{current_stage}}</b></div><div class="fact"><small>Tayyor bo‘lish muddati</small><b>{{deadline_text}}</b></div>
          <div class="fact"><small>Buyurtma turi</small><b>{{o['buyurtma_turi'] or 'To‘liq mebel'}}</b></div><div class="fact"><small>To‘lov sharti</small><b>{{o['tolov_sharti'] or 'Avans majburiy'}}</b></div>
          <div class="fact"><small>Rang</small><b><span class="swatch" style="background:{{rang_css}}"></span>{{o['rang'] or 'Belgilanmagan'}}{% if o['rang_kodi'] %} · {{o['rang_kodi']}}{% endif %}</b></div>
          <div class="fact"><small>Material</small><b>{{o['material'] or 'Belgilanmagan'}}</b></div><div class="fact"><small>Yaltiroqlik</small><b>{{o['yaltiroqlik'] or 'Belgilanmagan'}}</b></div><div class="fact"><small>Mas’ul xodim</small><b>{{o['masul_xodim'] or 'Belgilanmagan'}}</b></div>
        </div>{% if o['mijozga_izoh'] %}<div class="notice customer-note"><b>Mebel360° izohi:</b><br>{{o['mijozga_izoh']}}</div>{% endif %}{% if o['kechikish_sababi'] %}<div class="notice"><b>Muddat bo‘yicha izoh:</b><br>{{o['kechikish_sababi']}}</div>{% endif %}</section>
        <section class="card"><h3>💳 To‘lov holati</h3><div class="money-grid"><div class="money-box"><small>Umumiy narx</small><b>{{format_money(total,currency)}}</b></div><div class="money-box"><small>To‘langan</small><b style="color:#15803d">{{format_money(paid,currency)}}</b></div><div class="money-box"><small>Qoldiq</small><b class="balance">{{format_money(remaining,currency)}}</b></div></div>
        {% if o['tolov_sharti']=='Avans talab qilinmaydi — ishonchli mijoz' %}<div class="notice customer-note">🤝 Ishonchli mijoz: avans talab qilinmaydi. Rasmiy muddat rahbar tasdig‘idan keyin boshlanadi.</div>{% elif remaining>0 %}<div class="notice">Buyurtma tayyor bo‘lishidan oldin <b>{{format_money(remaining,currency)}}</b> qoldiq to‘lov mavjud.</div>{% else %}<div class="notice customer-note">✅ Buyurtma to‘lovi to‘liq amalga oshirilgan.</div>{% endif %}{% if o['taklif_amal_sana'] %}<div class="notice"><b>Narx taklifi amal qiladi:</b> {{o['taklif_amal_sana']}} gacha.</div>{% endif %}{% if o['avans_muddat_sana'] and not workflow_started %}<div class="notice"><b>Avans muddati:</b> {{o['avans_muddat_sana']}} gacha.</div>{% endif %}{% if contact %}<div class="actions"><a class="btn" href="tel:{{contact}}">📞 Qo‘ng‘iroq</a>{% if whatsapp %}<a class="btn green" href="{{whatsapp}}" target="_blank">💬 Savol berish</a>{% endif %}</div>{% endif %}</section>
      </div>
      <div class="grid">
        <section class="card"><h3>🛠 Buyurtma bosqichlari</h3><div class="stage-list">{% for s in stages %}<div class="stage {{'done' if s['bajarildi'] else 'current' if s['bosqich']==current_stage else ''}}"><div class="dot">{{'✓' if s['bajarildi'] else loop.index}}</div><div><b>{{s['bosqich']}}</b><small>{% if s['bajarildi'] %}Yakunlandi{% if s['tugash_vaqti'] %} · {{s['tugash_vaqti']}}{% endif %}{% else %}Navbatda{% endif %}{% if s['izoh'] %}<br>{{s['izoh']}}{% endif %}</small></div></div>{% endfor %}</div></section>
        <section class="card delivery"><h3>🚚 Yetkazib berish</h3>{% if delivery %}{% set h=delivery['holat'] or 'Rejalashtirilgan' %}<div class="delivery-status">{% if h in ['Yo‘lga chiqdim','Yo‘lga chiqdi'] %}🚚 Shofyor yo‘lda{% elif h in ['Yetib keldim','Yetib keldi'] %}📍 Shofyor yetib keldi{% elif h in ['Yetkazib berdim','Yetkazib berildi','Yetkazildi'] %}✅ Buyurtma yetkazib berildi{% else %}🕒 Yetkazish rejalashtirilgan{% endif %}</div><p><b>Shofyor:</b> {{(delivery['haydovchi_ism'] or '')+' '+(delivery['haydovchi_familiya'] or '')}}</p><p><b>Mashina:</b> {{delivery['mashina'] or 'Ko‘rsatilmagan'}}</p>{% if delivery['yolga_chiqdi'] %}<p><b>Yo‘lga chiqdi:</b> {{delivery['yolga_chiqdi']}}</p>{% endif %}{% if delivery['yetib_keldi'] %}<p><b>Yetib keldi:</b> {{delivery['yetib_keldi']}}</p>{% endif %}{% if delivery['topshirildi'] %}<p><b>Topshirildi:</b> {{delivery['topshirildi']}}</p>{% endif %}{% else %}<div class="empty">Hali yetkazish rejalashtirilmagan.</div>{% endif %}</section>
      </div>
      {% if media %}<section class="card" style="margin-top:14px"><h3>📷 Ish jarayonidan rasm va videolar</h3><div class="media-grid">{% for m in media %}<div class="media-item">{% if m.kind=='image' %}<a href="{{m.havola}}" target="_blank"><img src="{{m.havola}}" alt="Buyurtma rasmi"></a>{% elif m.kind=='video' %}<video controls preload="metadata" src="{{m.havola}}"></video>{% else %}<div class="media-caption"><a class="btn" href="{{m.havola}}" target="_blank">🔗 Faylni ochish</a></div>{% endif %}<div class="media-caption"><b>{{m.turi}}</b>{% if m.izoh %}<div>{{m.izoh}}</div>{% endif %}<small>{{m.created_at}}</small></div></div>{% endfor %}</div></section>{% endif %}
      {% if approvals %}<section class="card" style="margin-top:14px"><h3>✅ Mijoz tasdiqlashi</h3>{% for a in approvals %}<div class="approval {{'pending' if a['holat']=='Kutilmoqda' else 'ask' if a['holat']=='Savolim bor' else 'ok'}}"><b>{{a['turi']}}</b><p>{{a['izoh'] or 'Tasdiqlashingiz kutilmoqda.'}}</p><div><b>Holat:</b> {{a['holat']}}</div>{% if a['holat']=='Kutilmoqda' %}<form method="post" action="/kuzatuv/{{o['tracking_token']}}/tasdiq"><input type="hidden" name="approval_id" value="{{a['id']}}"><div class="actions"><button class="btn green" name="action" value="Tasdiqlandi">Tasdiqlayman</button><button class="btn orange" name="action" value="Savolim bor">Savolim bor</button></div></form>{% endif %}</div>{% endfor %}</section>{% endif %}
      <section class="card" style="margin-top:14px"><h3>🕘 O‘zgarishlar tarixi</h3>{% if timeline %}<div class="timeline">{% for e in timeline[:20] %}<div class="event"><b>{{e.title}}</b><div>{{e.detail}}</div><small>{{e.time}}</small></div>{% endfor %}</div>{% else %}<div class="empty">Hozircha tarix mavjud emas.</div>{% endif %}</section><div class="footer">Ma’lumotlar avtomatik yangilanadi · Mebel360°</div>
    </main>
    <script>(function(){const el=document.getElementById('countdown'),deadline={{deadline_ms}},ready={{'true' if ready else 'false'}},wait={{countdown_wait|tojson}};function pad(n){return String(n).padStart(2,'0')}function tick(){if(!el)return;if(ready){el.textContent='Buyurtmangiz tayyor';el.className='countdown green';return}if(!deadline){el.textContent=wait||'Muddat belgilanmagan';el.className='countdown yellow';return}const diff=deadline-Date.now(),abs=Math.abs(diff),days=Math.floor(abs/86400000),hours=Math.floor(abs%86400000/3600000),mins=Math.floor(abs%3600000/60000),secs=Math.floor(abs%60000/1000),value=`${days} kun ${pad(hours)}:${pad(mins)}:${pad(secs)}`;if(diff<0){el.textContent=value+' kechikdi';el.className='countdown red'}else if(diff<86400000){el.textContent=value+' qoldi';el.className='countdown orange'}else if(diff<3*86400000){el.textContent=value+' qoldi';el.className='countdown yellow'}else{el.textContent=value+' qoldi';el.className='countdown green'}}tick();setInterval(tick,1000)})();</script></body></html>"""
    return render_template_string(html,o=order,stages=stages,pct=pct,current_stage=current_stage,
        ready=ready,deadline_ms=deadline_ms,deadline_text=deadline_text,total=total,paid=paid,
        remaining=remaining,money=_money,format_money=_currency_money,currency=currency,
        delivery=delivery,media=media,approvals=approvals,
        timeline=timeline,contact=contact,whatsapp=whatsapp,rang_css=rang_css,
        workflow_started=workflow_started,countdown_wait=countdown_wait,
        message=request.args.get("xabar",""))


@app.route("/buyurtma/<int:oid>/qr.png")
def order_qr(oid):
    c=get_db(); row=c.execute("SELECT tracking_token FROM buyurtmalar WHERE id=?",(oid,)).fetchone(); c.close()
    if not row:return "Topilmadi",404
    url=request.url_root.rstrip('/')+url_for('public_track',token=row['tracking_token'])
    img=qrcode.make(url); out=io.BytesIO(); img.save(out,format='PNG'); out.seek(0)
    return send_file(out,mimetype='image/png',download_name=f'buyurtma_{oid}_qr.png')

@app.route("/buyurtma/<int:oid>/shartnoma-yaratish", methods=["POST"])
def order_contract_regenerate(oid):
    try:
        result=generate_order_contract(oid)
        return jsonify({"status":"ok","versiya":result["version"]})
    except Exception as e:
        return jsonify({"message":str(e)}),400


@app.route("/buyurtma/<int:oid>/shartnoma.docx")
def order_contract_docx(oid):
    row=latest_order_contract(oid)
    if not row:
        try:
            generate_order_contract(oid)
            row=latest_order_contract(oid)
        except Exception as e:
            return "Shartnoma yaratishda xato: "+str(e),500
    if not row or not os.path.exists(row["docx_fayl"]):
        return "Shartnoma Word fayli topilmadi",404
    return send_file(row["docx_fayl"],as_attachment=True,
                     download_name=f"shartnoma_{oid}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/buyurtma/<int:oid>/shartnoma.pdf")
def order_contract_pdf(oid):
    row=latest_order_contract(oid)
    if not row:
        try:
            generate_order_contract(oid)
            row=latest_order_contract(oid)
        except Exception as e:
            return "Shartnoma yaratishda xato: "+str(e),500
    if not row or not os.path.exists(row["pdf_fayl"]):
        return "Shartnoma PDF fayli topilmadi",404
    return send_file(row["pdf_fayl"],mimetype="application/pdf",
                     as_attachment=False,download_name=f"shartnoma_{oid}.pdf")


@app.route("/buyurtma/<int:oid>/chek.pdf")
def order_receipt_pdf(oid):
    order,stages,pays=_order_bundle(oid)
    if not order:return "Topilmadi",404
    out=io.BytesIO(); p=canvas.Canvas(out,pagesize=A4); w,h=A4
    p.setFont('Helvetica-Bold',20); p.drawCentredString(w/2,h-60,'Mebel360° - TOLOV CHEKI')
    p.setFont('Helvetica',12); y=h-110
    currency=_currency_code(order['valyuta'] if 'valyuta' in order.keys() else 'UZS')
    paid=float(order['oldindan_tolov'] or 0); remaining=max(0,float(order['umumiy_narx'] or 0)-paid)
    lines=[f"Buyurtma: {order['kod']}",f"Mijoz: {order['mijoz']}",f"Mahsulot: {order['mahsulot']}",
           f"Kelishuv valyutasi: {currency}",f"Umumiy summa: {_currency_money(order['umumiy_narx'],currency)}",
           f"Jami to'langan: {_currency_money(paid,currency)}",f"Qoldiq: {_currency_money(remaining,currency)}",
           f"Kurs tartibi: {order['kurs_tartibi'] if 'kurs_tartibi' in order.keys() else 'Tolov kunidagi kurs'}",
           f"Chek sanasi: {_tashkent_today()}"]
    for line in lines:
        ascii_line=(str(line).replace("‘","'").replace("’","'").replace("–","-").replace("—","-"))
        p.drawString(70,y,ascii_line); y-=25
    if pays:
        p.setFont('Helvetica-Bold',11); p.drawString(70,y-5,"Oxirgi to'lovlar:"); y-=25
        p.setFont('Helvetica',9)
        for pay in pays[-6:]:
            pc=_currency_code(pay['tolov_valyutasi'] if 'tolov_valyutasi' in pay.keys() else currency)
            received=pay['qabul_qilingan_summa'] if 'qabul_qilingan_summa' in pay.keys() and pay['qabul_qilingan_summa'] else pay['miqdor']
            line=f"{pay['sana']}: {_currency_money(received,pc)}"
            if float(pay['kurs'] or 0)>1: line+=f" | 1 USD={_money(pay['kurs'])} UZS"
            p.drawString(70,y,line.replace("‘","'").replace("’","'")); y-=17
    p.drawString(70,y-20,'Rahmat! Mebel360° xizmatidan foydalanganingiz uchun.')
    p.save(); out.seek(0); return send_file(out,mimetype='application/pdf',as_attachment=True,download_name=f"chek_{order['kod']}.pdf")

@app.route("/api/buyurtma/<int:oid>/link")
def order_public_link(oid):
    c=get_db(); row=c.execute("SELECT tracking_token FROM buyurtmalar WHERE id=?",(oid,)).fetchone(); c.close()
    if not row:return jsonify({'message':'Topilmadi'}),404
    return jsonify({'url':request.url_root.rstrip('/')+url_for('public_track',token=row['tracking_token'])})

@app.route("/shartnoma-namuna")
def contract_template_download():
    try:
        base_dir=os.path.dirname(os.path.abspath(__file__))
        file_path=os.path.join(base_dir,"Mebel_Shartnoma_Tuzatilgan.docx")
        if not os.path.exists(file_path):
            return "Word shartnoma fayli topilmadi: "+file_path,404
        with open(file_path,"rb") as f:
            data=f.read()
        return Response(
            data,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition":"attachment; filename=Mebel_Shartnoma.docx"}
        )
    except Exception as e:
        return "Word shartnomani ochishda xato: "+str(e),500

@app.route("/shartnoma-pdf")
def contract_pdf_download():
    try:
        base_dir=os.path.dirname(os.path.abspath(__file__))
        file_path=os.path.join(base_dir,"Mebel_Shartnoma_Yakuniy.pdf")
        if not os.path.exists(file_path):
            return "PDF shartnoma fayli topilmadi: "+file_path,404
        with open(file_path,"rb") as f:
            data=f.read()
        return Response(
            data,
            mimetype="application/pdf",
            headers={"Content-Disposition":"inline; filename=Mebel_Shartnoma.pdf"}
        )
    except Exception as e:
        return "PDF shartnomani ochishda xato: "+str(e),500

@app.route("/backup")
def backup_db():
    if not os.path.exists(DB_NAME):
        return jsonify({"message":"Baza topilmadi"}),404
    log_action('backup_downloaded', f'user={session.get("user","")}')
    return send_file(DB_NAME, as_attachment=True, download_name=f"pharm_mebel_backup_{_tashkent_today()}.db")

@app.route("/api/audit")
def audit_get():
    c=get_db(); rows=c.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT 200").fetchall(); c.close(); return jsonify([dict(r) for r in rows])


# ---------- MIJOZ SERVISI / KAFOLAT / YETKAZISH ----------
@app.route("/api/buyurtma/<int:oid>/baho", methods=["GET","POST"])
def order_rating(oid):
    c=get_db()
    if request.method=="POST":
        d=jdata(); vals=[max(1,min(5,int(d.get(k) or 5))) for k in ("sifat","muddat","muomala","yetkazish","montaj")]
        c.execute("INSERT INTO mijoz_baholari(buyurtma_id,sifat,muddat,muomala,yetkazish,montaj,izoh,reklama_ruxsat) VALUES(?,?,?,?,?,?,?,?,?)",
                  (oid,*vals,d.get("izoh",""),1 if d.get("reklama_ruxsat") else 0)); c.commit(); c.close(); return jsonify({"status":"ok"})
    rows=c.execute("SELECT *,ROUND((sifat+muddat+muomala+yetkazish+montaj)/5.0,1) umumiy FROM mijoz_baholari WHERE buyurtma_id=? ORDER BY id DESC",(oid,)).fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/servis", methods=["GET","POST"])
def service_api():
    c=get_db()
    if request.method=="POST":
        d=jdata(); c.execute("INSERT INTO servis_murojaatlari(buyurtma_id,turi,muammo,holat,servis_sana,usta,izoh) VALUES(?,?,?,?,?,?,?)",
          (int(d["buyurtma_id"]),d.get("turi","Servis"),d.get("muammo",""),d.get("holat","Qabul qilindi"),d.get("servis_sana",""),d.get("usta",""),d.get("izoh",""))); c.commit(); c.close(); return jsonify({"status":"ok"})
    rows=c.execute("""SELECT s.*,b.kod,b.mijoz FROM servis_murojaatlari s JOIN buyurtmalar b ON b.id=s.buyurtma_id ORDER BY s.id DESC""").fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/servis/<int:sid>", methods=["POST"])
def service_update(sid):
    d=jdata(); c=get_db(); c.execute("UPDATE servis_murojaatlari SET holat=?,servis_sana=?,usta=?,izoh=? WHERE id=?",(d.get("holat","Ko‘rib chiqilmoqda"),d.get("servis_sana",""),d.get("usta",""),d.get("izoh",""),sid)); c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/qoshimcha-ish", methods=["GET","POST"])
def extras_api():
    c=get_db()
    if request.method=="POST":
        d=jdata(); c.execute("INSERT INTO qoshimcha_ishlar(buyurtma_id,nomi,summa,qoshimcha_kun,tasdiq,izoh) VALUES(?,?,?,?,?,?)",(int(d["buyurtma_id"]),d.get("nomi",""),float(d.get("summa") or 0),int(d.get("qoshimcha_kun") or 0),1 if d.get("tasdiq") else 0,d.get("izoh",""))); c.commit(); c.close(); return jsonify({"status":"ok"})
    rows=c.execute("SELECT q.*,b.kod,b.mijoz FROM qoshimcha_ishlar q JOIN buyurtmalar b ON b.id=q.buyurtma_id ORDER BY q.id DESC").fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/yetkazish", methods=["GET","POST"])
def delivery_api():
    c=get_db()
    if request.method=="POST":
        d=jdata(); c.execute("INSERT INTO yetkazishlar(buyurtma_id,haydovchi_id,sana,navbat,mashina,holat,benzin,yol_xarajati,izoh) VALUES(?,?,?,?,?,?,?,?,?)",(int(d["buyurtma_id"]),int(d["haydovchi_id"]) if d.get("haydovchi_id") else None,d.get("sana") or _tashkent_today(),int(d.get("navbat") or 1),d.get("mashina",""),d.get("holat","Rejalashtirilgan"),float(d.get("benzin") or 0),float(d.get("yol_xarajati") or 0),d.get("izoh",""))); c.commit(); c.close(); return jsonify({"status":"ok"})
    rows=c.execute("""SELECT y.*,b.kod,b.mijoz,b.telefon,b.manzil,b.lokatsiya,i.ism haydovchi_ism,i.familiya haydovchi_familiya FROM yetkazishlar y JOIN buyurtmalar b ON b.id=y.buyurtma_id LEFT JOIN ishchilar i ON i.id=y.haydovchi_id ORDER BY y.sana DESC,y.navbat""").fetchall(); c.close(); return jsonify([dict(r) for r in rows])

@app.route("/api/yetkazish/<int:yid>/holat", methods=["POST"])
def delivery_status(yid):
    d=jdata(); holat=d.get("holat","Rejalashtirilgan"); now=_tashkent_now().strftime('%Y-%m-%d %H:%M')
    field={"Yo‘lga chiqdim":"yolga_chiqdi","Yetib keldim":"yetib_keldi","Yetkazib berdim":"topshirildi"}.get(holat)
    c=get_db(); c.execute("UPDATE yetkazishlar SET holat=? WHERE id=?",(holat,yid))
    if field: c.execute(f"UPDATE yetkazishlar SET {field}=? WHERE id=?",(now,yid))
    c.commit(); c.close(); return jsonify({"status":"ok"})

@app.route("/api/mijoz-xulosa")
def customer_summary():
    c=get_db(); row=c.execute("""SELECT COUNT(*) buyurtmalar,COALESCE(SUM(umumiy_narx),0) jami_summa,COALESCE(SUM(oldindan_tolov),0) tushgan,
      COALESCE(AVG((sifat+muddat+muomala+yetkazish+montaj)/5.0),0) reyting FROM buyurtmalar LEFT JOIN mijoz_baholari ON mijoz_baholari.buyurtma_id=buyurtmalar.id""").fetchone(); c.close(); return jsonify(dict(row))

# ---------- JAMLANMA / REYTING ----------
@app.route("/api/jami")
def totals():
    start=request.args.get("start") or "1900-01-01"
    end=request.args.get("end") or "2999-12-31"
    c=get_db()
    rows=c.execute("""
    SELECT i.id,i.ism,i.familiya,i.lavozim,i.sifat_ball,i.tezlik_ball,i.intizom_ball,
    COALESCE(a.kun,0) ish_kunlari,COALESCE(a.soat,0) jami_soat,
    COALESCE(n.miqdor,0) jami_miqdor,COALESCE(n.haq,0) ish_haqi,
    COALESCE(j.jarima,0) jarima,COALESCE(t.tolov,0) tolangan,COALESCE(b.bonus,0) bonus,
    ROUND(COALESCE(n.haq,0)+COALESCE(b.bonus,0)-COALESCE(j.jarima,0)-COALESCE(t.tolov,0),2) qoldiq,
    ROUND((i.sifat_ball+i.tezlik_ball+i.intizom_ball)/3.0 +
          MIN(COALESCE(n.miqdor,0)/100.0,5),2) reyting
    FROM ishchilar i
    LEFT JOIN (SELECT ishchi_id,COUNT(DISTINCT sana) kun,ROUND(SUM(ish_soatlari),2) soat
               FROM keldi_ketdi WHERE sana BETWEEN ? AND ? GROUP BY ishchi_id) a ON a.ishchi_id=i.id
    LEFT JOIN (SELECT ishchi_id,ROUND(SUM(miqdor),2) miqdor,ROUND(SUM(jami_haq),2) haq
               FROM ish_natijalari WHERE sana BETWEEN ? AND ? GROUP BY ishchi_id) n ON n.ishchi_id=i.id
    LEFT JOIN (SELECT ishchi_id,ROUND(SUM(miqdor),2) jarima
               FROM jarimalar WHERE sana BETWEEN ? AND ? GROUP BY ishchi_id) j ON j.ishchi_id=i.id
    LEFT JOIN (SELECT ishchi_id,ROUND(SUM(miqdor),2) tolov
               FROM tolovlar WHERE sana BETWEEN ? AND ? GROUP BY ishchi_id) t ON t.ishchi_id=i.id
    LEFT JOIN (SELECT ishchi_id,ROUND(SUM(miqdor),2) bonus
               FROM bonuslar WHERE sana BETWEEN ? AND ? GROUP BY ishchi_id) b ON b.ishchi_id=i.id
    WHERE i.faol=1 ORDER BY reyting DESC,i.ism
    """,(start,end,start,end,start,end,start,end,start,end)).fetchall()
    c.close(); return jsonify([dict(r) for r in rows])


@app.route("/api/dashboard")
def dashboard():
    c=get_db(); today=_tashkent_today()
    data={
      "workers":c.execute("SELECT COUNT(*) FROM ishchilar WHERE faol=1").fetchone()[0],
      "orders":c.execute("SELECT COUNT(*) FROM buyurtmalar WHERE holat!='Yetkazildi'").fetchone()[0],
      "hours":c.execute("SELECT COALESCE(SUM(ish_soatlari),0) FROM keldi_ketdi WHERE sana=?",(today,)).fetchone()[0],
      "production":c.execute("SELECT COALESCE(SUM(miqdor),0) FROM ish_natijalari WHERE sana=?",(today,)).fetchone()[0],
      "km":c.execute("SELECT COALESCE(SUM(masofa_km),0) FROM safarlar WHERE sana=?",(today,)).fetchone()[0],
      "low_stock":c.execute("SELECT COUNT(*) FROM ombor WHERE qoldiq<=min_qoldiq").fetchone()[0]
    }
    c.close(); return jsonify(data)


@app.route("/export/jami.csv")
def export_csv():
    start=request.args.get("start") or "1900-01-01"; end=request.args.get("end") or "2999-12-31"
    c=get_db()
    rows=c.execute("""SELECT i.ism||' '||i.familiya ishchi,i.lavozim,
    COALESCE(SUM(n.miqdor),0) miqdor,COALESCE(SUM(n.jami_haq),0) ish_haqi
    FROM ishchilar i LEFT JOIN ish_natijalari n ON n.ishchi_id=i.id AND n.sana BETWEEN ? AND ?
    WHERE i.faol=1 GROUP BY i.id ORDER BY ishchi""",(start,end)).fetchall()
    c.close()
    out=io.StringIO(); w=csv.writer(out); w.writerow(["Ishchi","Lavozim","Miqdor","Ish haqi"])
    for r in rows:w.writerow(list(r))
    return Response(out.getvalue(),mimetype="text/csv",
                    headers={"Content-Disposition":"attachment; filename=jami_hisob.csv"})


# ---------- ISHCHI RO'YXATDAN O'TISH VA SHAXSIY KABINET ----------
def normalize_phone(value):
    digits=''.join(ch for ch in str(value or '') if ch.isdigit())
    if digits.startswith('998') and len(digits)==12:
        return '+'+digits
    if len(digits)==9:
        return '+998'+digits
    return '+'+digits if digits else ''


def make_otp():
    return f"{secrets.randbelow(1000000):06d}"


@app.route('/ishchi/royxat', methods=['GET','POST'])
def worker_register():
    msg=''; error=''; demo_code=''
    if request.method=='POST':
        phone=normalize_phone(request.form.get('telefon'))
        if len(phone)<10:
            error='Telefon raqamini to‘g‘ri kiriting.'
        else:
            c=get_db()
            existing=c.execute('SELECT id FROM ishchi_akkauntlari WHERE telefon=?',(phone,)).fetchone()
            if existing:
                error='Bu telefon raqami avval ro‘yxatdan o‘tgan.'
            else:
                code=make_otp()
                expires=datetime.now().timestamp()+600
                c.execute('DELETE FROM ishchi_otp WHERE telefon=?',(phone,))
                c.execute('INSERT INTO ishchi_otp(telefon,kod_hash,muddati) VALUES(?,?,?)',
                          (phone,generate_password_hash(code),str(expires)))
                c.commit(); c.close()
                session['worker_pending_phone']=phone
                # Haqiqiy SMS provayder ulanmaguncha kod admin panelida ko‘rinadi.
                if os.environ.get('DEV_SHOW_OTP','0')=='1':
                    demo_code=code
                msg='Tasdiqlash kodi yaratildi. Kodni kiriting.'
    return render_template_string(WORKER_REGISTER_HTML,msg=msg,error=error,demo_code=demo_code)


@app.route('/ishchi/kod', methods=['GET','POST'])
def worker_verify():
    phone=session.get('worker_pending_phone','')
    if not phone:
        return redirect(url_for('worker_register'))
    error=''
    if request.method=='POST':
        code=request.form.get('kod','').strip()
        login=request.form.get('login','').strip()
        password=request.form.get('password','')
        ism=request.form.get('ism','').strip()
        familiya=request.form.get('familiya','').strip()
        if len(login)<3 or not ism:
            error='Ism va kamida 3 belgili login kiriting.'
        elif _weak_password(password):
            error='Parol kamida 8 belgidan iborat bo‘lib, harf va raqamni birga o‘z ichiga olishi kerak.'
        else:
            c=get_db()
            row=c.execute('SELECT * FROM ishchi_otp WHERE telefon=? AND ishlatildi=0 ORDER BY id DESC LIMIT 1',(phone,)).fetchone()
            if not row or float(row['muddati']) < datetime.now().timestamp() or not check_password_hash(row['kod_hash'],code):
                error='Kod noto‘g‘ri yoki muddati tugagan.'
            elif c.execute('SELECT 1 FROM ishchi_akkauntlari WHERE login=?',(login,)).fetchone():
                error='Bu login band. Boshqa login tanlang.'
            else:
                # telefon bazadagi ishchiga mos tushsa bog‘laymiz, aks holda yangi ishchi yozuvi ochiladi
                worker=c.execute('SELECT id FROM ishchilar WHERE telefon=? AND faol=1 LIMIT 1',(phone,)).fetchone()
                if worker:
                    worker_id=worker['id']
                else:
                    cur=c.execute('INSERT INTO ishchilar(ism,familiya,telefon,lavozim) VALUES(?,?,?,?)',
                                  (ism,familiya,phone,'Ishchi'))
                    worker_id=cur.lastrowid
                c.execute('INSERT INTO ishchi_akkauntlari(ishchi_id,telefon,login,parol_hash,tasdiqlangan,admin_tasdiq) VALUES(?,?,?,?,1,0)',
                          (worker_id,phone,login,generate_password_hash(password)))
                c.execute('UPDATE ishchi_otp SET ishlatildi=1 WHERE id=?',(row['id'],))
                c.commit(); c.close()
                session.pop('worker_pending_phone',None)
                return render_template_string(WORKER_WAIT_HTML)
            c.close()
    return render_template_string(WORKER_VERIFY_HTML,telefon=phone,error=error)


@app.route('/ishchi/login', methods=['GET','POST'])
def worker_login():
    error=''
    if request.method=='POST':
        login=request.form.get('login','').strip()
        password=request.form.get('password','')
        locked, wait_sec = _is_login_locked('w:'+login)
        if locked:
            error=f'Juda ko‘p xato urinish. {wait_sec} soniyadan so‘ng qayta urinib ko‘ring.'
        else:
            c=get_db(); row=c.execute('SELECT * FROM ishchi_akkauntlari WHERE login=? AND faol=1',(login,)).fetchone(); c.close()
            if not row or not check_password_hash(row['parol_hash'],password):
                _register_failed_login('w:'+login)
                error='Login yoki parol xato.'
            elif not row['admin_tasdiq']:
                error='Admin hali akkauntingizni tasdiqlamagan.'
            else:
                _clear_login_attempts('w:'+login)
                session.clear(); session['worker_account_id']=row['id']; session['worker_id']=row['ishchi_id']
                log_action('worker_login', f'login={login}')
                return redirect(url_for('worker_dashboard'))
    return render_template_string(WORKER_LOGIN_HTML,error=error)


@app.route('/ishchi/logout')
def worker_logout():
    session.clear(); return redirect(url_for('worker_login'))


@app.route('/ishchi/kabinet')
def worker_dashboard():
    wid=session.get('worker_id')
    c=get_db()
    worker=c.execute('SELECT * FROM ishchilar WHERE id=?',(wid,)).fetchone()
    tasks=c.execute('SELECT * FROM ishchi_topshiriqlari WHERE ishchi_id=? ORDER BY CASE WHEN holat="Ishlayapti" THEN 0 WHEN holat="Yangi" THEN 1 ELSE 2 END,sana DESC,id DESC LIMIT 50',(wid,)).fetchall()
    active_task=c.execute('SELECT * FROM ishchi_topshiriqlari WHERE ishchi_id=? AND holat="Ishlayapti" ORDER BY id DESC LIMIT 1',(wid,)).fetchone()
    worker_state='Ishlayapti' if active_task else 'Bo‘sh'
    stats=c.execute('''SELECT COUNT(DISTINCT sana) kun,COALESCE(SUM(ish_soatlari),0) soat FROM keldi_ketdi WHERE ishchi_id=?''',(wid,)).fetchone()
    result=c.execute('SELECT COALESCE(SUM(miqdor),0) miqdor FROM ish_natijalari WHERE ishchi_id=?',(wid,)).fetchone()
    rating=round(((worker['sifat_ball']+worker['tezlik_ball']+worker['intizom_ball'])/3.0)+min(float(result['miqdor'] or 0)/100.0,5),2)
    c.close()
    return render_template_string(WORKER_DASHBOARD_HTML,worker=worker,tasks=tasks,stats=stats,rating=rating,worker_state=worker_state,active_task=active_task)


@app.route('/ishchi/topshiriq/<int:tid>/boshlash', methods=['POST'])
def worker_task_start(tid):
    wid=session.get('worker_id')
    now=_tashkent_now().strftime('%Y-%m-%d %H:%M:%S')
    c=get_db()
    other=c.execute('SELECT id FROM ishchi_topshiriqlari WHERE ishchi_id=? AND holat="Ishlayapti" AND id<>? LIMIT 1',(wid,tid)).fetchone()
    if other:
        c.close()
        flash('Avval ishlayotgan topshiriqni tugating.')
        return redirect(url_for('worker_dashboard'))
    c.execute('''UPDATE ishchi_topshiriqlari
                 SET holat='Ishlayapti',progress=CASE WHEN progress<1 THEN 1 ELSE progress END,
                     boshlandi_vaqt=CASE WHEN boshlandi_vaqt='' OR boshlandi_vaqt IS NULL THEN ? ELSE boshlandi_vaqt END
                 WHERE id=? AND ishchi_id=?''',(now,tid,wid))
    c.commit(); c.close()
    return redirect(url_for('worker_dashboard'))


@app.route('/ishchi/topshiriq/<int:tid>/tugatish', methods=['POST'])
def worker_task_finish(tid):
    wid=session.get('worker_id')
    now=_tashkent_now().strftime('%Y-%m-%d %H:%M:%S')
    c=get_db()
    c.execute('''UPDATE ishchi_topshiriqlari
                 SET holat='Tayyor',progress=100,tugadi_vaqt=?
                 WHERE id=? AND ishchi_id=?''',(now,tid,wid))
    c.commit(); c.close()
    return redirect(url_for('worker_dashboard'))


@app.route('/ishchi/topshiriq/<int:tid>', methods=['POST'])
def worker_task_update(tid):
    wid=session.get('worker_id')
    progress=max(0,min(100,int(request.form.get('progress') or 0)))
    status=request.form.get('holat','Jarayonda')
    c=get_db(); c.execute('UPDATE ishchi_topshiriqlari SET progress=?,holat=? WHERE id=? AND ishchi_id=?',(progress,status,tid,wid)); c.commit(); c.close()
    return redirect(url_for('worker_dashboard'))


@app.route('/ishchi-boshqaruv', methods=['GET','POST'])
def worker_admin():
    c=get_db()
    if request.method=='POST':
        action=request.form.get('action')
        if action=='approve':
            aid=int(request.form['account_id']); c.execute('UPDATE ishchi_akkauntlari SET admin_tasdiq=1 WHERE id=?',(aid,)); log_action('worker_approved', f'account_id={aid}')
        elif action=='block':
            aid=int(request.form['account_id']); c.execute('UPDATE ishchi_akkauntlari SET faol=0 WHERE id=?',(aid,)); log_action('worker_blocked', f'account_id={aid}')
        elif action=='task':
            c.execute('''INSERT INTO ishchi_topshiriqlari(ishchi_id,buyurtma_kodi,ish_turi,tavsif,holat,progress,sana,tugash_sana)
                         VALUES(?,?,?,?,?,?,?,?)''',(
                int(request.form['ishchi_id']),request.form.get('buyurtma_kodi','').strip(),request.form.get('ish_turi','').strip(),
                request.form.get('tavsif','').strip(),'Yangi',0,request.form.get('sana') or _tashkent_today(),request.form.get('tugash_sana','')
            ))
        c.commit()
    accounts=c.execute('''SELECT a.*,i.ism,i.familiya,i.lavozim FROM ishchi_akkauntlari a LEFT JOIN ishchilar i ON i.id=a.ishchi_id ORDER BY a.id DESC''').fetchall()
    workers=c.execute('SELECT id,ism,familiya,lavozim FROM ishchilar WHERE faol=1 ORDER BY ism').fetchall()
    worker_states=c.execute('''SELECT i.id,i.ism,i.familiya,i.lavozim,
        CASE WHEN t.id IS NULL THEN 'Bo‘sh' ELSE 'Ishlayapti' END ish_holati,
        COALESCE(t.buyurtma_kodi,'') buyurtma_kodi,
        COALESCE(t.ish_turi,'') ish_turi,
        COALESCE(t.boshlandi_vaqt,'') boshlandi_vaqt
        FROM ishchilar i
        LEFT JOIN ishchi_topshiriqlari t ON t.id=(
            SELECT t2.id FROM ishchi_topshiriqlari t2
            WHERE t2.ishchi_id=i.id AND t2.holat='Ishlayapti'
            ORDER BY t2.id DESC LIMIT 1
        )
        WHERE i.faol=1 ORDER BY ish_holati DESC,i.ism''').fetchall()
    tasks=c.execute('''SELECT t.*,i.ism,i.familiya FROM ishchi_topshiriqlari t JOIN ishchilar i ON i.id=t.ishchi_id ORDER BY t.id DESC LIMIT 100''').fetchall()
    otp_rows=c.execute('SELECT telefon,created_at FROM ishchi_otp WHERE ishlatildi=0 ORDER BY id DESC LIMIT 20').fetchall()
    c.close()
    return render_template_string(WORKER_ADMIN_HTML,accounts=accounts,workers=workers,worker_states=worker_states,tasks=tasks,otp_rows=otp_rows,today=_tashkent_today())



# ---------- SHOFYOR KABINETI ----------
@app.route('/shofyor/royxat', methods=['GET','POST'])
def driver_register():
    error=''
    msg=''
    if request.method=='POST':
        ism=request.form.get('ism','').strip()
        familiya=request.form.get('familiya','').strip()
        telefon=normalize_phone(request.form.get('telefon',''))
        login=request.form.get('login','').strip()
        password=request.form.get('password','')
        if not ism or len(login)<3 or len(telefon)<10:
            error='Ism, telefon va kamida 3 belgili login kiriting.'
        elif _weak_password(password):
            error='Parol kamida 8 belgidan iborat bo‘lib, harf va raqamni birga o‘z ichiga olishi kerak.'
        else:
            c=get_db()
            if c.execute("SELECT 1 FROM shofyor_akkauntlari WHERE login=?",(login,)).fetchone():
                error='Bu login band.'
            elif c.execute("SELECT 1 FROM shofyor_akkauntlari WHERE telefon=?",(telefon,)).fetchone():
                error='Bu telefon avval ro‘yxatdan o‘tgan.'
            else:
                worker=c.execute("SELECT id FROM ishchilar WHERE telefon=? AND faol=1",(telefon,)).fetchone()
                if worker:
                    wid=worker['id']
                    c.execute("UPDATE ishchilar SET lavozim='Shofyor' WHERE id=?",(wid,))
                else:
                    cur=c.execute("INSERT INTO ishchilar(ism,familiya,telefon,lavozim) VALUES(?,?,?,'Shofyor')",
                                  (ism,familiya,telefon))
                    wid=cur.lastrowid
                c.execute("INSERT INTO shofyor_akkauntlari(ishchi_id,login,parol_hash,telefon,faol,admin_tasdiq) VALUES(?,?,?,?,1,0)",
                          (wid,login,generate_password_hash(password),telefon))
                c.commit()
                msg='Ro‘yxatdan o‘tdingiz. Endi rahbar tasdiqlaydi.'
            c.close()
    return render_template_string(DRIVER_REGISTER_HTML,error=error,msg=msg)

@app.route('/shofyor/login', methods=['GET','POST'])
def driver_login():
    error=''
    if request.method=='POST':
        login=request.form.get('login','').strip()
        password=request.form.get('password','')
        locked, wait_sec = _is_login_locked('d:'+login)
        if locked:
            error=f'Juda ko‘p xato urinish. {wait_sec} soniyadan so‘ng qayta urinib ko‘ring.'
        else:
            c=get_db()
            row=c.execute("SELECT a.*,i.ism,i.familiya FROM shofyor_akkauntlari a JOIN ishchilar i ON i.id=a.ishchi_id WHERE a.login=? AND a.faol=1",(login,)).fetchone()
            c.close()
            if not row or not check_password_hash(row['parol_hash'],password):
                _register_failed_login('d:'+login)
                error='Login yoki parol xato.'
            elif not row['admin_tasdiq']:
                error='Rahbar hali akkauntingizni tasdiqlamagan.'
            else:
                _clear_login_attempts('d:'+login)
                session.clear()
                session['driver_account_id']=row['id']
                session['driver_id']=row['ishchi_id']
                log_action('driver_login', f'login={login}')
                return redirect(url_for('driver_dashboard'))
    return render_template_string(DRIVER_LOGIN_HTML,error=error)

@app.route('/shofyor/logout')
def driver_logout():
    session.clear()
    return redirect(url_for('driver_login'))

@app.route('/shofyor/kabinet')
def driver_dashboard():
    did=session.get('driver_id')
    c=get_db()
    driver=c.execute("SELECT * FROM ishchilar WHERE id=?",(did,)).fetchone()
    deliveries=c.execute("SELECT y.*,b.kod,b.mijoz,b.telefon,b.manzil,b.mahsulot,b.lokatsiya,b.moljal,b.qavat,b.lift,b.katta_mashina,b.izoh buyurtma_izoh FROM yetkazishlar y JOIN buyurtmalar b ON b.id=y.buyurtma_id WHERE y.haydovchi_id=? ORDER BY CASE WHEN y.holat='Yetkazib berildi' THEN 1 ELSE 0 END,y.navbat,y.id DESC",(did,)).fetchall()
    c.close()
    return render_template_string(DRIVER_DASHBOARD_HTML,driver=driver,deliveries=deliveries)

@app.route('/shofyor/yetkazish/<int:yid>')
def driver_delivery_detail(yid):
    did=session.get('driver_id')
    c=get_db()
    row=c.execute("SELECT y.*,b.kod,b.mijoz,b.telefon,b.manzil,b.mahsulot,b.lokatsiya,b.moljal,b.qavat,b.lift,b.katta_mashina,b.izoh buyurtma_izoh FROM yetkazishlar y JOIN buyurtmalar b ON b.id=y.buyurtma_id WHERE y.id=? AND y.haydovchi_id=?",(yid,did)).fetchone()
    c.close()
    if not row:
        return 'Yetkazish topilmadi',404
    return render_template_string(DRIVER_DETAIL_HTML,x=row)

@app.route('/shofyor/yetkazish/<int:yid>/holat', methods=['POST'])
def driver_delivery_status(yid):
    did=session.get('driver_id')
    action=request.form.get('action','')
    now=_tashkent_now().strftime('%Y-%m-%d %H:%M')
    mapping={'yolga':('Yo‘lga chiqdi','yolga_chiqdi'),'yetib':('Yetib keldi','yetib_keldi'),'yetkazildi':('Yetkazib berildi','yetkazildi')}
    if action not in mapping:
        return redirect(url_for('driver_delivery_detail',yid=yid))
    status,col=mapping[action]
    c=get_db()
    c.execute(f"UPDATE yetkazishlar SET holat=?, {col}=? WHERE id=? AND haydovchi_id=?",(status,now,yid,did))
    c.commit(); c.close()
    return redirect(url_for('driver_delivery_detail',yid=yid))

@app.route('/shofyor-boshqaruv', methods=['GET','POST'])
def driver_admin():
    c=get_db()
    msg=''
    if request.method=='POST':
        action=request.form.get('action')
        try:
            if action=='account':
                wid=int(request.form['ishchi_id'])
                login=request.form.get('login','').strip()
                password=request.form.get('password','')
                if _weak_password(password):
                    msg='Xato: parol kamida 8 belgidan iborat bo‘lib, harf va raqamni birga o‘z ichiga olishi kerak.'
                else:
                    existing=c.execute("SELECT id FROM shofyor_akkauntlari WHERE ishchi_id=?",(wid,)).fetchone()
                    if existing:
                        c.execute("UPDATE shofyor_akkauntlari SET login=?,parol_hash=?,faol=1,admin_tasdiq=1 WHERE ishchi_id=?",(login,generate_password_hash(password),wid))
                    else:
                        c.execute("INSERT INTO shofyor_akkauntlari(ishchi_id,login,parol_hash,faol,admin_tasdiq) VALUES(?,?,?,1,1)",(wid,login,generate_password_hash(password)))
                    msg='Shofyor akkaunti saqlandi.'
                    log_action('driver_account_set', f'ishchi_id={wid},login={login}')
            elif action=='approve_driver':
                aid=int(request.form['account_id'])
                c.execute("UPDATE shofyor_akkauntlari SET admin_tasdiq=1,faol=1 WHERE id=?",(aid,))
                log_action('driver_approved', f'account_id={aid}')
                msg='Shofyor tasdiqlandi.'
            elif action=='block_driver':
                aid=int(request.form['account_id'])
                c.execute("UPDATE shofyor_akkauntlari SET faol=0 WHERE id=?",(aid,))
                log_action('driver_blocked', f'account_id={aid}')
                msg='Shofyor bloklandi.'
            elif action=='delivery':
                c.execute("INSERT INTO yetkazishlar(buyurtma_id,haydovchi_id,sana,qadoq_soni,navbat,izoh) VALUES(?,?,?,?,?,?)",(int(request.form['buyurtma_id']),int(request.form['shofyor_id']),_tashkent_today(),int(request.form.get('qadoq_soni') or 1),int(request.form.get('navbat') or 1),request.form.get('izoh','')))
                log_action('delivery_assigned', f"buyurtma_id={request.form['buyurtma_id']},shofyor_id={request.form['shofyor_id']}")
                msg='Yetkazish shofyorga biriktirildi.'
            c.commit()
        except Exception as e:
            c.rollback()
            msg='Xato: '+str(e)
    drivers=c.execute("SELECT id,ism,familiya,lavozim FROM ishchilar WHERE faol=1 AND (lower(lavozim) LIKE '%shof%' OR lower(lavozim) LIKE '%haydov%') ORDER BY ism").fetchall()
    orders=c.execute("SELECT id,kod,mijoz,mahsulot FROM buyurtmalar WHERE holat!='Yetkazildi' ORDER BY id DESC").fetchall()
    assigned=c.execute("SELECT y.*,b.kod,b.mahsulot,i.ism,i.familiya FROM yetkazishlar y JOIN buyurtmalar b ON b.id=y.buyurtma_id JOIN ishchilar i ON i.id=y.haydovchi_id ORDER BY y.id DESC LIMIT 100").fetchall()
    driver_accounts=c.execute("SELECT a.*,i.ism,i.familiya FROM shofyor_akkauntlari a JOIN ishchilar i ON i.id=a.ishchi_id ORDER BY a.id DESC").fetchall()
    c.close()
    return render_template_string(DRIVER_ADMIN_HTML,drivers=drivers,orders=orders,assigned=assigned,driver_accounts=driver_accounts,msg=msg)





@app.route("/pro-boshqaruv")
def pro_admin_page():
    c=get_db()
    stats={
        "orders":c.execute("SELECT COUNT(*) FROM buyurtmalar").fetchone()[0],
        "services":c.execute("SELECT COUNT(*) FROM servis_murojaatlari WHERE holat!='Yopildi'").fetchone()[0],
        "ratings":c.execute("SELECT COUNT(*) FROM baholar").fetchone()[0],
        "messages":c.execute("SELECT COUNT(*) FROM tizim_xabarlari WHERE oqildi=0").fetchone()[0],
        "reservations":c.execute("SELECT COUNT(*) FROM ombor_rezervlari WHERE holat='Rezerv'").fetchone()[0]
    }
    orders=c.execute("SELECT id,kod,mijoz,mahsulot,holat,umumiy_narx,oldindan_tolov FROM buyurtmalar ORDER BY id DESC LIMIT 30").fetchall()
    c.close()
    return render_template_string(PRO_ADMIN_HTML,stats=stats,orders=orders)


# ---------- MENEJER VA KONSTRUKTOR KABINETLARI ----------
def _staff_account(login, role):
    c = get_db()
    row = c.execute(
        "SELECT * FROM xodim_akkauntlari WHERE login=? AND rol=? LIMIT 1",
        ((login or "").strip(), role),
    ).fetchone()
    c.close()
    return row


def _staff_login(role, role_label, home_endpoint):
    error = ""
    if request.method == "POST":
        login_value = (request.form.get("login") or "").strip()
        password = request.form.get("password") or ""
        lock_key = f"staff:{role}:{login_value}"
        locked, wait_sec = _is_login_locked(lock_key)
        account = _staff_account(login_value, role)
        if locked:
            minutes = max(1, (wait_sec + 59) // 60)
            error = f"Juda ko‘p xato urinish. Taxminan {minutes} daqiqadan so‘ng qayta urinib ko‘ring."
        elif account and int(account["faol"] or 0) == 1 and check_password_hash(account["parol_hash"], password):
            _clear_login_attempts(lock_key)
            session.clear()
            session["staff_account_id"] = account["id"]
            session["staff_role"] = role
            session["staff_name"] = account["ism"]
            session["staff_login"] = account["login"]
            log_action(f"{role}_login", f"login={account['login']}")
            return redirect(url_for(home_endpoint))
        else:
            _register_failed_login(lock_key)
            error = "Login yoki parol xato."
    return render_template_string(
        STAFF_LOGIN_HTML,
        error=error,
        role_label=role_label,
        role=role,
        icon="🧑‍💼" if role == "menejer" else "📐",
    )


@app.route('/menejer/login', methods=['GET', 'POST'])
def manager_login():
    return _staff_login('menejer', 'Menejer', 'manager_dashboard')


@app.route('/menejer/logout')
def manager_logout():
    log_action('menejer_logout', f"login={session.get('staff_login', '')}")
    session.clear()
    return redirect(url_for('manager_login'))


@app.route('/konstruktor/login', methods=['GET', 'POST'])
def constructor_login():
    return _staff_login('konstruktor', 'Konstruktor', 'constructor_dashboard')


@app.route('/konstruktor/logout')
def constructor_logout():
    log_action('konstruktor_logout', f"login={session.get('staff_login', '')}")
    session.clear()
    return redirect(url_for('constructor_login'))


@app.route('/xodim-akkauntlari', methods=['GET', 'POST'])
def staff_accounts_admin():
    msg = ""
    c = get_db()
    if request.method == 'POST':
        action = request.form.get('action', 'save')
        try:
            if action == 'save':
                ism = (request.form.get('ism') or '').strip()
                login_value = (request.form.get('login') or '').strip()
                role = request.form.get('rol') or ''
                password = request.form.get('password') or ''
                if role not in {'menejer', 'konstruktor'}:
                    raise ValueError('Rol noto‘g‘ri tanlandi.')
                if len(ism) < 2 or len(login_value) < 3:
                    raise ValueError('Ism va loginni to‘liq kiriting.')
                if _weak_password(password):
                    raise ValueError('Parol kamida 8 belgi bo‘lib, harf va raqamni birga o‘z ichiga olsin.')
                existing = c.execute('SELECT id FROM xodim_akkauntlari WHERE login=?', (login_value,)).fetchone()
                if existing:
                    c.execute('''UPDATE xodim_akkauntlari
                                 SET ism=?,parol_hash=?,rol=?,faol=1,updated_at=? WHERE id=?''',
                              (ism, generate_password_hash(password), role,
                               _tashkent_now().strftime('%Y-%m-%d %H:%M:%S'), existing['id']))
                else:
                    c.execute('''INSERT INTO xodim_akkauntlari(ism,login,parol_hash,rol,faol)
                                 VALUES(?,?,?,?,1)''',
                              (ism, login_value, generate_password_hash(password), role))
                msg = 'Akkaunt saqlandi.'
                log_action('staff_account_saved', f'login={login_value},role={role}')
            elif action == 'toggle':
                aid = int(request.form['account_id'])
                c.execute('UPDATE xodim_akkauntlari SET faol=CASE WHEN faol=1 THEN 0 ELSE 1 END,updated_at=? WHERE id=?',
                          (_tashkent_now().strftime('%Y-%m-%d %H:%M:%S'), aid))
                msg = 'Akkaunt holati o‘zgartirildi.'
                log_action('staff_account_toggled', f'id={aid}')
            c.commit()
        except Exception as e:
            c.rollback()
            msg = 'Xato: ' + str(e)
    accounts = c.execute('SELECT * FROM xodim_akkauntlari ORDER BY rol,ism').fetchall()
    c.close()
    return render_template_string(STAFF_ADMIN_HTML, accounts=accounts, msg=msg)


@app.route('/menejer/kabinet', methods=['GET', 'POST'])
def manager_dashboard():
    c = get_db()
    if request.method == 'POST':
        action = request.form.get('action', '')
        try:
            if action == 'add_order':
                code = (request.form.get('kod') or '').strip()
                customer = (request.form.get('mijoz') or '').strip()
                if not code or not customer:
                    raise ValueError('Buyurtma kodi va mijoz ismi majburiy.')
                order_type=request.form.get('buyurtma_turi') or 'To‘liq mebel'
                payment_condition=request.form.get('tolov_sharti') or 'Avans majburiy'
                order_currency=_currency_code(request.form.get('valyuta') or 'UZS')
                initial_received=float(request.form.get('oldindan_tolov') or 0)
                initial_currency=_currency_code(request.form.get('boshlangich_tolov_valyutasi') or order_currency)
                initial_rate=float(request.form.get('boshlangich_kurs') or 0)
                initial_payment=0.0
                initial_uzs=0.0
                if initial_received>0:
                    initial_payment,initial_uzs,initial_rate=_payment_conversion(order_currency,initial_currency,initial_received,initial_rate)
                cur = c.execute('''INSERT INTO buyurtmalar(
                    kod,mijoz,telefon,manzil,mahsulot,umumiy_narx,oldindan_tolov,
                    boshlanish_sana,tugash_sana,taxminiy_sana,taxminiy_vaqt,holat,izoh,
                    tracking_token,masul_xodim,material,rang,buyurtma_turi,tolov_sharti,
                    avans_talab,avans_muddat_sana,taklif_amal_sana,ishlab_chiqarish_kun,
                    muddat_tartibi,rasmiy_muddat_vaqt,valyuta,kurs_tartibi,oxirgi_kurs)
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                    (code, customer, request.form.get('telefon',''), request.form.get('manzil',''),
                     request.form.get('mahsulot',''), float(request.form.get('umumiy_narx') or 0),
                     initial_payment, request.form.get('boshlanish_sana',''),
                     request.form.get('tugash_sana',''), request.form.get('taxminiy_sana',''),
                     _safe_time(request.form.get('taxminiy_vaqt')), request.form.get('holat','Yangi'),
                     request.form.get('izoh',''), secrets.token_urlsafe(8), session.get('staff_name','Menejer'),
                     request.form.get('material',''), request.form.get('rang',''), order_type,payment_condition,
                     float(request.form.get('avans_talab') or 0),request.form.get('avans_muddat_sana',''),
                     request.form.get('taklif_amal_sana',''),int(request.form.get('ishlab_chiqarish_kun') or 0),
                     'Tasdiqdan keyin',_safe_time(request.form.get('taxminiy_vaqt')),order_currency,'To‘lov kunidagi kurs',initial_rate if initial_rate>1 else 0))
                order_id = cur.lastrowid
                if initial_received>0:
                    c.execute('''INSERT INTO buyurtma_tolovlari
                        (buyurtma_id,sana,miqdor,turi,izoh,tolov_valyutasi,kurs,qabul_qilingan_summa,buyurtma_summa,uzs_ekvivalent)
                        VALUES(?,?,?,?,?,?,?,?,?,?)''',(order_id,_tashkent_today(),initial_payment,'Avans','Menejer buyurtma yaratganda kiritdi',initial_currency,initial_rate,initial_received,initial_payment,initial_uzs))
                workflow=_workflow_stages(order_type,payment_condition,request.form.get('maxsus_bosqichlar',''))
                c.executemany('INSERT INTO buyurtma_bosqichlari(buyurtma_id,bosqich) VALUES(?,?)',
                              [(order_id, stage) for stage in workflow])
                if payment_condition=='Avans talab qilinmaydi — ishonchli mijoz':
                    _complete_order_stage(c,order_id,'Avans talab qilinmaydi','Ishonchli mijoz uchun avans talab qilinmaydi')
                    c.execute("UPDATE buyurtmalar SET holat='Rahbar tasdiqlashi kutilmoqda' WHERE id=?",(order_id,))
                new_order=c.execute('SELECT * FROM buyurtmalar WHERE id=?',(order_id,)).fetchone()
                if _payment_ready(new_order):
                    pay_stage='Qisman avans olindi' if payment_condition=='Qisman avans' else 'Avans olindi'
                    _complete_order_stage(c,order_id,pay_stage,'Menejer buyurtma yaratganda to‘lov kiritdi')
                    _start_official_deadline(c,order_id,'Avans qabul qilindi — rasmiy muddat boshlandi')
                c.commit()
                try:
                    generate_order_contract(order_id)
                except Exception:
                    pass
                log_action('manager_order_added', f'order_id={order_id},code={code}')
                flash('✅ Yangi buyurtma saqlandi.')
            elif action == 'payment':
                order_id = int(request.form['buyurtma_id'])
                order_row=c.execute('SELECT * FROM buyurtmalar WHERE id=?',(order_id,)).fetchone()
                received=float(request.form.get('miqdor') or 0)
                payment_currency=_currency_code(request.form.get('tolov_valyutasi') or order_row['valyuta'] or 'UZS')
                amount,uzs_equivalent,rate=_payment_conversion(order_row['valyuta'],payment_currency,received,float(request.form.get('kurs') or 0))
                c.execute('''INSERT INTO buyurtma_tolovlari
                    (buyurtma_id,sana,miqdor,turi,izoh,tolov_valyutasi,kurs,qabul_qilingan_summa,buyurtma_summa,uzs_ekvivalent)
                    VALUES(?,?,?,?,?,?,?,?,?,?)''',(order_id,request.form.get('sana') or _tashkent_today(),amount,request.form.get('turi','To‘lov'),request.form.get('izoh',''),payment_currency,rate,received,amount,uzs_equivalent))
                c.execute('UPDATE buyurtmalar SET oldindan_tolov=oldindan_tolov+?,oxirgi_kurs=CASE WHEN ?>1 THEN ? ELSE oxirgi_kurs END WHERE id=?', (amount,rate,rate,order_id))
                updated=c.execute('SELECT * FROM buyurtmalar WHERE id=?',(order_id,)).fetchone()
                if _payment_ready(updated):
                    condition=str(updated['tolov_sharti'] or 'Avans majburiy')
                    pay_stage='Qisman avans olindi' if condition=='Qisman avans' else 'Avans olindi'
                    _complete_order_stage(c,order_id,pay_stage,'Menejer to‘lov kiritdi')
                    _start_official_deadline(c,order_id,'Avans qabul qilindi — rasmiy muddat boshlandi')
                c.commit()
                log_action('manager_payment_added', f'order_id={order_id},amount={amount}')
                flash('✅ To‘lov qo‘shildi.')
            elif action == 'status':
                order_id = int(request.form['buyurtma_id'])
                status = request.form.get('holat') or 'Yangi'
                allowed = {'Yangi','Jarayonda','Tayyor','Yetkazishga tayyor','Yetkazildi','Yopildi'}
                if status not in allowed:
                    raise ValueError('Holat noto‘g‘ri.')
                c.execute('UPDATE buyurtmalar SET holat=? WHERE id=?', (status, order_id))
                c.commit()
                log_action('manager_order_status', f'order_id={order_id},status={status}')
                flash('✅ Buyurtma holati yangilandi.')
            else:
                raise ValueError('Noma’lum amal.')
        except Exception as e:
            c.rollback()
            flash('❌ ' + str(e))
        c.close()
        return redirect(url_for('manager_dashboard'))

    orders = c.execute('''SELECT b.*,
        ROUND(MAX(0,b.umumiy_narx-b.oldindan_tolov),2) qoldiq,
        COALESCE(ROUND(100.0*SUM(CASE WHEN bs.bajarildi=1 THEN 1 ELSE 0 END)/NULLIF(COUNT(bs.id),0),1),0) progress
        FROM buyurtmalar b LEFT JOIN buyurtma_bosqichlari bs ON bs.buyurtma_id=b.id
        GROUP BY b.id ORDER BY b.id DESC LIMIT 200''').fetchall()
    stats = c.execute('''SELECT COUNT(*) jami,
        SUM(CASE WHEN holat NOT IN ('Yetkazildi','Yopildi') THEN 1 ELSE 0 END) faol,
        COALESCE(SUM(CASE WHEN valyuta='USD' AND oxirgi_kurs>0 THEN umumiy_narx*oxirgi_kurs WHEN valyuta='UZS' THEN umumiy_narx ELSE 0 END),0) summa,
        COALESCE((SELECT SUM(uzs_ekvivalent) FROM buyurtma_tolovlari),0) tushum FROM buyurtmalar''').fetchone()
    c.close()
    return render_template_string(MANAGER_DASHBOARD_HTML, orders=orders, stats=stats,
                                  today=_tashkent_now().date().isoformat(), staff_name=session.get('staff_name','Menejer'))


CONSTRUCTOR_ALLOWED_EXTENSIONS = {
    'sto','dxf','dwg','skp','art','nc','cnc','tap','mpr','pdf','txt','csv','xlsx','zip','rar','7z','jpg','jpeg','png','webp'
}


@app.route('/konstruktor/kabinet', methods=['GET', 'POST'])
def constructor_dashboard():
    c = get_db()
    selected_id = request.args.get('order_id', type=int)
    if request.method == 'POST':
        action = request.form.get('action', '')
        selected_id = request.form.get('order_id', type=int)
        try:
            if action == 'stage':
                stage_id = int(request.form['stage_id'])
                done = 1 if request.form.get('bajarildi') == '1' else 0
                now_text = _tashkent_now().strftime('%Y-%m-%d %H:%M:%S')
                c.execute('''UPDATE buyurtma_bosqichlari SET bajarildi=?,
                    boshlanish_vaqti=CASE WHEN ?=1 AND COALESCE(boshlanish_vaqti,'')='' THEN ? ELSE boshlanish_vaqti END,
                    tugash_vaqti=CASE WHEN ?=1 THEN ? ELSE '' END,
                    ishchi=?,izoh=? WHERE id=?''',
                    (done, done, now_text, done, now_text, session.get('staff_name','Konstruktor'),
                     request.form.get('izoh',''), stage_id))
                flash('✅ Bosqich yangilandi.')
                log_action('constructor_stage_updated', f'stage_id={stage_id},done={done}')
            elif action == 'upload':
                if not selected_id:
                    raise ValueError('Buyurtmani tanlang.')
                uploaded = request.files.get('fayl')
                if not uploaded or not uploaded.filename:
                    raise ValueError('Fayl tanlanmagan.')
                original = secure_filename(uploaded.filename)
                if not original or '.' not in original:
                    raise ValueError('Fayl nomi yoki turi noto‘g‘ri.')
                ext = original.rsplit('.', 1)[1].lower()
                if ext not in CONSTRUCTOR_ALLOWED_EXTENSIONS:
                    raise ValueError('Bu turdagi faylga ruxsat berilmagan.')
                stored = f"{selected_id}_{_tashkent_now().strftime('%Y%m%d_%H%M%S')}_{secrets.token_hex(3)}_{original}"
                uploaded.save(os.path.join(CONSTRUCTOR_UPLOAD_DIR, stored))
                c.execute('INSERT INTO buyurtma_hujjatlari(buyurtma_id,nomi,fayl_nomi) VALUES(?,?,?)',
                          (selected_id, request.form.get('nomi') or original, stored))
                flash('✅ Konstruktor fayli yuklandi.')
                log_action('constructor_file_uploaded', f'order_id={selected_id},file={stored}')
            elif action == 'order_status':
                status = request.form.get('holat') or 'Jarayonda'
                c.execute('UPDATE buyurtmalar SET holat=? WHERE id=?', (status, selected_id))
                flash('✅ Buyurtma holati yangilandi.')
            else:
                raise ValueError('Noma’lum amal.')
            c.commit()
        except Exception as e:
            c.rollback()
            flash('❌ ' + str(e))
        c.close()
        return redirect(url_for('constructor_dashboard', order_id=selected_id or ''))

    orders = c.execute('''SELECT b.*,
        COALESCE(ROUND(100.0*SUM(CASE WHEN bs.bajarildi=1 THEN 1 ELSE 0 END)/NULLIF(COUNT(bs.id),0),1),0) progress
        FROM buyurtmalar b LEFT JOIN buyurtma_bosqichlari bs ON bs.buyurtma_id=b.id
        GROUP BY b.id ORDER BY b.id DESC LIMIT 200''').fetchall()
    if not selected_id and orders:
        selected_id = orders[0]['id']
    selected_order = None
    stages = []
    files = []
    if selected_id:
        selected_order = c.execute('SELECT * FROM buyurtmalar WHERE id=?', (selected_id,)).fetchone()
        stages = c.execute('SELECT * FROM buyurtma_bosqichlari WHERE buyurtma_id=? ORDER BY id', (selected_id,)).fetchall()
        files = c.execute('SELECT * FROM buyurtma_hujjatlari WHERE buyurtma_id=? ORDER BY id DESC', (selected_id,)).fetchall()
    c.close()
    return render_template_string(CONSTRUCTOR_DASHBOARD_HTML, orders=orders, selected_order=selected_order,
                                  stages=stages, files=files, staff_name=session.get('staff_name','Konstruktor'))


@app.route('/konstruktor/fayl/<path:filename>')
def constructor_file_download(filename):
    return send_from_directory(CONSTRUCTOR_UPLOAD_DIR, filename, as_attachment=True)


STAFF_LOGIN_HTML = r"""
<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Mebel360° — {{role_label}} kirishi</title>
<style>
*{box-sizing:border-box}body{margin:0;min-height:100vh;display:grid;place-items:center;padding:18px;font-family:Arial;background:linear-gradient(145deg,#07142f,#123b8f,#0784c8)}
.box{width:min(430px,100%);background:#fff;border-radius:24px;padding:30px;box-shadow:0 28px 75px #02081766}.icon{width:66px;height:66px;border-radius:19px;display:grid;place-items:center;font-size:31px;background:#dbeafe;margin-bottom:18px}h1{margin:0 0 7px;color:#0f172a}.muted{color:#64748b;font-size:14px;line-height:1.5}label{display:block;font-weight:800;font-size:13px;margin:15px 0 6px}input{width:100%;height:50px;border:1px solid #cbd5e1;border-radius:12px;padding:0 14px;font-size:15px}button{width:100%;height:50px;border:0;border-radius:12px;background:#2563eb;color:#fff;font-weight:900;margin-top:19px;cursor:pointer}.err{margin-top:12px;color:#b91c1c;background:#fef2f2;border:1px solid #fecaca;padding:10px;border-radius:10px}.links{display:flex;justify-content:space-between;gap:10px;margin-top:18px;font-size:13px}.links a{color:#2563eb;text-decoration:none;font-weight:700}
</style></head><body><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><div class="icon">{{icon}}</div><h1>{{role_label}} kabineti</h1><p class="muted">Faqat sizga biriktirilgan vazifalar va ruxsat berilgan bo‘limlar ochiladi.</p><label>Login</label><input name="login" required autofocus autocomplete="username"><label>Parol</label><input name="password" type="password" required autocomplete="current-password"><button>Kirish →</button>{% if error %}<div class="err">⚠ {{error}}</div>{% endif %}<div class="links"><a href="/login">← Rahbar kirishi</a>{% if role=='menejer' %}<a href="/konstruktor/login">Konstruktor</a>{% else %}<a href="/menejer/login">Menejer</a>{% endif %}</div></form></body></html>
"""

STAFF_ADMIN_HTML = r"""
<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Menejer va Konstruktor akkauntlari</title>
<style>*{box-sizing:border-box}body{margin:0;font-family:Arial;background:#eef3f8;color:#172033}.head{background:linear-gradient(135deg,#0f1b33,#2563eb);color:#fff;padding:18px}.wrap{max-width:1100px;margin:auto;padding:16px}.box{background:#fff;border-radius:16px;padding:18px;box-shadow:0 8px 24px #0f172a18;margin-bottom:14px}.grid{display:grid;grid-template-columns:360px 1fr;gap:14px}input,select{width:100%;padding:11px;border:1px solid #cbd5e1;border-radius:9px;margin:6px 0 11px}button,.btn{border:0;border-radius:9px;padding:10px 14px;background:#2563eb;color:#fff;font-weight:800;text-decoration:none;cursor:pointer}.red{background:#dc2626}.green{background:#16a34a}table{width:100%;border-collapse:collapse}th,td{padding:10px;border-bottom:1px solid #e5e7eb;text-align:left}.msg{padding:10px;border-radius:10px;background:#eff6ff;color:#1d4ed8;margin-bottom:12px}@media(max-width:800px){.grid{grid-template-columns:1fr}.box{overflow:auto}}</style></head><body><div class="head"><div class="wrap" style="padding:0"><b>🔐 Menejer va Konstruktor akkauntlari</b><a class="btn" style="float:right" href="/dashboard">Bosh sahifa</a></div></div><div class="wrap">{% if msg %}<div class="msg">{{msg}}</div>{% endif %}<div class="grid"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="save"><h3>Akkaunt yaratish yoki parolni yangilash</h3><label>Ism</label><input name="ism" required><label>Login</label><input name="login" required minlength="3"><label>Rol</label><select name="rol"><option value="menejer">Menejer</option><option value="konstruktor">Konstruktor</option></select><label>Parol</label><input name="password" type="password" minlength="8" required><button>Saqlash</button></form><div class="box"><h3>Mavjud akkauntlar</h3><table><tr><th>Ism</th><th>Login</th><th>Rol</th><th>Holat</th><th>Amal</th></tr>{% for a in accounts %}<tr><td>{{a['ism']}}</td><td>{{a['login']}}</td><td>{{'Menejer' if a['rol']=='menejer' else 'Konstruktor'}}</td><td>{{'Faol' if a['faol'] else 'Bloklangan'}}</td><td><form method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="toggle"><input type="hidden" name="account_id" value="{{a['id']}}"><button class="{{'red' if a['faol'] else 'green'}}">{{'Bloklash' if a['faol'] else 'Faollashtirish'}}</button></form></td></tr>{% else %}<tr><td colspan="5">Hali akkaunt yaratilmagan.</td></tr>{% endfor %}</table></div></div></div></body></html>
"""

MANAGER_DASHBOARD_HTML = r"""
<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Menejer kabineti</title>
<style>*{box-sizing:border-box}body{margin:0;font-family:Arial;background:#eef3f8;color:#172033}.head{background:linear-gradient(135deg,#0f1b33,#2563eb);color:#fff;padding:18px}.wrap{max-width:1250px;margin:auto;padding:16px}.box,.card{background:#fff;border-radius:16px;padding:17px;box-shadow:0 8px 24px #0f172a18;margin-bottom:14px}.cards{display:grid;grid-template-columns:repeat(4,1fr);gap:12px}.card b{font-size:25px;color:#2563eb}.grid{display:grid;grid-template-columns:370px 1fr;gap:14px}input,select,textarea{width:100%;padding:10px;border:1px solid #cbd5e1;border-radius:9px;margin:5px 0 10px}button,.btn{border:0;border-radius:9px;padding:10px 13px;background:#2563eb;color:#fff;font-weight:800;text-decoration:none;cursor:pointer}.red{background:#dc2626}.green{background:#16a34a}table{width:100%;border-collapse:collapse;font-size:13px}th,td{padding:9px;border-bottom:1px solid #e5e7eb;text-align:left;vertical-align:top}.progress{height:8px;background:#e2e8f0;border-radius:10px;overflow:hidden;min-width:90px}.progress i{display:block;height:100%;background:#16a34a}.messages{padding:10px;border-radius:10px;background:#eff6ff;color:#1d4ed8;margin-bottom:12px}@media(max-width:900px){.grid{grid-template-columns:1fr}.cards{grid-template-columns:1fr 1fr}.table{overflow:auto}}@media(max-width:500px){.cards{grid-template-columns:1fr}}</style></head><body><div class="head"><div class="wrap" style="padding:0"><b>🧑‍💼 Menejer: {{staff_name}}</b><a class="btn red" style="float:right" href="/menejer/logout">Chiqish</a></div></div><div class="wrap">{% with messages=get_flashed_messages() %}{% if messages %}<div class="messages">{{messages[0]}}</div>{% endif %}{% endwith %}<div class="cards"><div class="card"><span>JAMI BUYURTMA</span><br><b>{{stats['jami'] or 0}}</b></div><div class="card"><span>FAOL</span><br><b>{{stats['faol'] or 0}}</b></div><div class="card"><span>JAMI SUMMA (UZS)</span><br><b>{{'{:,.0f}'.format(stats['summa'] or 0)}}</b></div><div class="card"><span>TUSHUM (UZS)</span><br><b>{{'{:,.0f}'.format(stats['tushum'] or 0)}}</b></div></div><div class="grid"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="add_order"><h3>Yangi buyurtma</h3><label>Kod</label><input name="kod" placeholder="AB-001" required><label>Mijoz</label><input name="mijoz" required><label>Telefon</label><input name="telefon"><label>Manzil</label><input name="manzil"><label>Mahsulot</label><input name="mahsulot"><label>Buyurtma turi</label><select name="buyurtma_turi"><option>To‘liq mebel</option><option>Faqat oyna</option><option>MDF fasad</option><option>Ta’mirlash</option><option>Aralash buyurtma</option></select><label>Material</label><input name="material"><label>Rang</label><input name="rang"><label>Kelishuv valyutasi</label><select name="valyuta"><option value="UZS">UZS — so‘m</option><option value="USD">USD — AQSh dollari</option></select><label>Umumiy narx</label><input type="number" step="0.01" name="umumiy_narx" value="0"><label>Avans</label><input type="number" step="0.01" name="oldindan_tolov" value="0"><label>Avans valyutasi</label><select name="boshlangich_tolov_valyutasi"><option value="UZS">UZS</option><option value="USD">USD</option></select><label>To‘lov kunidagi 1 USD kursi</label><input type="number" step="0.01" name="boshlangich_kurs" value="0"><label>To‘lov sharti</label><select name="tolov_sharti"><option>Avans majburiy</option><option>Qisman avans</option><option>Avans talab qilinmaydi — ishonchli mijoz</option><option>Muddatli to‘lov</option><option>Shartnoma asosida to‘lov</option></select><label>Talab qilinadigan avans</label><input type="number" name="avans_talab" value="0"><label>Avans muddati</label><input type="date" name="avans_muddat_sana"><label>Narx taklifi amal qilish sanasi</label><input type="date" name="taklif_amal_sana"><label>Ishlab chiqarish muddati (kun)</label><input type="number" name="ishlab_chiqarish_kun" value="0"><label>Boshlanish</label><input type="date" name="boshlanish_sana" value="{{today}}"><label>Tugash</label><input type="date" name="tugash_sana"><label>Taxminiy tayyor</label><input type="date" name="taxminiy_sana"><label>Taxminiy soat</label><input type="time" name="taxminiy_vaqt" value="18:00"><label>Izoh</label><textarea name="izoh"></textarea><button>Buyurtmani saqlash</button></form><div class="box table"><h3>Buyurtmalar</h3><table><tr><th>Kod / mijoz</th><th>Mahsulot</th><th>Summa / qoldiq</th><th>Jarayon</th><th>Holat</th><th>To‘lov</th></tr>{% for o in orders %}<tr><td><b>{{o['kod']}}</b><br>{{o['mijoz']}}<br><small>{{o['telefon']}}</small></td><td>{{o['mahsulot'] or '-'}}<br><small>{{o['material'] or ''}} {{o['rang'] or ''}}</small></td><td>{% if o['valyuta']=='USD' %}${% endif %}{{'{:,.2f}'.format(o['umumiy_narx'] or 0)}} <small>{{o['valyuta'] or 'UZS'}}</small><br><b style="color:#dc2626">Qoldiq: {% if o['valyuta']=='USD' %}${% endif %}{{'{:,.2f}'.format(o['qoldiq'] or 0)}} {{o['valyuta'] or 'UZS'}}</b></td><td><div class="progress"><i style="width:{{o['progress']}}%"></i></div><small>{{o['progress']}}%</small></td><td><form method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="status"><input type="hidden" name="buyurtma_id" value="{{o['id']}}"><select name="holat"><option selected>{{o['holat']}}</option><option>Yangi</option><option>Jarayonda</option><option>Tayyor</option><option>Yetkazishga tayyor</option><option>Yetkazildi</option><option>Yopildi</option></select><button class="green">Yangilash</button></form></td><td><form method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="payment"><input type="hidden" name="buyurtma_id" value="{{o['id']}}"><input type="date" name="sana" value="{{today}}"><input type="number" step="0.01" name="miqdor" placeholder="Qabul qilingan summa" required><select name="tolov_valyutasi"><option value="UZS">UZS</option><option value="USD">USD</option></select><input type="number" step="0.01" name="kurs" placeholder="1 USD kursi"><input name="izoh" placeholder="Izoh"><button>Qo‘shish</button></form></td></tr>{% else %}<tr><td colspan="6">Buyurtma yo‘q.</td></tr>{% endfor %}</table></div></div></div></body></html>
"""

CONSTRUCTOR_DASHBOARD_HTML = r"""
<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Konstruktor kabineti</title>
<style>*{box-sizing:border-box}body{margin:0;font-family:Arial;background:#eef3f8;color:#172033}.head{background:linear-gradient(135deg,#0f1b33,#7c3aed);color:#fff;padding:18px}.wrap{max-width:1250px;margin:auto;padding:16px}.box,.card{background:#fff;border-radius:16px;padding:17px;box-shadow:0 8px 24px #0f172a18;margin-bottom:14px}.layout{display:grid;grid-template-columns:340px 1fr;gap:14px}.order{display:block;text-decoration:none;color:inherit;border-left:5px solid #7c3aed}.order.active{background:#f5f3ff}.progress{height:8px;background:#e2e8f0;border-radius:10px;overflow:hidden}.progress i{display:block;height:100%;background:#16a34a}input,select,textarea{width:100%;padding:10px;border:1px solid #cbd5e1;border-radius:9px;margin:5px 0 10px}button,.btn{border:0;border-radius:9px;padding:10px 13px;background:#7c3aed;color:#fff;font-weight:800;text-decoration:none;cursor:pointer}.red{background:#dc2626}.green{background:#16a34a}.stage{border:1px solid #e2e8f0;border-radius:12px;padding:12px;margin-bottom:9px}.stage.done{border-color:#86efac;background:#f0fdf4}.files a{color:#2563eb;font-weight:700}.messages{padding:10px;border-radius:10px;background:#f5f3ff;color:#6d28d9;margin-bottom:12px}@media(max-width:850px){.layout{grid-template-columns:1fr}}</style></head><body><div class="head"><div class="wrap" style="padding:0"><b>📐 Konstruktor: {{staff_name}}</b><a class="btn red" style="float:right" href="/konstruktor/logout">Chiqish</a></div></div><div class="wrap">{% with messages=get_flashed_messages() %}{% if messages %}<div class="messages">{{messages[0]}}</div>{% endif %}{% endwith %}<div class="layout"><div>{% for o in orders %}<a class="card order {% if selected_order and selected_order['id']==o['id'] %}active{% endif %}" href="/konstruktor/kabinet?order_id={{o['id']}}"><b>{{o['kod']}} — {{o['mahsulot'] or 'Mahsulot'}}</b><p>{{o['mijoz']}}</p><div class="progress"><i style="width:{{o['progress']}}%"></i></div><small>{{o['progress']}}% · {{o['holat']}}</small></a>{% else %}<div class="card">Buyurtma yo‘q.</div>{% endfor %}</div><div>{% if selected_order %}<div class="box"><h2>{{selected_order['kod']}} — {{selected_order['mahsulot']}}</h2><p><b>Mijoz:</b> {{selected_order['mijoz']}} · <b>Material:</b> {{selected_order['material'] or '-'}} · <b>Rang:</b> {{selected_order['rang'] or '-'}}</p><p><b>O‘lcham:</b> {{selected_order['olcham'] or '-'}} · <b>Soni:</b> {{selected_order['soni'] or 1}}</p><form method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="order_status"><input type="hidden" name="order_id" value="{{selected_order['id']}}"><label>Buyurtma holati</label><select name="holat"><option selected>{{selected_order['holat']}}</option><option>Chizma tayyorlanmoqda</option><option>Mijoz tasdiqlashi kutilmoqda</option><option>Material tayyorlanmoqda</option><option>Jarayonda</option><option>Tayyor</option></select><button>Holatni yangilash</button></form></div><div class="box"><h3>Chizma, kroy va CNC fayli</h3><form method="post" enctype="multipart/form-data"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="upload"><input type="hidden" name="order_id" value="{{selected_order['id']}}"><label>Fayl nomi/izohi</label><input name="nomi" placeholder="PRO100 chizma / CNC fayli"><label>Fayl</label><input type="file" name="fayl" required><button class="green">Faylni yuklash</button></form><div class="files">{% for f in files %}<p>📎 <a href="/konstruktor/fayl/{{f['fayl_nomi']}}">{{f['nomi']}}</a> <small>({{f['created_at']}})</small></p>{% else %}<p>Hali fayl yuklanmagan.</p>{% endfor %}</div></div><div class="box"><h3>Ishlab chiqarish bosqichlari</h3>{% for s in stages %}<form class="stage {% if s['bajarildi'] %}done{% endif %}" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="stage"><input type="hidden" name="order_id" value="{{selected_order['id']}}"><input type="hidden" name="stage_id" value="{{s['id']}}"><b>{{'✅' if s['bajarildi'] else '⬜'}} {{s['bosqich']}}</b><p><small>Boshlanish: {{s['boshlanish_vaqti'] or '-'}} · Tugash: {{s['tugash_vaqti'] or '-'}}</small></p><input name="izoh" value="{{s['izoh'] or ''}}" placeholder="Izoh"><select name="bajarildi"><option value="0" {% if not s['bajarildi'] %}selected{% endif %}>Jarayonda</option><option value="1" {% if s['bajarildi'] %}selected{% endif %}>Bajarildi</option></select><button>Saqlash</button></form>{% endfor %}</div>{% else %}<div class="box">Buyurtmani tanlang.</div>{% endif %}</div></div></div></body></html>
"""


ADMIN_SETUP_HTML = r"""
<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Mebel360° - Birinchi xavfsiz sozlash</title><style>body{margin:0;background:linear-gradient(135deg,#0f1b33,#16a34a);font-family:Arial;display:grid;place-items:center;min-height:100vh}.box{background:#fff;padding:28px;border-radius:18px;width:min(410px,92%);box-shadow:0 20px 50px #0005}h2{margin-top:0}input{width:100%;padding:12px;margin:8px 0;border:1px solid #cbd5e1;border-radius:9px;box-sizing:border-box}button{width:100%;padding:12px;border:0;border-radius:9px;background:#16a34a;color:#fff;font-weight:700}.err{color:#b91c1c;font-size:13px}.note{font-size:13px;color:#475569;line-height:1.45}</style></head><body><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h2>🔐 Birinchi xavfsiz sozlash</h2><input name="user" placeholder="Admin login" value="admin" minlength="3" required><input name="password" type="password" placeholder="Yangi parol" minlength="8" required><input name="confirm" type="password" placeholder="Parolni takrorlang" minlength="8" required><button>Admin akkauntini yaratish</button><div class="err">{{error}}</div></form></body></html>
"""

LOGIN_HTML = r"""
<!doctype html>
<html lang="uz">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Mebel360° — Rahbar kirishi</title>
<style>
:root{
  --navy:#0b1736;
  --blue:#2563eb;
  --blue2:#38bdf8;
  --green:#16a34a;
  --text:#172033;
  --muted:#64748b;
  --line:#dbe5f0;
  --danger:#b91c1c;
}
*{box-sizing:border-box}
html,body{margin:0;min-height:100%;font-family:Arial,Helvetica,sans-serif;color:var(--text)}
body{
  min-height:100vh;
  display:grid;
  place-items:center;
  padding:24px;
  background:
    radial-gradient(circle at 12% 16%,rgba(56,189,248,.28),transparent 31%),
    radial-gradient(circle at 88% 82%,rgba(37,99,235,.30),transparent 34%),
    linear-gradient(145deg,#07142f 0%,#123b8f 50%,#0b5fa8 100%);
  overflow-x:hidden;
}
body:before,body:after{
  content:"";position:fixed;border-radius:999px;filter:blur(2px);pointer-events:none;
  background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.12)
}
body:before{width:320px;height:320px;left:-100px;bottom:-120px}
body:after{width:230px;height:230px;right:-65px;top:-60px}
.login-shell{
  width:min(900px,100%);
  display:grid;
  grid-template-columns:1.05fr .95fr;
  border:1px solid rgba(255,255,255,.22);
  border-radius:28px;
  overflow:hidden;
  background:rgba(255,255,255,.96);
  box-shadow:0 32px 90px rgba(2,8,23,.38);
  animation:show .55s ease both;
}
.brand-panel{
  position:relative;
  min-height:560px;
  padding:48px 42px;
  display:flex;
  flex-direction:column;
  justify-content:space-between;
  color:#fff;
  background:
    radial-gradient(circle at 78% 20%,rgba(56,189,248,.34),transparent 30%),
    linear-gradient(150deg,#0b1736,#154aa8 62%,#0784c8);
  overflow:hidden;
}
.brand-panel:after{
  content:"360°";
  position:absolute;
  right:-18px;
  bottom:18px;
  font-size:132px;
  line-height:1;
  font-weight:900;
  color:rgba(255,255,255,.055);
  letter-spacing:-10px;
}
.brand-top{position:relative;z-index:1}
.logo-mark{
  width:66px;height:66px;border-radius:20px;
  display:grid;place-items:center;
  background:linear-gradient(145deg,#fff,#dbeafe);
  color:#154aa8;font-size:31px;font-weight:900;
  box-shadow:0 15px 35px rgba(2,8,23,.25)
}
.brand-title{margin:28px 0 8px;font-size:42px;line-height:1.05;font-weight:900;letter-spacing:-1px}
.brand-sub{font-size:18px;line-height:1.55;color:#dbeafe;max-width:390px}
.brand-features{position:relative;z-index:1;display:grid;gap:11px;margin-top:28px}
.feature{display:flex;align-items:center;gap:10px;font-size:14px;color:#e0f2fe}
.feature i{width:27px;height:27px;border-radius:9px;display:grid;place-items:center;background:rgba(255,255,255,.13);font-style:normal}
.form-panel{padding:48px 44px;display:flex;align-items:center;background:#fff}
.form-wrap{width:100%;max-width:360px;margin:auto}
.role-badge{display:inline-flex;align-items:center;gap:7px;padding:7px 11px;border-radius:999px;background:#eff6ff;color:#1d4ed8;font-size:12px;font-weight:800}
h1{font-size:30px;margin:17px 0 7px;letter-spacing:-.5px}
.intro{margin:0 0 25px;color:var(--muted);font-size:14px;line-height:1.5}
label{display:block;margin:13px 0 6px;font-size:13px;font-weight:800;color:#334155}
.field{position:relative}
.field input{
  width:100%;height:50px;padding:0 46px 0 43px;
  border:1px solid var(--line);border-radius:13px;
  background:#f8fafc;color:#0f172a;font-size:15px;outline:none;
  transition:.18s ease;
}
.field input:focus{border-color:#60a5fa;background:#fff;box-shadow:0 0 0 4px rgba(37,99,235,.11)}
.field-icon{position:absolute;left:15px;top:50%;transform:translateY(-50%);font-size:17px;opacity:.72}
.show-pass{position:absolute;right:8px;top:50%;transform:translateY(-50%);border:0;background:transparent;color:#475569;font-size:12px;font-weight:800;padding:8px;cursor:pointer}
.login-btn{
  width:100%;height:51px;margin-top:20px;border:0;border-radius:13px;
  color:#fff;font-size:15px;font-weight:900;cursor:pointer;
  background:linear-gradient(100deg,#2563eb,#0284c7);
  box-shadow:0 13px 28px rgba(37,99,235,.25);
  transition:.18s ease;
}
.login-btn:hover{transform:translateY(-1px);box-shadow:0 16px 32px rgba(37,99,235,.30)}
.login-btn:active{transform:translateY(0)}
.error-box{margin-top:13px;padding:10px 12px;border-radius:10px;background:#fef2f2;color:var(--danger);font-size:13px;font-weight:700;border:1px solid #fecaca}
.other-title{display:flex;align-items:center;gap:10px;margin:24px 0 13px;color:#94a3b8;font-size:11px;font-weight:900;text-transform:uppercase;letter-spacing:.8px}
.other-title:before,.other-title:after{content:"";height:1px;background:#e2e8f0;flex:1}
.role-grid{display:grid;grid-template-columns:1fr 1fr;gap:9px}
.role-link{
  min-height:67px;padding:10px;border:1px solid #e2e8f0;border-radius:13px;
  text-decoration:none;color:#334155;background:#f8fafc;
  display:flex;align-items:center;gap:9px;transition:.18s ease
}
.role-link:hover{border-color:#93c5fd;background:#eff6ff;transform:translateY(-1px)}
.role-icon{width:35px;height:35px;border-radius:10px;display:grid;place-items:center;background:#dbeafe;font-size:17px;flex:none}
.role-link.driver .role-icon{background:#dcfce7}
.role-link b{display:block;font-size:12px}.role-link small{display:block;color:#64748b;font-size:10px;margin-top:3px}
.creator-note{margin-top:20px;padding-top:16px;border-top:1px solid #e2e8f0;text-align:center;line-height:1.35}
.creator-system{font-size:13px;font-weight:900;color:#334155;letter-spacing:.15px}
.creator-name{margin-top:5px;font-size:12px;font-weight:900;color:#1d4ed8}
.creator-story{max-width:330px;margin:5px auto 0;color:#7c8ca0;font-size:10.5px;line-height:1.45}
@keyframes show{from{opacity:0;transform:translateY(12px) scale(.985)}to{opacity:1;transform:none}}
@media(max-width:760px){
  body{padding:13px;display:block}
  .login-shell{grid-template-columns:1fr;max-width:480px;margin:12px auto;border-radius:22px}
  .brand-panel{min-height:auto;padding:26px 25px 24px}
  .brand-title{font-size:29px;margin-top:18px}.brand-sub{font-size:14px}
  .brand-features{display:none}.brand-panel:after{font-size:85px;bottom:-5px}
  .logo-mark{width:52px;height:52px;border-radius:16px;font-size:24px}
  .form-panel{padding:29px 22px 25px}
  h1{font-size:26px}.role-grid{grid-template-columns:1fr 1fr}
}
@media(max-width:390px){.role-grid{grid-template-columns:1fr}.role-link{min-height:58px}}
</style>
</head>
<body>
<main class="login-shell">
  <section class="brand-panel">
    <div class="brand-top">
      <div class="logo-mark">M</div>
      <div class="brand-title">Mebel360°</div>
      <div class="brand-sub">Korxona, buyurtmalar, ishchilar va omborni yagona tizimda boshqaring.</div>
    </div>
    <div class="brand-features">
      <div class="feature"><i>✓</i><span>Buyurtmalar va ishlab chiqarish nazorati</span></div>
      <div class="feature"><i>✓</i><span>Menejer, Konstruktor, Ishchi va Shofyor kabinetlari</span></div>
      <div class="feature"><i>✓</i><span>Ombor, xarajat va hisobotlar</span></div>
    </div>
  </section>

  <section class="form-panel">
    <form class="form-wrap" method="post">
      <input type="hidden" name="csrf_token" value="{{csrf_token()}}">
      <span class="role-badge">◆ Rahbar kabineti</span>
      <h1>Xush kelibsiz</h1>
      <p class="intro">Dastur boshqaruv paneliga kirish uchun login va parolingizni kiriting.</p>

      <label for="user">Login</label>
      <div class="field">
        <span class="field-icon">👤</span>
        <input id="user" name="user" placeholder="Login" value="admin" autocomplete="username" required autofocus>
      </div>

      <label for="password">Parol</label>
      <div class="field">
        <span class="field-icon">🔑</span>
        <input id="password" name="password" type="password" placeholder="Parol" autocomplete="current-password" required>
        <button class="show-pass" type="button" onclick="togglePassword()" id="showPassword">Ko‘rsatish</button>
      </div>

      <button class="login-btn" type="submit">Kirish →</button>
      {% if error %}<div class="error-box">⚠ {{error}}</div>{% endif %}

      <div class="other-title">Boshqa kabinetlar</div>
      <div class="role-grid">
        <a class="role-link" href="/menejer/login"><span class="role-icon">🧑‍💼</span><span><b>Menejer kirishi</b><small>Buyurtmalar va mijozlar</small></span></a>
        <a class="role-link" href="/konstruktor/login"><span class="role-icon">📐</span><span><b>Konstruktor kirishi</b><small>Chizma, kroy va CNC</small></span></a>
        <a class="role-link" href="/ishchi/login"><span class="role-icon">👷</span><span><b>Ishchi kirishi</b><small>Kabinetga kirish</small></span></a>
        <a class="role-link" href="/ishchi/royxat"><span class="role-icon">＋</span><span><b>Ishchi ro‘yxati</b><small>Yangi akkaunt</small></span></a>
        <a class="role-link driver" href="/shofyor/login"><span class="role-icon">🚚</span><span><b>Shofyor kirishi</b><small>Kabinetga kirish</small></span></a>
        <a class="role-link driver" href="/shofyor/royxat"><span class="role-icon">＋</span><span><b>Shofyor ro‘yxati</b><small>Yangi akkaunt</small></span></a>
      </div>
      <div class="creator-note">
        <div class="creator-system">Mebel360° boshqaruv tizimi</div>
        <div class="creator-name">Zuhriddin Ubaydullayev</div>
        <div class="creator-story">22 yillik amaliy tajriba va yo‘l qo‘yilgan xatolardan olingan saboqlar asosida shakllangan tizim</div>
      </div>
    </form>
  </section>
</main>
<script>
function togglePassword(){
  const input=document.getElementById('password');
  const button=document.getElementById('showPassword');
  const visible=input.type==='text';
  input.type=visible?'password':'text';
  button.textContent=visible?'Ko‘rsatish':'Yashirish';
}
</script>
</body>
</html>
"""


WORKER_BASE_STYLE = """
<style>
*{box-sizing:border-box}body{margin:0;font-family:Arial;background:#eef3f8;color:#182235}.head{background:linear-gradient(135deg,#0f1b33,#2563eb);color:white;padding:18px}.wrap{max-width:1050px;margin:auto;padding:16px}.box,.card{background:white;border-radius:16px;padding:18px;box-shadow:0 8px 24px #0f172a18;margin-bottom:14px}input,select,textarea{width:100%;padding:11px;border:1px solid #cbd5e1;border-radius:9px;margin:5px 0 10px}button,.btn{display:inline-block;border:0;border-radius:9px;padding:10px 14px;background:#2563eb;color:white;font-weight:700;text-decoration:none;cursor:pointer}.green{background:#16a34a}.red{background:#dc2626}.muted{color:#64748b;font-size:13px}.err{color:#b91c1c}.ok{color:#166534}.grid{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}.stat b{font-size:25px;color:#2563eb}.task{border-left:5px solid #2563eb}.bar{height:10px;background:#e2e8f0;border-radius:20px;overflow:hidden}.bar i{display:block;height:100%;background:#16a34a}table{width:100%;border-collapse:collapse;font-size:13px}th,td{padding:8px;border-bottom:1px solid #e5e7eb;text-align:left}@media(max-width:700px){.grid{grid-template-columns:1fr}.wrap{padding:9px}}
</style>
"""

DRIVER_REGISTER_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Shofyor ro‘yxatdan o‘tishi</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>🚚 Shofyor ro‘yxatdan o‘tishi</b></div><div class="wrap"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h2>Yangi akkaunt</h2><label>Ism</label><input name="ism" required><label>Familiya</label><input name="familiya"><label>Telefon</label><input name="telefon" placeholder="+998901234567" required><label>Login</label><input name="login" required><label>Parol</label><input name="password" type="password" minlength="8" required><button>Ro‘yxatdan o‘tish</button><p class="ok">{{msg}}</p><p class="err">{{error}}</p><p><a href="/shofyor/login">Akkauntim bor — kirish</a></p></form></div></body></html>"""

DRIVER_LOGIN_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Shofyor kirishi</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>🚚 Mebel360° — Shofyor kabineti</b></div><div class="wrap"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h2>Kirish</h2><input name="login" placeholder="Login" required><input name="password" type="password" placeholder="Parol" required><button>Kirish</button><p class="err">{{error}}</p><p><a href="/shofyor/royxat">Yangi shofyor — ro‘yxatdan o‘tish</a></p><p><a href="/login">Rahbar kirishi</a></p></form></div></body></html>"""

DRIVER_DASHBOARD_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Shofyor kabineti</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><div class="wrap" style="padding:0"><b>🚚 {{driver['ism']}} {{driver['familiya']}}</b><a class="btn red" style="float:right" href="/shofyor/logout">Chiqish</a></div></div><div class="wrap"><h2>Menga biriktirilgan yuklar</h2>{% for x in deliveries %}<a href="/shofyor/yetkazish/{{x['id']}}" style="text-decoration:none;color:inherit"><div class="card task"><div style="display:flex;justify-content:space-between;gap:10px"><div><b style="font-size:21px">{{x['kod']}}</b><p>{{x['mahsulot']}}</p><p class="muted">{{x['manzil']}} · {{x['qavat'] or '-'}}-qavat · Lift: {{x['lift'] or '-'}}</p></div><div><span class="btn {% if x['holat']=='Yetkazib berildi' %}green{% endif %}">{{x['holat']}}</span></div></div></div></a>{% else %}<div class="card">Hozircha yuk biriktirilmagan.</div>{% endfor %}</div></body></html>"""

DRIVER_DETAIL_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{{x['kod']}}</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><div class="wrap" style="padding:0"><b>📦 {{x['kod']}}</b><a class="btn" style="float:right" href="/shofyor/kabinet">Orqaga</a></div></div><div class="wrap"><div class="card"><h2>{{x['mahsulot']}}</h2><p><b>Mijoz:</b> {{x['mijoz']}}</p><p><b>Telefon:</b> <a href="tel:{{x['telefon']}}">{{x['telefon']}}</a></p><p><b>Manzil:</b> {{x['manzil']}}</p><p><b>Mo‘ljal:</b> {{x['moljal'] or '-'}}</p><p><b>Qavat:</b> {{x['qavat'] or '-'}}</p><p><b>Lift:</b> {{x['lift'] or '-'}}</p><p><b>Katta mashina:</b> {{x['katta_mashina'] or '-'}}</p><p><b>Qadoqlar:</b> {{x['qadoq_soni']}} ta</p><p><b>Yetkazish navbati:</b> {{x['navbat']}}</p><p><b>Izoh:</b> {{x['izoh'] or x['buyurtma_izoh'] or '-'}}</p>{% if x['lokatsiya'] %}
<div class="card" style="background:#f8fafc">
<h3>📍 Xarita tanlang</h3>
<p class="muted">Lokatsiyani o‘zingiz xohlagan xaritada oching.</p>
<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px">
<a class="btn green" href="https://www.google.com/maps/search/?api=1&query={{x['lokatsiya']|urlencode}}" target="_blank">Google Maps</a>
<a class="btn" style="background:#ef4444" href="https://yandex.com/maps/?text={{x['lokatsiya']|urlencode}}" target="_blank">Yandex Maps</a>
</div></div>
{% endif %}</div><div class="card"><h3>Holat: {{x['holat']}}</h3><form method="post" action="/shofyor/yetkazish/{{x['id']}}/holat"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><button name="action" value="yolga">🚚 Yo‘lga chiqdim</button><button name="action" value="yetib" class="green">📍 Yetib keldim</button><button name="action" value="yetkazildi" style="background:#7c3aed">✅ Yetkazib berdim</button></form><p class="muted">Yo‘lga chiqdi: {{x['yolga_chiqdi'] or '-'}}</p><p class="muted">Yetib keldi: {{x['yetib_keldi'] or '-'}}</p><p class="muted">Yetkazildi: {{x['yetkazildi'] or '-'}}</p></div></div></body></html>"""

DRIVER_ADMIN_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Shofyor boshqaruvi</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>🚚 Shofyor boshqaruvi</b><a class="btn" style="float:right" href="/dashboard">ERP bosh sahifa</a></div><div class="wrap"><p class="ok">{{msg}}</p><div class="grid"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h3>Shofyor login yaratish</h3><input type="hidden" name="action" value="account"><label>Shofyor</label><select name="ishchi_id" required>{% for d in drivers %}<option value="{{d['id']}}">{{d['ism']}} {{d['familiya']}}</option>{% endfor %}</select><label>Login</label><input name="login" required><label>Parol</label><input name="password" type="password" minlength="8" required><button>Akkaunt yaratish</button></form><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h3>Yuk biriktirish</h3><input type="hidden" name="action" value="delivery"><label>Buyurtma</label><select name="buyurtma_id" required>{% for o in orders %}<option value="{{o['id']}}">{{o['kod']}} — {{o['mahsulot']}}</option>{% endfor %}</select><label>Shofyor</label><select name="shofyor_id" required>{% for d in drivers %}<option value="{{d['id']}}">{{d['ism']}} {{d['familiya']}}</option>{% endfor %}</select><label>Qadoq soni</label><input type="number" name="qadoq_soni" value="1" min="1"><label>Yetkazish navbati</label><input type="number" name="navbat" value="1" min="1"><label>Izoh</label><textarea name="izoh"></textarea><button>Biriktirish</button></form></div><div class="box" style="overflow:auto"><h3>Shofyor akkauntlari</h3><table><tr><th>Shofyor</th><th>Telefon</th><th>Login</th><th>Holat</th><th>Amal</th></tr>{% for a in driver_accounts %}<tr><td>{{a['ism']}} {{a['familiya']}}</td><td>{{a['telefon'] or '-'}}</td><td>{{a['login']}}</td><td>{% if a['admin_tasdiq'] %}✅ Tasdiqlangan{% else %}⏳ Kutilmoqda{% endif %}</td><td>{% if not a['admin_tasdiq'] %}<form method="post" style="display:inline"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="approve_driver"><input type="hidden" name="account_id" value="{{a['id']}}"><button class="green">Tasdiqlash</button></form>{% endif %}<form method="post" style="display:inline"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="block_driver"><input type="hidden" name="account_id" value="{{a['id']}}"><button class="red">Bloklash</button></form></td></tr>{% endfor %}</table></div><div class="box" style="overflow:auto"><h3>Biriktirilgan yuklar</h3><table><tr><th>Buyurtma</th><th>Mahsulot</th><th>Shofyor</th><th>Qadoq</th><th>Holat</th></tr>{% for a in assigned %}<tr><td>{{a['kod']}}</td><td>{{a['mahsulot']}}</td><td>{{a['ism']}} {{a['familiya']}}</td><td>{{a['qadoq_soni']}}</td><td>{{a['holat']}}</td></tr>{% endfor %}</table></div></div></body></html>"""


WORKER_REGISTER_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Ishchi ro‘yxati</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>👷 Ishchi ro‘yxatdan o‘tishi</b></div><div class="wrap"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h2>Telefon raqamingiz</h2><p class="muted">Masalan: +998 90 123 45 67</p><input name="telefon" required placeholder="+998901234567"><button>Kod olish</button><p class="ok">{{msg}}</p><p class="err">{{error}}</p>{% if demo_code %}<div class="card"><b>Sinov kodi: {{demo_code}}</b><p class="muted">Haqiqiy SMS xizmati ulanmaguncha shu koddan foydalaning.</p><a class="btn green" href="/ishchi/kod">Kodni kiritish</a></div>{% endif %}<p><a href="/ishchi/login">Akkauntim bor — kirish</a></p></form></div></body></html>"""

WORKER_VERIFY_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Kodni tasdiqlash</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>🔐 Telefonni tasdiqlash</b></div><div class="wrap"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><p>{{telefon}}</p><label>Kod</label><input name="kod" inputmode="numeric" maxlength="6" required><label>Ism</label><input name="ism" required><label>Familiya</label><input name="familiya"><label>Yangi login</label><input name="login" required><label>Yangi parol</label><input name="password" type="password" minlength="8" required><button>Ro‘yxatdan o‘tish</button><p class="err">{{error}}</p></form></div></body></html>"""

WORKER_WAIT_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Kutilmoqda</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="wrap"><div class="box"><h2>✅ Ro‘yxatdan o‘tdingiz</h2><p>Endi administrator akkauntingizni tasdiqlaydi. Tasdiqlangach login va parolingiz bilan kirasiz.</p><a class="btn" href="/ishchi/login">Kirish sahifasi</a></div></div></body></html>"""

WORKER_LOGIN_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Ishchi kirishi</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>🏭 Mebel360° — Ishchi kabineti</b></div><div class="wrap"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h2>Kirish</h2><input name="login" placeholder="Login" required><input name="password" type="password" placeholder="Parol" required><button>Kirish</button><p class="err">{{error}}</p><p><a href="/ishchi/royxat">Yangi ro‘yxatdan o‘tish</a></p></form></div></body></html>"""

WORKER_DASHBOARD_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Ishchi kabineti</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><div class="wrap" style="padding:0"><b>👷 {{worker['ism']}} {{worker['familiya']}}</b> — {{worker['lavozim']}} <a class="btn red" style="float:right" href="/ishchi/logout">Chiqish</a></div></div><div class="wrap">
{% with messages = get_flashed_messages() %}{% if messages %}<div class="card err">{{messages[0]}}</div>{% endif %}{% endwith %}
<div class="grid"><div class="card stat"><span>HOZIRGI HOLAT</span><br><b style="color:{% if worker_state=='Ishlayapti' %}#16a34a{% else %}#64748b{% endif %}">{{worker_state}}</b></div><div class="card stat"><span>ISHLAGAN KUN</span><br><b>{{stats['kun'] or 0}}</b></div><div class="card stat"><span>JAMI SOAT</span><br><b>{{'%.1f'|format(stats['soat'] or 0)}}</b></div></div>
{% if active_task %}<div class="card" style="border-left:7px solid #16a34a"><h3>🟢 Hozir ishlayotgan ishim</h3><p><b>{{active_task['ish_turi']}}</b> — {{active_task['buyurtma_kodi'] or 'Buyurtmasiz'}}</p><p>{{active_task['tavsif']}}</p><p>Boshlangan vaqt: <b>{{active_task['boshlandi_vaqt']}}</b></p></div>{% endif %}
<h2>Topshiriqlarim</h2>{% for t in tasks %}<div class="card task"><b>{{t['ish_turi']}}</b> {% if t['buyurtma_kodi'] %}<span class="muted">— {{t['buyurtma_kodi']}}</span>{% endif %}<p>{{t['tavsif']}}</p><div class="bar"><i style="width:{{t['progress']}}%"></i></div><p><b>{{t['progress']}}%</b> · {{t['holat']}} · Reja: {{t['sana']}}{% if t['tugash_sana'] %} — {{t['tugash_sana']}}{% endif %}</p><p class="muted">Boshladi: {{t['boshlandi_vaqt'] or '-'}} · Tugatdi: {{t['tugadi_vaqt'] or '-'}}</p>
{% if t['holat']=='Yangi' or t['holat']=='Jarayonda' %}<form method="post" action="/ishchi/topshiriq/{{t['id']}}/boshlash"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><button class="green">▶ Ishni boshladim</button></form>{% elif t['holat']=='Ishlayapti' %}<form method="post" action="/ishchi/topshiriq/{{t['id']}}/tugatish"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><button style="background:#7c3aed">✅ Ishni tugatdim</button></form>{% endif %}
{% if t['holat']!='Tayyor' %}<form method="post" action="/ishchi/topshiriq/{{t['id']}}"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="holat" value="{{t['holat']}}"><label>Jarayon foizi</label><input type="number" name="progress" min="0" max="100" value="{{t['progress']}}"><button>Foizni yangilash</button></form>{% endif %}</div>{% else %}<div class="card">Hozircha topshiriq yo‘q. Holatingiz: <b>Bo‘sh</b>.</div>{% endfor %}</div></body></html>"""

WORKER_ADMIN_HTML = r"""<!doctype html><html lang="uz"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Ishchi boshqaruvi</title>"""+WORKER_BASE_STYLE+r"""</head><body><div class="head"><b>👥 Ishchi kabinetlarini boshqarish</b> <a class="btn" style="float:right" href="/dashboard">ERP bosh sahifa</a></div><div class="wrap">
<div class="box" style="overflow:auto"><h3>Ishchilarning hozirgi holati</h3><table><tr><th>Ishchi</th><th>Holat</th><th>Buyurtma</th><th>Ish turi</th><th>Boshlagan vaqt</th></tr>{% for w in worker_states %}<tr><td>{{w['ism']}} {{w['familiya']}}</td><td>{% if w['ish_holati']=='Ishlayapti' %}<b style="color:#16a34a">🟢 Ishlayapti</b>{% else %}<b style="color:#64748b">⚪ Bo‘sh</b>{% endif %}</td><td>{{w['buyurtma_kodi'] or '-'}}</td><td>{{w['ish_turi'] or '-'}}</td><td>{{w['boshlandi_vaqt'] or '-'}}</td></tr>{% endfor %}</table></div>
<div class="grid"><form class="box" method="post"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><h3>Yangi topshiriq</h3><input type="hidden" name="action" value="task"><label>Ishchi</label><select name="ishchi_id" required>{% for w in workers %}<option value="{{w['id']}}">{{w['ism']}} {{w['familiya']}} — {{w['lavozim']}}</option>{% endfor %}</select><label>Buyurtma kodi</label><input name="buyurtma_kodi" placeholder="AB 007"><label>Ish turi</label><input name="ish_turi" required placeholder="Kesish / Rover / Yig‘ish"><label>Topshiriq</label><textarea name="tavsif" placeholder="Nima ish qilishi kerakligini batafsil yozing"></textarea><label>Boshlanish rejasi</label><input type="date" name="sana" value="{{today}}"><label>Tugash rejasi</label><input type="date" name="tugash_sana"><button>Topshiriq berish</button></form><div class="box" style="grid-column:span 2;overflow:auto"><h3>Ro‘yxatdan o‘tganlar</h3><table><tr><th>Ishchi</th><th>Telefon</th><th>Login</th><th>Holat</th><th>Amal</th></tr>{% for a in accounts %}<tr><td>{{a['ism']}} {{a['familiya']}}</td><td>{{a['telefon']}}</td><td>{{a['login'] or ''}}</td><td>{% if a['admin_tasdiq'] %}✅ Tasdiqlangan{% else %}⏳ Kutilmoqda{% endif %}</td><td>{% if not a['admin_tasdiq'] %}<form method="post" style="display:inline"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="approve"><input type="hidden" name="account_id" value="{{a['id']}}"><button class="green">Tasdiqlash</button></form>{% endif %}<form method="post" style="display:inline"><input type="hidden" name="csrf_token" value="{{csrf_token()}}"><input type="hidden" name="action" value="block"><input type="hidden" name="account_id" value="{{a['id']}}"><button class="red">Bloklash</button></form></td></tr>{% endfor %}</table></div></div>
<div class="box" style="overflow:auto"><h3>Topshiriqlar tarixi</h3><table><tr><th>Ishchi</th><th>Buyurtma</th><th>Ish</th><th>Holat</th><th>Progress</th><th>Boshladi</th><th>Tugatdi</th></tr>{% for t in tasks %}<tr><td>{{t['ism']}} {{t['familiya']}}</td><td>{{t['buyurtma_kodi']}}</td><td>{{t['ish_turi']}}</td><td>{{t['holat']}}</td><td>{{t['progress']}}%</td><td>{{t['boshlandi_vaqt'] or '-'}}</td><td>{{t['tugadi_vaqt'] or '-'}}</td></tr>{% endfor %}</table></div></div></body></html>"""


HTML = r"""
<!doctype html>
<html lang="uz">
<head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Mebel360°</title>
<style>
:root{--nav:#18213f;--blue:#315bd8;--purple:#6d4aff;--bg:#f4f6fb;--card:#fff;--text:#172039;--muted:#68728a;--danger:#e33d61;--ok:#19a66a;--line:#e6e9f2;--soft:#f7f8fc}
*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;background:var(--bg);color:var(--text)}
header{background:linear-gradient(118deg,#232e7a 0%,#4248b8 48%,#7452d7 100%);color:#fff;padding:18px 22px 20px;box-shadow:0 12px 34px rgba(40,48,118,.24)}
.top{max-width:1680px;margin:auto;display:grid;grid-template-columns:minmax(310px,1fr) auto minmax(590px,1.5fr);align-items:center;gap:18px}
.brand{display:flex;align-items:center;gap:13px;min-width:0}.brand-logo{width:92px;height:64px;object-fit:contain;background:#fff;border-radius:17px;padding:5px;box-shadow:0 8px 24px rgba(10,18,64,.26)}.brand-text h1{margin:0;font-size:29px;letter-spacing:-.6px}.sub{opacity:.88;font-size:13px;margin-top:4px;line-height:1.35}.live-clock{min-width:190px;text-align:center;background:rgba(255,255,255,.13);border:1px solid rgba(255,255,255,.22);border-radius:17px;padding:10px 15px;backdrop-filter:blur(8px)}.live-clock-time{font-size:28px;font-weight:900;letter-spacing:2.2px;line-height:1.05;font-variant-numeric:tabular-nums}.live-clock-date{font-size:11px;opacity:.92;margin-top:5px;white-space:nowrap}.header-actions{display:flex;flex-wrap:wrap;justify-content:flex-end;gap:7px}.header-actions a{display:inline-flex;text-decoration:none}.header-actions button{white-space:nowrap;background:rgba(16,24,64,.55)!important;border:1px solid rgba(255,255,255,.16);box-shadow:none;padding:8px 11px;font-size:12px}.header-actions a:nth-last-child(2) button{background:#19a66a!important}.header-actions a:last-child button{background:#e33d61!important}
.wrap{max-width:1680px;margin:auto;padding:18px 22px 34px}.cards{display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:12px;margin-bottom:16px}.card,.panel{background:var(--card);border:1px solid rgba(225,228,239,.85);border-radius:18px;box-shadow:0 8px 24px rgba(33,43,84,.075);padding:16px}.card{position:relative;overflow:hidden;min-height:90px;padding-left:68px;display:flex;flex-direction:column;justify-content:center}.card:before{content:"";position:absolute;left:16px;top:22px;width:38px;height:38px;border-radius:12px;background:#eaf0ff}.card:nth-child(2):before{background:#e8f8f1}.card:nth-child(3):before{background:#f1eaff}.card:nth-child(4):before{background:#fff1e5}.card:nth-child(5):before{background:#e6f8fb}.card:nth-child(6):before{background:#ffe9ee}.card span{font-size:10px;color:var(--muted);font-weight:900;letter-spacing:.35px}.card b{display:block;font-size:27px;color:var(--blue);margin-top:5px}.card:nth-child(2) b{color:#14915c}.card:nth-child(3) b{color:#7047e8}.card:nth-child(4) b{color:#df7a20}.card:nth-child(5) b{color:#1497ad}.card:nth-child(6) b{color:#dc3156}
.tabs{display:flex;gap:5px;overflow-x:auto;flex-wrap:nowrap;margin-bottom:15px;padding:8px;background:#fff;border:1px solid var(--line);border-radius:15px;box-shadow:0 5px 18px rgba(33,43,84,.055);scrollbar-width:thin}.tabs button{flex:0 0 auto;background:transparent;color:#59637a;padding:9px 12px;border-radius:10px;font-size:12px}.tabs button.active{background:linear-gradient(135deg,#315bd8,#684ce2);color:#fff;box-shadow:0 5px 14px rgba(67,77,194,.24)}
button{border:0;border-radius:10px;padding:9px 13px;background:var(--blue);color:#fff;font-weight:800;cursor:pointer;transition:.18s ease}button:hover{transform:translateY(-1px);filter:brightness(.97)}button:disabled{opacity:.55;cursor:not-allowed;transform:none}.tab{display:none}.tab.active{display:block}.grid{display:grid;grid-template-columns:380px minmax(0,1fr);gap:15px}#orders .grid{grid-template-columns:410px minmax(0,1fr)}
h3{margin:0 0 13px;font-size:17px;letter-spacing:-.2px}label{display:block;font-size:12px;font-weight:800;margin-top:9px;color:#3f4961}input,select,textarea{width:100%;margin-top:5px;padding:10px 11px;border:1px solid #d9deea;border-radius:10px;background:#fff;color:#172039;outline:none;transition:.16s}input:focus,select:focus,textarea:focus{border-color:#6b72e8;box-shadow:0 0 0 3px rgba(91,97,219,.12)}textarea{min-height:70px;resize:vertical}form>button{width:100%;margin-top:13px}.form-section{padding:13px;border:1px solid #e9ebf3;border-radius:14px;background:#fafbfe;margin-top:10px}.form-section:first-of-type{margin-top:0}.form-section-title{display:flex;align-items:center;justify-content:space-between;font-size:12px;font-weight:900;color:#39445d;margin-bottom:5px}.form-row{display:grid;grid-template-columns:1fr 1fr;gap:9px}.hint{font-size:11px;color:var(--muted);line-height:1.45;margin-top:6px}.currency-note{padding:10px 11px;border-radius:11px;background:#eef3ff;color:#2f4da7;font-size:11px;font-weight:700;line-height:1.45;margin-top:9px}details.advanced{border:1px solid #e4e7f0;border-radius:13px;padding:10px 12px;margin-top:10px;background:#fff}details.advanced summary{cursor:pointer;font-weight:900;font-size:12px;color:#42506d}.danger{background:var(--danger);padding:7px 10px}.ok{background:var(--ok)}.secondary{background:#eef1f8;color:#34415d}.violet{background:#6b4ce6}.outline{background:#fff;color:#3454b4;border:1px solid #cad4f4}.outline.red{color:#d43d58;border-color:#f0c5cd}.outline.green{color:#14895a;border-color:#b9e4d1}
.msg{min-height:18px;margin-top:8px;font-size:12px;color:#15803d}.badge{display:inline-flex;align-items:center;padding:5px 9px;border-radius:999px;background:#e9edff;color:#4a45ad;font-size:10px;font-weight:900}.badge.new{background:#eef2ff;color:#4f46e5}.badge.progressing{background:#fff4df;color:#b36a08}.badge.ready{background:#e8f8f0;color:#168358}.badge.done{background:#e9f8f5;color:#087c68}.badge.waiting{background:#fff0f2;color:#bd3650}.money{font-weight:800;color:#168358}.minus{font-weight:800;color:#be334f}.balance{font-weight:900;color:#d23f59}.stage{display:flex;align-items:center;gap:7px;padding:7px 0}.stage input{width:auto;margin:0}.low{background:#fff0f2!important}
.tablewrap{overflow:auto;max-height:680px;padding:0}.table-head{position:sticky;top:0;z-index:2;display:flex;align-items:center;justify-content:space-between;gap:10px;padding:15px 16px 12px;background:#fff;border-bottom:1px solid var(--line)}.table-tools{display:flex;gap:7px;align-items:center}.table-tools input,.table-tools select{width:auto;margin:0;padding:8px 10px;font-size:11px}.table-tools input{min-width:210px}table{width:100%;border-collapse:separate;border-spacing:0;font-size:11px}th,td{padding:11px 10px;border-bottom:1px solid #eceef4;text-align:left;vertical-align:middle;white-space:nowrap}th{position:sticky;top:64px;z-index:1;background:#f8f9fc;color:#626c82;font-size:10px;text-transform:uppercase;letter-spacing:.35px}tbody tr:hover{background:#fafbfe}.order-code{font-size:13px;font-weight:900;color:#23345f}.order-sub{display:block;color:var(--muted);font-size:10px;margin-top:3px}.progress-track{height:7px;min-width:115px;background:#e7eaf2;border-radius:999px;overflow:hidden;margin-bottom:5px}.progress-fill{height:100%;background:linear-gradient(90deg,#5d55e9,#8456e6);border-radius:999px}.action-group{display:flex;gap:5px;flex-wrap:wrap;max-width:210px}.action-group button{padding:6px 8px;font-size:10px;border-radius:8px}.doc-group{display:flex;gap:4px;flex-wrap:wrap;min-width:195px}.doc-group button{padding:6px 7px;font-size:9px;border-radius:7px}.price-main{font-weight:900;font-size:12px}.price-sub{display:block;color:var(--muted);font-size:9px;margin-top:3px}.empty-row{text-align:center!important;color:var(--muted);padding:30px!important}
.modal-shell{display:none;position:fixed;inset:0;background:rgba(17,24,54,.64);z-index:30;place-items:center;padding:14px;backdrop-filter:blur(4px)}.modal-card{width:min(580px,97vw);max-height:94vh;overflow:auto;background:#fff;border-radius:20px;padding:18px;box-shadow:0 30px 80px rgba(8,13,42,.35)}.modal-title{display:flex;align-items:center;justify-content:space-between;gap:12px}.modal-title h3{margin:0}.payment-preview{display:grid;grid-template-columns:repeat(2,1fr);gap:8px;margin-top:12px}.payment-preview div{padding:12px;border-radius:12px;background:#f5f7fc}.payment-preview small{display:block;color:var(--muted);font-size:10px;margin-bottom:4px}.payment-preview b{font-size:14px}.close-btn{width:auto!important;margin:0!important;background:#edf0f7;color:#34415d}
@media(max-width:1320px){.top{grid-template-columns:1fr auto}.header-actions{grid-column:1/-1;justify-content:flex-start}.cards{grid-template-columns:repeat(3,1fr)}.grid,#orders .grid{grid-template-columns:1fr}.tablewrap{max-height:none}}
@media(max-width:720px){header{padding:14px}.top{display:flex;flex-direction:column;align-items:stretch}.brand-logo{width:80px;height:58px}.brand-text h1{font-size:23px}.live-clock{width:100%}.header-actions{justify-content:flex-start}.wrap{padding:11px}.cards{grid-template-columns:repeat(2,1fr)}.card{padding-left:58px}.card:before{left:12px}.form-row{grid-template-columns:1fr}.table-head{align-items:stretch;flex-direction:column}.table-tools{width:100%;display:grid;grid-template-columns:1fr}.table-tools input,.table-tools select{width:100%;min-width:0}th{top:111px}.payment-preview{grid-template-columns:1fr}.modal-card{padding:14px}}
</style>
</head>
<body>
<header><div class="top">
<div class="brand">
  <img class="brand-logo" src="/static/mebel360-logo.png?v=20260722" alt="Mebel360° logosi">
  <div class="brand-text">
    <h1>Mebel360°</h1>
    <div class="sub">Ishchilar, buyurtmalar, ombor, ishlab chiqarish va moliyani yagona tizimda boshqaring</div>
  </div>
</div>
<div class="live-clock">
  <div id="pharmClock" class="live-clock-time">00:00:00</div>
  <div id="pharmDate" class="live-clock-date">Toshkent vaqti</div>
</div>
<div class="header-actions">
  <a href="/pro-boshqaruv"><button style="background:#0f766e">PRO boshqaruv</button></a>
  <a href="/xodim-akkauntlari"><button style="background:#2563eb">Menejer / Konstruktor</button></a>
  <a href="/shofyor-boshqaruv"><button style="background:#7c3aed">Shofyor boshqaruvi</button></a>
  <a href="/shartnoma-namuna"><button style="background:#0f766e">Shartnoma Word</button></a>
  <a href="/shartnoma-pdf" target="_blank"><button style="background:#0369a1">Shartnoma PDF</button></a>
  <a href="/backup"><button style="background:#16a34a">Backup</button></a>
  <a href="/logout"><button style="background:#dc2626">Chiqish</button></a>
</div>
</div></header>
<main class="wrap">
<div class="cards">
 <div class="card"><span>ISHCHILAR</span><b id="dWorkers">0</b></div>
 <div class="card"><span>FAOL BUYURTMALAR</span><b id="dOrders">0</b></div>
 <div class="card"><span>BUGUNGI SOAT</span><b id="dHours">0</b></div>
 <div class="card"><span>BUGUNGI ISH</span><b id="dProduction">0</b></div>
 <div class="card"><span>BUGUNGI KM</span><b id="dKm">0</b></div>
 <div class="card"><span>KAM QOLDIQ</span><b id="dLow">0</b></div>
</div>

<div class="tabs">
 <button data-tab="workers">Ishchilar</button>
 <button data-tab="attendance">Keldi-ketdi</button>
 <button data-tab="results">Ish natijasi</button>
 <button class="active" data-tab="orders">Buyurtmalar</button>
 <button data-tab="stock">Ombor</button>
 <button data-tab="trips">Shofyor</button>
 <button data-tab="payments">To‘lov</button>
 <button data-tab="penalties">Jarima</button>
 <button data-tab="totals">Jami/Reyting</button>
 <button data-tab="expenses">Xarajat</button>
 <button data-tab="bonuses">Bonus/Ta’til</button>
 <button data-tab="finished">Tayyor mahsulot</button>
 <button data-tab="finance">Foyda</button>
 <button data-tab="customerPro">Mijoz PRO</button>
 <button data-tab="service">Servis/Kafolat</button>
 <button data-tab="delivery">Yetkazish</button>
</div>

<section id="workers" class="tab"><div class="grid">
<div class="panel"><h3>Yangi ishchi</h3><form id="workerForm">
<label>Ism<input name="ism" required></label><label>Familiya<input name="familiya"></label>
<label>Telefon<input name="telefon"></label><label>Lavozim<input name="lavozim"></label>
<label>Ishga kirgan sana<input type="date" name="ishga_kirgan_sana"></label>
<label>Staj (yil)<input type="number" step="0.1" name="staj_yil" value="0"></label>
<label>Kunlik stavka<input type="number" name="kunlik_stavka" value="0"></label>
<label>Oylik maosh<input type="number" name="oylik_maosh" value="0"></label>
<label>Sifat balli (1-10)<input type="number" min="1" max="10" name="sifat_ball" value="5"></label>
<label>Tezlik balli (1-10)<input type="number" min="1" max="10" name="tezlik_ball" value="5"></label>
<label>Intizom balli (1-10)<input type="number" min="1" max="10" name="intizom_ball" value="5"></label>
<hr><h3>Maxfiy hujjatlar</h3>
<p style="font-size:12px;color:#64748b">Bu ma’lumotlarni faqat rahbar ko‘radi.</p>
<label>Pasport/ID seriya-raqami<input name="pasport"></label>
<label>JSHSHIR<input name="jshshir"></label>
<label>Tug‘ilgan sana<input type="date" name="tugilgan_sana"></label>
<label>Yashash manzili<textarea name="yashash_manzil"></textarea></label>
<label>Pasport berilgan sana<input type="date" name="pasport_berilgan_sana"></label>
<label>Pasportni bergan tashkilot<input name="pasport_bergan"></label>
<label>Favqulodda aloqa telefoni<input name="favqulodda_telefon"></label>
<label>Izoh<textarea name="izoh"></textarea></label><button>Saqlash</button><div class="msg"></div>
</form></div><div class="panel tablewrap"><h3>Ishchilar</h3><table><thead><tr><th>Ism</th><th>Lavozim</th><th>Staj</th><th>Kunlik</th><th>Oylik</th><th></th></tr></thead><tbody id="workersBody"></tbody></table></div>
</div></section>

<section id="attendance" class="tab"><div class="grid"><div class="panel"><h3>Avtomatik keldi-ketdi</h3>
<label>Ishchi<select class="workerSelect" id="autoWorkerSelect" required></select></label>
<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px">
<button type="button" class="ok" onclick="clockIn()">Hozir keldi</button>
<button type="button" style="background:#dc2626" onclick="clockOut()">Hozir ketdi</button>
</div><div id="attendanceAutoMsg" class="msg"></div>
<hr><h3>Qo‘lda kiritish</h3><form id="attendanceForm">
<label>Ishchi<select class="workerSelect" name="ishchi_id" required></select></label><label>Sana<input type="date" name="sana" required></label>
<label>Keldi<input type="time" name="keldi_vaqti" required></label><label>Ketdi<input type="time" name="ketdi_vaqti" required></label>
<button>Saqlash</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Ishchi</th><th>Sana</th><th>Keldi</th><th>Ketdi</th><th>Soat</th></tr></thead><tbody id="attendanceBody"></tbody></table></div></div></section>

<section id="results" class="tab"><div class="grid"><div class="panel"><h3>Ish natijasi</h3><form id="resultForm">
<label>Ishchi<select class="workerSelect" name="ishchi_id" required></select></label>
<label>Ish turi<select id="workTypeSelect" name="ish_turi_id" required></select></label>
<label>Sana<input type="date" name="sana" required></label><label>Miqdor<input type="number" step="0.1" name="miqdor" required></label>
<label>Birlik narxi<input type="number" name="birlik_narxi" value="0"></label>
<label>Buyurtma kodi<input name="buyurtma_kodi"></label><label>Izoh<textarea name="izoh"></textarea></label>
<button>Saqlash</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Ishchi</th><th>Ish turi</th><th>Sana</th><th>Miqdor</th><th>Haq</th><th>Buyurtma</th></tr></thead><tbody id="resultsBody"></tbody></table></div></div></section>

<section id="orders" class="tab active"><div class="grid">
<div class="panel"><h3>➕ Yangi buyurtma</h3><form id="orderForm">
<div class="form-section"><div class="form-section-title"><span>1. Mijoz va buyurtma</span></div>
<div class="form-row"><label>Kod<input name="kod" required placeholder="AB-001"></label><label>Mijoz<input name="mijoz" required placeholder="Mijoz ismi"></label></div>
<div class="form-row"><label>Telefon<input name="telefon" placeholder="+998 90 123 45 67"></label><label>Pasport/ID<input name="pasport_id"></label></div>
<label>Manzil<input name="manzil" placeholder="Yetkazish manzili"></label><label>Mahsulot<input name="mahsulot" placeholder="Masalan: dorixona vitrinalari"></label>
<label>Buyurtma turi<select name="buyurtma_turi" id="orderTypeSelect"><option>To‘liq mebel</option><option>Faqat oyna</option><option>MDF fasad</option><option>Ta’mirlash</option><option>Aralash buyurtma</option><option>Maxsus</option></select></label>
<label id="customStagesLabel" style="display:none">Maxsus bosqichlar<textarea name="maxsus_bosqichlar" placeholder="O‘lchov → Chizma → Ishlov → Sifat nazorati → Yakunlandi"></textarea></label></div>

<div class="form-section"><div class="form-section-title"><span>2. Narx va valyuta</span><span>UZS / USD</span></div>
<div class="form-row"><label>Kelishuv valyutasi<select name="valyuta" id="orderCurrency"><option value="UZS">UZS — so‘m</option><option value="USD">USD — AQSh dollari</option></select></label><label>Umumiy narx<input type="number" step="0.01" name="umumiy_narx" value="0" min="0"></label></div>
<div class="currency-note" id="currencyRuleNote">Hisob-kitob so‘mda yuritiladi.</div>
<div class="form-row"><label>Boshlang‘ich avans<input type="number" step="0.01" name="oldindan_tolov" value="0" min="0"></label><label>Avans qabul qilingan valyuta<select name="boshlangich_tolov_valyutasi" id="initialPaymentCurrency"><option value="UZS">UZS — so‘m</option><option value="USD">USD — dollar</option></select></label></div>
<div class="form-row" id="initialRateRow" style="display:none"><label>1 USD kursi<input type="number" step="0.01" name="boshlangich_kurs" id="initialRate" min="0" placeholder="Masalan: 12900"></label><label>Avans sanasi<input type="date" name="boshlangich_tolov_sana"></label></div>
<label>To‘lov sharti<select name="tolov_sharti" id="paymentConditionSelect"><option>Avans majburiy</option><option>Qisman avans</option><option>Avans talab qilinmaydi — ishonchli mijoz</option><option>Muddatli to‘lov</option><option>Shartnoma asosida to‘lov</option></select></label>
<div class="form-row"><label>Talab qilinadigan avans<input type="number" step="0.01" name="avans_talab" value="0" min="0"></label><label>To‘lov usuli<select name="tolov_usuli"><option>Naqd</option><option>Bank orqali</option><option>Click</option><option>Payme</option><option>Qarz</option></select></label></div>
<div class="hint">USD buyurtmada qoldiq dollar holatida saqlanadi. Mijoz so‘mda to‘lasa, har safar o‘sha kunning kursi kiritiladi.</div></div>

<div class="form-section"><div class="form-section-title"><span>3. Ishlab chiqarish</span></div>
<div class="form-row"><label>O‘lcham<input name="olcham" placeholder="2000×600×2400 mm"></label><label>Soni<input type="number" name="soni" value="1" min="1"></label></div>
<div class="form-row"><label>Material<input name="material" placeholder="MDF / LDSP / Akril"></label><label>Rang<input name="rang" placeholder="Oq marmar / Yashil"></label></div>
<div class="form-row"><label>Ishlab chiqarish muddati (kun)<input type="number" name="ishlab_chiqarish_kun" value="0" min="0"></label><label>Taxminiy tayyor sana<input type="date" name="taxminiy_sana"></label></div>
<div class="form-row"><label>Taxminiy soat<input type="time" name="taxminiy_vaqt" value="18:00"></label><label>Mas’ul xodim<input name="masul_xodim"></label></div></div>

<details class="advanced"><summary>Qo‘shimcha ma’lumotlar</summary>
<div class="form-row"><label>Rang kodi<input name="rang_kodi" placeholder="RAL 6018"></label><label>Yaltiroqlik<select name="yaltiroqlik"><option></option><option>Mat</option><option>Yarim mat</option><option>Yaltiroq</option></select></label></div>
<div class="form-row"><label>Avans muddati<input type="date" name="avans_muddat_sana"></label><label>Narx taklifi amal qilish sanasi<input type="date" name="taklif_amal_sana"></label></div>
<div class="form-row"><label>Chizma versiyasi<input type="number" name="chizma_versiya" value="1" min="1"></label><label>Bepul o‘zgartirish limiti<input type="number" name="bepul_ozgarish_limit" value="2" min="0"></label></div>
<input type="hidden" name="oraliq_tolov" value="0"><div class="hint">Keyingi barcha to‘lovlarni buyurtmalar jadvalidagi “To‘lov” tugmasi orqali kiriting.</div><label>Kafolat muddati<input name="kafolat_muddati" value="12 oy"></label>
<div class="form-row"><label>Yetkazish<select name="yetkazish"><option>Kiritilgan</option><option>Kiritilmagan</option><option>Alohida haq</option></select></label><label>Montaj<select name="montaj"><option>Kiritilgan</option><option>Kiritilmagan</option><option>Alohida haq</option></select></label></div>
<div class="form-row"><label>Boshlanish sana<input type="date" name="boshlanish_sana"></label><label>Tugash sana<input type="date" name="tugash_sana"></label></div>
<label>Mijozga aloqa telefoni<input name="aloqa_telefon" placeholder="+998..."></label><label>Mijozga ko‘rinadigan izoh<textarea name="mijozga_izoh"></textarea></label>
<label>Muddat o‘zgarishi / kechikish sababi<textarea name="kechikish_sababi"></textarea></label>
<div class="form-row"><label>Rasm/video havolasi<input name="media_havola" placeholder="https://..."></label><label>Media turi<select name="media_turi"><option>Rasm</option><option>Video</option><option>Fayl</option></select></label></div>
<label>Media izohi<input name="media_izoh"></label><label>Mijoz tasdiqlashi<select name="tasdiq_turi"><option value="">Kerak emas</option><option>Hammasini tasdiqlash</option><option>Rang tasdig‘i</option><option>Chizma tasdig‘i</option><option>Material tasdig‘i</option><option>Narx tasdig‘i</option><option>Muddat tasdig‘i</option></select></label>
<label>Tasdiq izohi<input name="tasdiq_izoh"></label><div class="form-row"><label>Holat<select name="holat"><option>Yangi</option><option>Jarayonda</option><option>Tayyor</option><option>Yetkazildi</option></select></label><label>Kechikish sababi<select name="kechikish_turi"><option value="">Yo‘q</option><option>Korxona</option><option>Mijoz</option><option>Material</option><option>Favqulodda holat</option></select></label></div>
<div class="form-row"><label>Kechikish chegirmasi (%/kun)<input type="number" step="0.1" name="kechikish_foiz" value="0"></label><label>Maksimal chegirma %<input type="number" step="0.1" name="maks_chegirma_foiz" value="20"></label></div>
<div class="form-row"><label>Keshbek %<input type="number" step="0.1" name="keshbek_foiz" value="0"></label><label>Keshbek summasi<input type="number" name="keshbek_summa" value="0"></label></div>
<div class="form-row"><label>Kafolat boshlanish<input type="date" name="kafolat_boshlanish"></label><label>Kafolat tugash<input type="date" name="kafolat_tugash"></label></div>
<label>Kafolat sharti<textarea name="kafolat_sharti"></textarea></label><label>Lokatsiya havolasi<input name="lokatsiya" placeholder="Google Maps havolasi"></label>
<div class="form-row"><label>Mo‘ljal<input name="moljal"></label><label>Qavat<input name="qavat"></label></div><div class="form-row"><label>Lift<select name="lift"><option></option><option>Bor</option><option>Yo‘q</option></select></label><label>Katta mashina<select name="katta_mashina"><option></option><option>Kira oladi</option><option>Kira olmaydi</option></select></label></div>
<label>Izoh<textarea name="izoh"></textarea></label></details>
<button class="violet">Buyurtmani saqlash</button><div class="msg"></div></form></div>

<div class="panel tablewrap"><div class="table-head"><div><h3 style="margin:0">Buyurtmalar</h3><div class="hint">Narx, jarayon, muddat va hujjatlar bir qatorda</div></div><div class="table-tools"><input id="orderSearch" placeholder="Kod, mijoz yoki mahsulot..."><select id="orderStatusFilter"><option value="">Barcha holatlar</option><option>Yangi</option><option>Jarayonda</option><option>Tayyor</option><option>Yetkazildi</option></select></div></div>
<table><thead><tr><th>Buyurtma</th><th>Sana</th><th>Holat</th><th>Jarayon</th><th>Umumiy</th><th>To‘langan</th><th>Qoldiq</th><th>Muddat</th><th>Mijoz oynasi</th><th>Hujjatlar</th></tr></thead><tbody id="ordersBody"></tbody></table></div>
</div></section>

<section id="stock" class="tab"><div class="grid"><div class="panel"><h3>Ombor harakati</h3><form id="stockForm">
<label>Material<select id="stockSelect" name="material_id" required></select></label><label>Sana<input type="date" name="sana" required></label>
<label>Turi<select name="turi"><option>Kirim</option><option>Chiqim</option></select></label><label>Miqdor<input type="number" step="0.1" name="miqdor" required></label>
<label>Buyurtma kodi<input name="buyurtma_kodi"></label><label>Izoh<input name="izoh"></label><button>Saqlash</button><div class="msg"></div></form></div>
<div class="panel tablewrap"><h3>Ombor qoldig‘i</h3><table><thead><tr><th>Material</th><th>Kategoriya</th><th>Qoldiq</th><th>Birlik</th><th>Min.</th></tr></thead><tbody id="stockBody"></tbody></table></div></div></section>

<section id="trips" class="tab"><div class="grid"><div class="panel"><h3>Shofyor safari</h3><form id="tripForm">
<label>Shofyor<select class="workerSelect" name="ishchi_id" required></select></label><label>Sana<input type="date" name="sana" required></label>
<label>Mashina<input name="mashina"></label><label>Qayerdan<input name="qayerdan"></label><label>Qayerga<input name="qayerga"></label>
<label>Masofa km<input type="number" step="0.1" name="masofa_km"></label><label>Sabab<input name="sabab"></label>
<label>Yoqilg‘i litr<input type="number" step="0.1" name="yonilgi"></label><label>Xarajat<input type="number" name="xarajat"></label>
<button>Saqlash</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Shofyor</th><th>Sana</th><th>Mashina</th><th>Yo‘nalish</th><th>Km</th><th>Sabab</th><th>Xarajat</th></tr></thead><tbody id="tripsBody"></tbody></table></div></div></section>

<section id="payments" class="tab"><div class="grid"><div class="panel"><h3>To‘lov</h3><form id="paymentForm">
<label>Ishchi<select class="workerSelect" name="ishchi_id" required></select></label><label>Sana<input type="date" name="sana" required></label>
<label>Summa<input type="number" name="miqdor" required></label><label>Turi<select name="turi"><option>Avans</option><option>Oylik</option><option>Bonus</option></select></label>
<label>Izoh<input name="tavsifi"></label><button>Saqlash</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Ishchi</th><th>Sana</th><th>Summa</th><th>Turi</th></tr></thead><tbody id="paymentsBody"></tbody></table></div></div></section>

<section id="penalties" class="tab"><div class="grid"><div class="panel"><h3>Jarima</h3><form id="penaltyForm">
<label>Ishchi<select class="workerSelect" name="ishchi_id" required></select></label><label>Sana<input type="date" name="sana" required></label>
<label>Summa<input type="number" name="miqdor" required></label><label>Sababi<input name="sababi"></label><button>Saqlash</button><div class="msg"></div></form></div>
<div class="panel tablewrap"><table><thead><tr><th>Ishchi</th><th>Sana</th><th>Summa</th><th>Sabab</th></tr></thead><tbody id="penaltiesBody"></tbody></table></div></div></section>

<section id="totals" class="tab"><div class="panel">
<div style="display:flex;gap:8px;align-items:end;flex-wrap:wrap;margin-bottom:12px">
<label style="margin:0">Boshlanish<input id="totalStart" type="date"></label><label style="margin:0">Tugash<input id="totalEnd" type="date"></label>
<button onclick="loadTotals()">Hisoblash</button><button onclick="setMonth()" style="background:#475569">Shu oy</button>
<a id="csvLink"><button style="background:#16a34a">Excel/CSV</button></a></div>
<div class="tablewrap"><table><thead><tr><th>O‘rin</th><th>Ishchi</th><th>Lavozim</th><th>Kun</th><th>Soat</th><th>Miqdor</th><th>Ish haqi</th><th>Jarima</th><th>Bonus</th><th>To‘langan</th><th>Qoldiq</th><th>Reyting</th></tr></thead><tbody id="totalsBody"></tbody></table></div>
</div></section>

<section id="expenses" class="tab">
<div class="grid">
<div class="panel"><h3>Yangi xarajat</h3><form id="expenseForm">
<label>Sana<input type="date" name="sana" required></label>
<label>Kategoriya<select name="kategoriya">
<option>Material</option><option>Furnitura</option><option>Ishlab chiqarish</option>
<option>Ishchi</option><option>Transport</option><option>Korxona</option>
<option>Ofis va reklama</option><option>Mijoz va buyurtma</option>
<option>Uy</option><option>Akam</option><option>Ota uchun</option><option>Ona uchun</option>
<option>Favqulodda</option><option>Boshqa</option>
</select></label>
<label>Xarajat nomi<select name="xarajat_nomi">
<option>MDF</option><option>DSP</option><option>LDSP</option><option>HDF</option><option>DVP</option>
<option>Akril</option><option>Fanera</option><option>Oyna</option><option>Shisha</option>
<option>Kromka</option><option>Profil</option><option>Bo‘yoq</option><option>Lak</option>
<option>Grunt</option><option>Shpaklyovka</option><option>Yelim</option><option>Silikon</option>
<option>Porolon</option><option>Mato</option><option>Ekokoja</option><option>Qadoqlash materiali</option>
<option>Petlya</option><option>Ruchka</option><option>Napravlyayushiy</option><option>Gazlift</option>
<option>Oyoq</option><option>G‘ildirak</option><option>Zamok</option><option>Magnit</option>
<option>Samorez</option><option>Bolt va gayka</option><option>Tortma mexanizmi</option>
<option>Ko‘tarma mexanizm</option><option>LED lenta</option><option>Blok pitaniya</option>
<option>Sim va elektr jihozlari</option><option>Raspil</option><option>Kromka urish</option>
<option>CNC Rover</option><option>Frezalash</option><option>Lazer</option><option>Teshish</option>
<option>Bo‘yash</option><option>Sayqalash</option><option>Oyna teshish</option>
<option>Oyna kesish</option><option>Payvandlash</option><option>Tashqi ustaga berilgan ish</option>
<option>Qayta ishlash xarajati</option><option>Yaroqsiz mahsulot</option><option>Material chiqindisi</option>
<option>Kunlik ish haqi</option><option>Oylik maosh</option><option>Soatbay haq</option>
<option>Ish hajmiga haq</option><option>Avans</option><option>Bonus</option>
<option>Qo‘shimcha ish haqi</option><option>Ortib ishlagan vaqt</option>
<option>Ovqat puli</option><option>Yo‘l puli</option><option>Telefon puli</option>
<option>Xizmat safari</option><option>Ishchi kiyimi</option><option>Himoya vositalari</option>
<option>Benzin</option><option>Gaz</option><option>Dizel</option><option>Moy almashtirish</option>
<option>Mashina ta’miri</option><option>Ehtiyot qismlar</option><option>Shina</option>
<option>Yuvish</option><option>Jarima</option><option>Yo‘l haqi</option><option>Parking</option>
<option>Yuk tashish</option><option>Yetkazib berish</option><option>Bozorga borish</option>
<option>Raspildan olib kelish</option><option>Mijoz uyiga borish</option>
<option>Elektr</option><option>Gaz</option><option>Suv</option><option>Ijara</option>
<option>Internet</option><option>Telefon</option><option>Qo‘riqlash</option><option>Tozalash</option>
<option>Chiqindi olib ketish</option><option>Stanok ta’miri</option><option>Asbob-uskunalar</option>
<option>Freza</option><option>Disk</option><option>Sverlo</option><option>Zımpara</option>
<option>Kompressor xarajati</option><option>Moylash vositalari</option><option>Texnika xavfsizligi</option>
<option>Kantselyariya</option><option>Printer qog‘ozi</option><option>Kartrij</option>
<option>Kompyuter ta’miri</option><option>Dastur xarajati</option><option>Hosting</option>
<option>Domen</option><option>SMS xizmati</option><option>Reklama</option>
<option>Instagram reklama</option><option>Telegram reklama</option><option>OLX reklama</option>
<option>Foto-video</option><option>Dizayn</option><option>Bank xizmati</option>
<option>Soliq</option><option>Buxgalteriya</option><option>Razmer olish</option>
<option>Chizma tayyorlash</option><option>Namuna tayyorlash</option><option>Montaj</option>
<option>Qavatga ko‘tarish</option><option>Servis</option><option>Kafolat ta’miri</option>
<option>Mijozga chegirma</option><option>Keshbek</option><option>Qaytarilgan pul</option>
<option>Kechikish kompensatsiyasi</option><option>Qo‘shimcha ish</option>
<option>Uy uchun</option><option>Akam</option><option>Ota uchun</option><option>Ona uchun</option>
<option>Korxona uchun</option><option>Qarzdorlik yopish</option><option>Favqulodda xarajat</option>
<option>Boshqa xarajat</option>
</select></label>
<label>Summa<input type="number" name="miqdor" required></label>
<label>To‘lov usuli<select name="tolov_usuli">
<option>Naqd</option><option>Click</option><option>Payme</option>
<option>Bank orqali</option><option>Plastik karta</option><option>Qarz</option>
</select></label>
<label>Buyurtma kodi<input name="buyurtma_kodi" placeholder="AB 007"></label>
<label>Kimga berildi<input name="kimga_berildi"></label>
<label>Chek rasmi yoki havolasi<input name="chek_havola"></label>
<label>Izoh<input name="tavsifi"></label>
<button>Saqlash</button><div class="msg"></div>
</form></div>
<div class="panel tablewrap"><table><thead><tr>
<th>Sana</th><th>Kategoriya</th><th>Xarajat nomi</th><th>Summa</th>
<th>To‘lov usuli</th><th>Buyurtma</th><th>Kimga</th><th>Izoh</th>
</tr></thead><tbody id="expensesBody"></tbody></table></div>
</div></section>

<section id="bonuses" class="tab"><div class="grid"><div class="panel"><h3>Bonus</h3><form id="bonusForm"><label>Ishchi<select class="workerSelect" name="ishchi_id" required></select></label><label>Sana<input type="date" name="sana" required></label><label>Summa<input type="number" name="miqdor" required></label><label>Sabab<input name="sababi"></label><button>Bonus saqlash</button><div class="msg"></div></form><hr><h3>Ta’til / holat</h3><form id="statusForm"><label>Ishchi<select class="workerSelect" name="ishchi_id" required></select></label><label>Sana<input type="date" name="sana" required></label><label>Turi<select name="turi"><option>Dam olish</option><option>Ta’til</option><option>Kasallik</option><option>Sababsiz</option></select></label><label>Izoh<input name="izoh"></label><button>Saqlash</button><div class="msg"></div></form></div>
<div class="panel tablewrap"><h3>Bonuslar</h3><table><thead><tr><th>Ishchi</th><th>Sana</th><th>Summa</th><th>Sabab</th></tr></thead><tbody id="bonusesBody"></tbody></table><h3 style="margin-top:20px">Ta’til va holatlar</h3><table><thead><tr><th>Ishchi</th><th>Sana</th><th>Turi</th><th>Izoh</th></tr></thead><tbody id="statusesBody"></tbody></table></div></div></section>

<section id="finished" class="tab"><div class="grid"><div class="panel"><h3>Tayyor mahsulot</h3><form id="finishedForm"><label>Nomi<input name="nomi" required></label><label>Kodi<input name="kodi"></label><label>Rang<input name="rang"></label><label>Miqdor<input type="number" step="0.1" name="miqdor" required></label><label>Birlik<select name="birlik"><option>dona</option><option>komplekt</option></select></label><label>Narx<input type="number" name="narx"></label><label>Izoh<input name="izoh"></label><button>Saqlash</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Nomi</th><th>Kod</th><th>Rang</th><th>Miqdor</th><th>Narx</th></tr></thead><tbody id="finishedBody"></tbody></table></div></div></section>

<section id="finance" class="tab"><div class="panel"><h3>Moliyaviy xulosa</h3><div style="display:flex;gap:8px;flex-wrap:wrap;align-items:end"><label>Boshlanish<input id="finStart" type="date"></label><label>Tugash<input id="finEnd" type="date"></label><button onclick="loadFinance()">Hisoblash</button></div><div class="cards" style="margin-top:15px"><div class="card"><span>KIRIM</span><b id="fIncome">0</b></div><div class="card"><span>XARAJAT</span><b id="fExpense">0</b></div><div class="card"><span>ISHCHI TO‘LOV</span><b id="fSalary">0</b></div><div class="card"><span>BONUS</span><b id="fBonus">0</b></div><div class="card"><span>SOF FOYDA</span><b id="fProfit">0</b></div></div></div></section>

<section id="customerPro" class="tab"><div class="grid"><div class="panel"><h3>Mijoz bahosi</h3><form id="ratingForm"><label>Buyurtma<select class="orderSelect" name="buyurtma_id" required></select></label><label>Mebel sifati (1-5)<input type="number" min="1" max="5" name="sifat" value="5"></label><label>Muddat (1-5)<input type="number" min="1" max="5" name="muddat" value="5"></label><label>Muomala (1-5)<input type="number" min="1" max="5" name="muomala" value="5"></label><label>Yetkazish (1-5)<input type="number" min="1" max="5" name="yetkazish" value="5"></label><label>Montaj (1-5)<input type="number" min="1" max="5" name="montaj" value="5"></label><label>Izoh<textarea name="izoh"></textarea></label><button>Bahoni saqlash</button><div class="msg"></div></form><hr><h3>Qo‘shimcha ish</h3><form id="extraForm"><label>Buyurtma<select class="orderSelect" name="buyurtma_id" required></select></label><label>Ish nomi<input name="nomi" required></label><label>Qo‘shimcha summa<input type="number" name="summa"></label><label>Qo‘shimcha kun<input type="number" name="qoshimcha_kun"></label><label>Izoh<input name="izoh"></label><button>Qo‘shish</button><div class="msg"></div></form></div><div class="panel tablewrap"><h3>Qo‘shimcha ishlar</h3><table><thead><tr><th>Buyurtma</th><th>Mijoz</th><th>Ish</th><th>Summa</th><th>Kun</th></tr></thead><tbody id="extrasBody"></tbody></table></div></div></section>
<section id="service" class="tab"><div class="grid"><div class="panel"><h3>Servis / kafolat murojaati</h3><form id="serviceForm"><label>Buyurtma<select class="orderSelect" name="buyurtma_id" required></select></label><label>Turi<select name="turi"><option>Kafolat</option><option>Servis</option><option>Pullik xizmat</option></select></label><label>Muammo<textarea name="muammo" required></textarea></label><label>Servis sanasi<input type="date" name="servis_sana"></label><label>Usta<input name="usta"></label><button>Qabul qilish</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Buyurtma</th><th>Mijoz</th><th>Turi</th><th>Muammo</th><th>Holat</th><th>Sana</th><th>Usta</th></tr></thead><tbody id="serviceBody"></tbody></table></div></div></section>
<section id="delivery" class="tab"><div class="grid"><div class="panel"><h3>Yetkazishni rejalash</h3><form id="deliveryForm"><label>Buyurtma<select class="orderSelect" name="buyurtma_id" required></select></label><label>Haydovchi<select class="workerSelect" name="haydovchi_id"></select></label><label>Sana<input type="date" name="sana" required></label><label>Navbat<input type="number" name="navbat" value="1"></label><label>Mashina<input name="mashina"></label><label>Benzin<input type="number" name="benzin"></label><label>Yo‘l xarajati<input type="number" name="yol_xarajati"></label><label>Izoh<input name="izoh"></label><button>Rejalash</button><div class="msg"></div></form></div><div class="panel tablewrap"><table><thead><tr><th>Sana</th><th>Navbat</th><th>Buyurtma</th><th>Mijoz</th><th>Haydovchi</th><th>Manzil</th><th>Holat</th><th>Amal</th></tr></thead><tbody id="deliveryBody"></tbody></table></div></div></section>

</main>

<div id="stageModal" style="display:none;position:fixed;inset:0;background:#0009;z-index:20;place-items:center;padding:12px">
<div class="panel" style="width:min(620px,97%);max-height:90vh;overflow:auto"><h3>Buyurtma bosqichlari</h3><div id="stageWorkflowInfo"></div><div id="stageList"></div><button onclick="closeStage()" style="margin-top:10px">Yopish</button></div>
</div>

<div id="paymentModal" class="modal-shell">
<div class="modal-card"><div class="modal-title"><div><h3>💳 Buyurtma to‘lovi</h3><div id="paymentOrderInfo" class="hint"></div></div><button type="button" class="close-btn" onclick="closePaymentModal()">Yopish</button></div>
<form id="orderPaymentForm"><input type="hidden" name="order_id" id="paymentOrderId"><input type="hidden" id="paymentOrderCurrency">
<div class="form-row"><label>To‘lov sanasi<input type="date" name="sana" id="paymentDate"></label><label>Qabul qilingan valyuta<select name="tolov_valyutasi" id="paymentCurrency"><option value="UZS">UZS — so‘m</option><option value="USD">USD — dollar</option></select></label></div>
<label>Qabul qilingan summa<input type="number" step="0.01" min="0" name="qabul_qilingan_summa" id="paymentReceived" required></label>
<label id="paymentRateLabel">1 USD kursi<input type="number" step="0.01" min="0" name="kurs" id="paymentRate" placeholder="Masalan: 12900"></label>
<div class="form-row"><label>To‘lov turi<select name="turi"><option>To‘lov</option><option>Avans</option><option>Oraliq to‘lov</option><option>Yakuniy to‘lov</option></select></label><label>Izoh<input name="izoh" placeholder="Naqd / bank orqali..."></label></div>
<div class="payment-preview"><div><small>Buyurtmadan ayriladi</small><b id="paymentOrderAmount">0</b></div><div><small>So‘mdagi ekvivalent</small><b id="paymentUzsAmount">0 so‘m</b></div></div>
<div class="currency-note">USD qatnashsa, kurs aynan shu to‘lov bilan saqlanadi va keyinchalik o‘zgarmaydi.</div>
<button type="submit" class="ok">To‘lovni saqlash</button><div class="msg"></div></form></div></div>

<div id="customerCardModal" style="display:none;position:fixed;inset:0;background:#0009;z-index:25;place-items:center;padding:12px">
<div class="panel" style="width:min(720px,97%);max-height:92vh;overflow:auto">
<div style="display:flex;justify-content:space-between;gap:10px;align-items:center"><h3>👤 Mijozga ko‘rinadigan karta</h3><button type="button" class="danger" onclick="closeCustomerCard()">Yopish</button></div>
<form id="customerCardForm"><input type="hidden" name="order_id" id="ccOrderId">
<div style="display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:8px">
<label>Taxminiy tayyor sana<input type="date" name="taxminiy_sana"></label><label>Taxminiy tayyor soat<input type="time" name="taxminiy_vaqt" value="18:00"></label>
<label>Rang<input name="rang" placeholder="Oq marmar"></label><label>Rang kodi<input name="rang_kodi" placeholder="RAL 6018 yoki #22aa66"></label>
<label>Material<input name="material"></label><label>Yaltiroqlik<select name="yaltiroqlik"><option></option><option>Mat</option><option>Yarim mat</option><option>Yaltiroq</option></select></label>
<label>Mas’ul xodim<input name="masul_xodim"></label><label>Aloqa telefoni<input name="aloqa_telefon" placeholder="+998..."></label>
<label>Talab qilinadigan avans<input type="number" name="avans_talab" min="0"></label><label>Ishlab chiqarish muddati (kun)<input type="number" name="ishlab_chiqarish_kun" min="0"></label>
<label>Avans muddati<input type="date" name="avans_muddat_sana"></label><label>Narx taklifi amal qilish sanasi<input type="date" name="taklif_amal_sana"></label>
<label>Chizma versiyasi<input type="number" name="chizma_versiya" min="1"></label><label>Bepul o‘zgartirish limiti<input type="number" name="bepul_ozgarish_limit" min="0"></label>
</div>
<label>Mijozga ko‘rinadigan izoh<textarea name="mijozga_izoh"></textarea></label>
<label>Muddat o‘zgarishi / kechikish sababi<textarea name="kechikish_sababi"></textarea></label>
<label>Muddatni o‘zgartirish sababi<input name="muddat_ozgarish_sababi" placeholder="Muddat o‘zgarsa tarixga yoziladi"></label>
<hr><h3>📷 Rasm yoki video qo‘shish</h3><div style="display:grid;grid-template-columns:1fr 1fr;gap:8px"><label>Havola<input name="media_havola" placeholder="https://..."></label><label>Turi<select name="media_turi"><option>Rasm</option><option>Video</option><option>Fayl</option></select></label></div><label>Izoh<input name="media_izoh" placeholder="Bo‘yash yakunlandi"></label>
<hr><h3>✅ Mijoz tasdiqlashi</h3><div style="display:grid;grid-template-columns:1fr 1fr;gap:8px"><label>Turi<select name="tasdiq_turi"><option value="">Yangi so‘rov yo‘q</option><option>Hammasini tasdiqlash</option><option>Rang tasdig‘i</option><option>Chizma tasdig‘i</option><option>Material tasdig‘i</option><option>Narx tasdig‘i</option><option>Muddat tasdig‘i</option></select></label><label>Izoh<input name="tasdiq_izoh" placeholder="Tasdiqlaysizmi?"></label></div>
<button type="submit" class="ok">Saqlash</button><div class="msg"></div></form>
<div id="customerCardInfo" style="margin-top:14px"></div>
</div></div>

<script>
const $=s=>document.querySelector(s),$$=s=>document.querySelectorAll(s);
function setupOrderWorkflowForm(){const t=$('#orderTypeSelect'),c=$('#customStagesLabel');if(!t||!c)return;const sync=()=>{c.style.display=t.value==='Maxsus'?'block':'none'};t.addEventListener('change',sync);sync()}
function setupOrderCurrencyForm(){
 const oc=$('#orderCurrency'),pc=$('#initialPaymentCurrency'),row=$('#initialRateRow'),note=$('#currencyRuleNote');if(!oc||!pc)return;
 const sync=()=>{const usd=oc.value==='USD'||pc.value==='USD';row.style.display=usd?'grid':'none';note.textContent=oc.value==='USD'?'Kelishuv USDda saqlanadi. So‘mda qilingan har bir to‘lov to‘lov kunidagi kurs bo‘yicha dollardan ayriladi.':'Hisob-kitob so‘mda yuritiladi.'};
 oc.addEventListener('change',()=>{pc.value=oc.value;sync()});pc.addEventListener('change',sync);sync();
}
function tashkentDate(){const p=new Intl.DateTimeFormat('en-CA',{timeZone:'Asia/Tashkent',year:'numeric',month:'2-digit',day:'2-digit'}).formatToParts(new Date());const v=Object.fromEntries(p.map(x=>[x.type,x.value]));return `${v.year}-${v.month}-${v.day}`}const today=tashkentDate();$$('input[type=date]').forEach(x=>x.value=today);
$$('.tabs button').forEach(b=>b.onclick=()=>{$$('.tabs button').forEach(x=>x.classList.remove('active'));$$('.tab').forEach(x=>x.classList.remove('active'));b.classList.add('active');$('#'+b.dataset.tab).classList.add('active')});
function esc(s){return String(s===undefined||s===null?'':s).replace(/[&<>"']/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]))}
async function api(u,o){const r=await fetch(u,o),j=await r.json();if(!r.ok)throw new Error(j.message||'Xato');return j}
function fj(f){return Object.fromEntries(new FormData(f).entries())}
function bind(id,url){const f=$(id);f.onsubmit=async e=>{e.preventDefault();const m=f.querySelector('.msg');try{await api(url,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(fj(f))});m.textContent='✅ Saqlandi';await refresh()}catch(x){m.textContent='❌ '+x.message}}}
function money(v){return Number(v||0).toLocaleString('uz-UZ',{maximumFractionDigits:2})}

async function loadWorkers(){const a=await api('/api/ishchilar');$('#workersBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.lavozim||'')}</td><td>${esc(x.staj_yil)} yil</td><td>${money(x.kunlik_stavka)}</td><td>${money(x.oylik_maosh)}</td><td><button class="danger" onclick="delWorker(${esc(x.id)})">O‘chirish</button></td></tr>`).join('');$$('.workerSelect').forEach(s=>s.innerHTML='<option value="">Tanlang</option>'+a.map(x=>`<option value="${esc(x.id)}">${esc(x.ism)} ${esc(x.familiya||'')}</option>`).join(''))}
async function delWorker(i){if(confirm('O‘chirasizmi?')){await api('/api/ishchilar/'+i,{method:'DELETE'});refresh()}}
async function loadTypes(){const a=await api('/api/ish-turlari');$('#workTypeSelect').innerHTML=a.map(x=>`<option value="${esc(x.id)}">${esc(x.kategoriya)} — ${esc(x.nomi)} (${esc(x.birlik)})</option>`).join('')}
async function loadAttendance(){const a=await api('/api/keldi-ketdi');$('#attendanceBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.sana)}</td><td>${esc(x.keldi_vaqti)}</td><td>${esc(x.ketdi_vaqti)}</td><td>${esc(x.ish_soatlari)}</td></tr>`).join('')}
async function loadResults(){const a=await api('/api/natijalar');$('#resultsBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.ish_turi)}</td><td>${esc(x.sana)}</td><td>${esc(x.miqdor)} ${esc(x.birlik)}</td><td>${money(x.jami_haq)}</td><td>${esc(x.buyurtma_kodi||'')}</td></tr>`).join('')}
let orderRows=[];
function amountText(value,currency){const n=Number(value||0);return currency==='USD'?'$'+n.toLocaleString('en-US',{maximumFractionDigits:2}):money(n)+' so‘m'}
function statusClass(status){const s=String(status||'');if(s.includes('Yangi'))return'new';if(s.includes('Jarayon'))return'progressing';if(s.includes('Tayyor'))return'ready';if(s.includes('Yetkaz')||s.includes('Yop'))return'done';return'waiting'}
function renderOrders(){
 const q=($('#orderSearch')?.value||'').toLowerCase().trim(),status=$('#orderStatusFilter')?.value||'';
 const rows=orderRows.filter(x=>(!status||String(x.holat||'').includes(status))&&(!q||`${x.kod} ${x.mijoz} ${x.mahsulot}`.toLowerCase().includes(q)));
 $('#ordersBody').innerHTML=rows.length?rows.map(x=>{const c=x.valyuta||'UZS',rate=Number(x.oxirgi_kurs||0),approx=c==='USD'&&rate>0?`≈ ${money(Number(x.qoldiq||0)*rate)} so‘m`:'';return `<tr>
 <td><span class="order-code">${esc(x.kod)}</span><span class="order-sub">${esc(x.mijoz)} · ${esc(x.mahsulot||'Mebel')}</span></td>
 <td>${esc(String(x.created_at||'').slice(0,10)||'-')}<span class="order-sub">${esc(x.buyurtma_turi||'To‘liq mebel')}</span></td>
 <td><span class="badge ${statusClass(x.holat)}">${esc(x.holat)}</span><span class="order-sub">${esc(x.joriy_bosqich||'')}</span></td>
 <td><div class="progress-track"><div class="progress-fill" style="width:${Math.min(100,Number(x.progress||0))}%"></div></div><b>${esc(x.progress||0)}%</b><span class="order-sub">${esc(x.bosqich_bajarildi||0)}/${esc(x.bosqich_jami||0)} bosqich</span></td>
 <td><span class="price-main">${amountText(x.umumiy_narx,c)}</span><span class="order-sub">${c}</span></td>
 <td><span class="price-main" style="color:#168358">${amountText(x.oldindan_tolov,c)}</span>${c==='USD'&&Number(x.tolangan_uzs||0)>0?`<span class="price-sub">${money(x.tolangan_uzs)} so‘m qabul qilindi</span>`:''}</td>
 <td><span class="price-main balance">${amountText(x.qoldiq,c)}</span><span class="price-sub">${esc(approx||'')}</span></td>
 <td>${esc(x.muddat_matn||'Belgilanmagan')}<span class="order-sub">${esc(x.tolov_sharti||'')}</span></td>
 <td><div class="action-group"><button class="violet" onclick="openStage(${esc(x.id)})">Bosqich</button><button class="ok" onclick="addOrderPayment(${esc(x.id)})">To‘lov</button><button class="outline green" onclick="openCustomerCard(${esc(x.id)})">Sozlash</button><button class="outline" onclick="openTrack(${esc(x.id)})">Ochish</button><button class="outline" onclick="copyTrack(${esc(x.id)})">Link</button></div></td>
 <td><div class="doc-group"><a href="/buyurtma/${esc(x.id)}/shartnoma.docx"><button class="outline">Word</button></a><a href="/buyurtma/${esc(x.id)}/shartnoma.pdf" target="_blank"><button class="outline red">PDF</button></a><button class="outline" onclick="regenContract(${esc(x.id)})">Yangilash</button><a href="/buyurtma/${esc(x.id)}/chek.pdf" target="_blank"><button class="outline green">Chek</button></a><a href="/buyurtma/${esc(x.id)}/qr.png" target="_blank"><button class="ok">QR</button></a></div></td>
 </tr>`}).join(''):'<tr><td colspan="10" class="empty-row">Mos buyurtma topilmadi.</td></tr>';
}
async function loadOrders(){orderRows=await api('/api/buyurtmalar');$$('.orderSelect').forEach(s=>s.innerHTML='<option value="">Tanlang</option>'+orderRows.map(x=>`<option value="${esc(x.id)}">${esc(x.kod)} — ${esc(x.mijoz)}</option>`).join(''));renderOrders()}
let activeStageOrderId=null;
async function openStage(id){activeStageOrderId=id;const data=await api('/api/buyurtma/'+id+'/workflow'),a=data.stages||[],o=data.order||{};const done=a.filter(x=>x.bajarildi).length,foiz=a.length?Math.round(done*1000/a.length)/10:0;const trusted=['Avans talab qilinmaydi — ishonchli mijoz','Muddatli to‘lov','Shartnoma asosida to‘lov'].includes(o.tolov_sharti);$('#stageWorkflowInfo').innerHTML=`<div style="margin-bottom:10px;padding:12px;border-radius:12px;background:#f1f5f9"><b>${esc(o.buyurtma_turi||'To‘liq mebel')}</b><br><small>To‘lov: ${esc(o.tolov_sharti||'Avans majburiy')} · Rasmiy muddat: ${esc(o.deadline_text||'Hali boshlanmagan')}</small>${trusted&&!Number(o.rahbar_tasdiq)?`<br><button class="ok" style="margin-top:10px" onclick="managerApprove(${esc(id)})">✅ Rahbar tasdiqladi</button>`:''}</div>`;$('#stageList').innerHTML=`<div style="margin-bottom:10px;padding:10px;border-radius:10px;background:#eff6ff;color:#1d4ed8;font-weight:800">Jarayon: ${foiz}% — ${done}/${a.length} bosqich</div>`+a.map(x=>`<label class="stage"><input type="checkbox" ${x.bajarildi?'checked':''} onchange="toggleStage(${esc(x.id)},this)"> ${esc(x.bosqich)}</label>`).join('');$('#stageModal').style.display='grid'}
function closeStage(){$('#stageModal').style.display='none';activeStageOrderId=null}
async function managerApprove(id){if(!confirm('Ishonchli mijoz uchun avanssiz ishni boshlashni tasdiqlaysizmi?'))return;try{await api(`/api/buyurtma/${id}/rahbar-tasdiq`,{method:'POST',headers:{'Content-Type':'application/json'},body:'{}'});await loadOrders();await loadProgress();await openStage(id)}catch(e){alert('Tasdiqlanmadi: '+e.message)}}
async function toggleStage(id,box){const old=!box.checked;box.disabled=true;try{const x=await api('/api/buyurtma-bosqich/'+id,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({bajarildi:box.checked})});await loadOrders();await loadProgress();await loadDashboard();await openStage(x.buyurtma_id)}catch(e){box.checked=old;alert('Bosqich saqlanmadi: '+e.message)}finally{box.disabled=false}}
async function loadStock(){const a=await api('/api/ombor');$('#stockSelect').innerHTML=a.map(x=>`<option value="${esc(x.id)}">${esc(x.nomi)} (${esc(x.birlik)})</option>`).join('');$('#stockBody').innerHTML=a.map(x=>`<tr class="${Number(x.qoldiq)<=Number(x.min_qoldiq)?'low':''}"><td>${esc(x.nomi)}</td><td>${esc(x.kategoriya)}</td><td>${esc(x.qoldiq)}</td><td>${esc(x.birlik)}</td><td>${esc(x.min_qoldiq)}</td></tr>`).join('')}
async function loadTrips(){const a=await api('/api/safarlar');$('#tripsBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.sana)}</td><td>${esc(x.mashina||'')}</td><td>${esc(x.qayerdan||'')} → ${esc(x.qayerga||'')}</td><td>${esc(x.masofa_km)}</td><td>${esc(x.sabab||'')}</td><td>${money(x.xarajat)}</td></tr>`).join('')}
async function loadPayments(){const a=await api('/api/tolovlar');$('#paymentsBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.sana)}</td><td>${money(x.miqdor)}</td><td>${esc(x.turi)}</td></tr>`).join('')}
async function loadPenalties(){const a=await api('/api/jarimalar');$('#penaltiesBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.sana)}</td><td class="minus">${money(x.miqdor)}</td><td>${esc(x.sababi||'')}</td></tr>`).join('')}
async function loadDashboard(){const x=await api('/api/dashboard');$('#dWorkers').textContent=x.workers;$('#dOrders').textContent=x.orders;$('#dHours').textContent=x.hours;$('#dProduction').textContent=x.production;$('#dKm').textContent=x.km;$('#dLow').textContent=x.low_stock}
function setMonth(){const d=new Date(),y=d.getFullYear(),m=String(d.getMonth()+1).padStart(2,'0');$('#totalStart').value=`${y}-${m}-01`;$('#totalEnd').value=new Date(y,d.getMonth()+1,0).toISOString().slice(0,10);loadTotals()}
async function loadTotals(){const s=$('#totalStart').value||'1900-01-01',e=$('#totalEnd').value||'2999-12-31';$('#csvLink').href=`/export/jami.csv?start=${s}&end=${e}`;const a=await api(`/api/jami?start=${s}&end=${e}`);$('#totalsBody').innerHTML=a.map((x,i)=>`<tr><td>${i+1}</td><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.lavozim||'')}</td><td>${esc(x.ish_kunlari)}</td><td>${esc(x.jami_soat)}</td><td>${esc(x.jami_miqdor)}</td><td class="money">${money(x.ish_haqi)}</td><td class="minus">${money(x.jarima)}</td><td class="money">${money(x.bonus)}</td><td>${money(x.tolangan)}</td><td class="balance">${money(x.qoldiq)}</td><td>${esc(x.reyting)}</td></tr>`).join('')}

async function loadExpenses(){const a=await api('/api/xarajatlar');$('#expensesBody').innerHTML=a.map(x=>`<tr><td>${esc(x.sana)}</td><td>${esc(x.kategoriya)}</td><td>${esc(x.xarajat_nomi||'')}</td><td class="minus">${money(x.miqdor)}</td><td>${esc(x.tolov_usuli||'Naqd')}</td><td>${esc(x.buyurtma_kodi||'')}</td><td>${esc(x.kimga_berildi||'')}</td><td>${esc(x.tavsifi||'')}</td></tr>`).join('')}
async function clockIn(){const id=$('#autoWorkerSelect').value,m=$('#attendanceAutoMsg');if(!id){m.textContent='Ishchini tanlang';return}try{const x=await api('/api/davomat/keldi',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({ishchi_id:id})});m.textContent=`✅ Keldi: ${esc(x.sana)} ${esc(x.vaqt)}`;loadAttendance();loadDashboard()}catch(e){m.textContent='❌ '+e.message}}
async function clockOut(){const id=$('#autoWorkerSelect').value,m=$('#attendanceAutoMsg');if(!id){m.textContent='Ishchini tanlang';return}try{const x=await api('/api/davomat/ketdi',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({ishchi_id:id})});m.textContent=`✅ Ketdi: ${esc(x.sana)} ${esc(x.vaqt)} — ${esc(x.ish_soatlari)} soat`;loadAttendance();loadDashboard()}catch(e){m.textContent='❌ '+e.message}}
async function loadBonuses(){const a=await api('/api/bonuslar');$('#bonusesBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.sana)}</td><td class="money">${money(x.miqdor)}</td><td>${esc(x.sababi||'')}</td></tr>`).join('')}
async function loadStatuses(){const a=await api('/api/ishchi-holatlari');$('#statusesBody').innerHTML=a.map(x=>`<tr><td>${esc(x.ism)} ${esc(x.familiya||'')}</td><td>${esc(x.sana)}</td><td>${esc(x.turi)}</td><td>${esc(x.izoh||'')}</td></tr>`).join('')}
async function loadFinished(){const a=await api('/api/tayyor-mahsulot');$('#finishedBody').innerHTML=a.map(x=>`<tr><td>${esc(x.nomi)}</td><td>${esc(x.kodi||'')}</td><td>${esc(x.rang||'')}</td><td>${esc(x.miqdor)} ${esc(x.birlik)}</td><td>${money(x.narx)}</td></tr>`).join('')}
async function loadFinance(){const s=$('#finStart').value||'1900-01-01',e=$('#finEnd').value||'2999-12-31',x=await api(`/api/moliyaviy-xulosa?start=${s}&end=${e}`);$('#fIncome').textContent=money(x.kirim);$('#fExpense').textContent=money(x.xarajat);$('#fSalary').textContent=money(x.ishchi_tolov);$('#fBonus').textContent=money(x.bonus);$('#fProfit').textContent=money(x.sof_foyda)}
function paymentPreview(){const oc=$('#paymentOrderCurrency').value,pc=$('#paymentCurrency').value,received=Number($('#paymentReceived').value||0),rate=Number($('#paymentRate').value||0);let orderAmount=0,uzs=0;if(oc===pc)orderAmount=received;else if(oc==='USD'&&pc==='UZS'&&rate>0)orderAmount=received/rate;else if(oc==='UZS'&&pc==='USD'&&rate>0)orderAmount=received*rate;uzs=pc==='UZS'?received:(rate>0?received*rate:0);$('#paymentRateLabel').style.display=(oc==='USD'||pc==='USD')?'block':'none';$('#paymentOrderAmount').textContent=amountText(orderAmount,oc);$('#paymentUzsAmount').textContent=amountText(uzs,'UZS')}
function addOrderPayment(id){const o=orderRows.find(x=>Number(x.id)===Number(id));if(!o)return;$('#paymentOrderId').value=id;$('#paymentOrderCurrency').value=o.valyuta||'UZS';$('#paymentOrderInfo').textContent=`${o.kod} · Qoldiq: ${amountText(o.qoldiq,o.valyuta||'UZS')}`;$('#paymentDate').value=today;$('#paymentCurrency').value=o.valyuta||'UZS';$('#paymentReceived').value='';$('#paymentRate').value=Number(o.oxirgi_kurs||0)>0?o.oxirgi_kurs:'';$('#orderPaymentForm').querySelector('.msg').textContent='';paymentPreview();$('#paymentModal').style.display='grid'}
function closePaymentModal(){$('#paymentModal').style.display='none'}
async function regenContract(id){try{const x=await api(`/buyurtma/${id}/shartnoma-yaratish`,{method:'POST'});alert('Shartnoma yangilandi. Versiya: '+x.versiya)}catch(e){alert('Xato: '+e.message)}}
async function copyTrack(id){const x=await api(`/api/buyurtma/${id}/link`);try{await navigator.clipboard.writeText(x.url);alert('Mijoz kuzatuv havolasi nusxalandi')}catch(e){prompt('Havolani nusxalang:',x.url)}}
let activeCustomerOrderId=null;
async function openTrack(id){const x=await api(`/api/buyurtma/${id}/link`);window.open(x.url,'_blank')}
async function openCustomerCard(id){
  activeCustomerOrderId=id;const data=await api(`/api/buyurtma/${id}/mijoz-kartasi`),o=data.order,f=$('#customerCardForm');
  $('#ccOrderId').value=id;['taxminiy_sana','taxminiy_vaqt','rang','rang_kodi','material','yaltiroqlik','mijozga_izoh','kechikish_sababi','aloqa_telefon','masul_xodim','avans_talab','avans_muddat_sana','taklif_amal_sana','ishlab_chiqarish_kun','chizma_versiya','bepul_ozgarish_limit'].forEach(n=>{if(f.elements[n])f.elements[n].value=o[n]||''});
  ['muddat_ozgarish_sababi','media_havola','media_izoh','tasdiq_izoh'].forEach(n=>{if(f.elements[n])f.elements[n].value=''});if(f.elements.media_turi)f.elements.media_turi.value='Rasm';if(f.elements.tasdiq_turi)f.elements.tasdiq_turi.value='';
  const media=data.media||[],approvals=data.approvals||[],history=data.deadline_history||[];
  $('#customerCardInfo').innerHTML=`<div style="padding:12px;border-radius:12px;background:#eff6ff"><b>${esc(o.kod)} mijoz oynasi</b><div style="margin-top:8px">${esc(o.buyurtma_turi||'To‘liq mebel')} · ${esc(o.tolov_sharti||'Avans majburiy')}</div><div style="margin-top:5px">Rasm/video: ${media.length} ta · Tasdiqlar: ${approvals.length} ta · Muddat o‘zgarishi: ${history.length} ta</div></div>`+(approvals.length?`<h3 style="margin-top:12px">Tasdiqlar</h3>${approvals.map(a=>`<div style="padding:8px;border-bottom:1px solid #e2e8f0"><b>${esc(a.turi)}</b> — ${esc(a.holat)}<br><small>${esc(a.izoh||'')}</small></div>`).join('')}`:'');
  $('#customerCardModal').style.display='grid';
}
function closeCustomerCard(){$('#customerCardModal').style.display='none';activeCustomerOrderId=null}
async function loadProgress(){return}
const paymentForm=$('#orderPaymentForm');
['paymentCurrency','paymentReceived','paymentRate'].forEach(id=>{const el=$('#'+id);if(el)el.addEventListener('input',paymentPreview)});
paymentForm.onsubmit=async e=>{e.preventDefault();const m=paymentForm.querySelector('.msg'),id=$('#paymentOrderId').value,d=fj(paymentForm);delete d.order_id;try{const x=await api(`/api/buyurtma/${id}/tolovlar`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(d)});m.textContent=`✅ Saqlandi: ${amountText(x.buyurtma_summa,x.buyurtma_valyutasi)}`;await refresh();setTimeout(closePaymentModal,700)}catch(err){m.textContent='❌ '+err.message}};
const orderSearch=$('#orderSearch'),orderStatusFilter=$('#orderStatusFilter');if(orderSearch)orderSearch.addEventListener('input',renderOrders);if(orderStatusFilter)orderStatusFilter.addEventListener('change',renderOrders);

const customerCardForm=$('#customerCardForm');customerCardForm.onsubmit=async e=>{e.preventDefault();const m=customerCardForm.querySelector('.msg');try{const d=fj(customerCardForm),id=d.order_id;delete d.order_id;await api(`/api/buyurtma/${id}/mijoz-kartasi`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(d)});m.textContent='✅ Mijoz kartasi yangilandi';await loadOrders();await openCustomerCard(id)}catch(x){m.textContent='❌ '+x.message}};


async function loadExtras(){const a=await api('/api/qoshimcha-ish');$('#extrasBody').innerHTML=a.map(x=>`<tr><td>${esc(x.kod)}</td><td>${esc(x.mijoz)}</td><td>${esc(x.nomi)}</td><td>${money(x.summa)}</td><td>${esc(x.qoshimcha_kun)}</td></tr>`).join('')}
async function loadService(){const a=await api('/api/servis');$('#serviceBody').innerHTML=a.map(x=>`<tr><td>${esc(x.kod)}</td><td>${esc(x.mijoz)}</td><td>${esc(x.turi)}</td><td>${esc(x.muammo)}</td><td><span class="badge">${esc(x.holat)}</span></td><td>${esc(x.servis_sana||'')}</td><td>${esc(x.usta||'')}</td></tr>`).join('')}
async function loadDelivery(){const a=await api('/api/yetkazish');$('#deliveryBody').innerHTML=a.map(x=>`<tr><td>${esc(x.sana)}</td><td>${esc(x.navbat)}</td><td>${esc(x.kod)}</td><td>${esc(x.mijoz)}</td><td>${esc(x.haydovchi_ism||'')} ${esc(x.haydovchi_familiya||'')}</td><td>${x.lokatsiya?`<a href="${esc(x.lokatsiya)}" target="_blank">Xarita</a>`:esc(x.manzil||'')}</td><td>${esc(x.holat)}</td><td><button onclick="deliveryState(${esc(x.id)},'Yo‘lga chiqdim')">Yo‘lga</button> <button class="ok" onclick="deliveryState(${esc(x.id)},'Yetkazib berdim')">Topshirildi</button></td></tr>`).join('')}
async function deliveryState(id,holat){await api(`/api/yetkazish/${id}/holat`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({holat})});loadDelivery()}

async function refresh(){await Promise.all([loadWorkers(),loadTypes(),loadAttendance(),loadResults(),loadOrders(),loadStock(),loadTrips(),loadPayments(),loadPenalties(),loadDashboard(),loadTotals(),loadExpenses(),loadBonuses(),loadStatuses(),loadFinished(),loadFinance(),loadProgress(),loadExtras(),loadService(),loadDelivery()])}
const rf=$('#ratingForm');rf.onsubmit=async e=>{e.preventDefault();const d=fj(rf),id=d.buyurtma_id;delete d.buyurtma_id;const m=rf.querySelector('.msg');try{await api(`/api/buyurtma/${id}/baho`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(d)});m.textContent='✅ Baho saqlandi'}catch(x){m.textContent='❌ '+x.message}};
bind('#extraForm','/api/qoshimcha-ish');bind('#serviceForm','/api/servis');bind('#deliveryForm','/api/yetkazish');
bind('#workerForm','/api/ishchilar');bind('#attendanceForm','/api/keldi-ketdi');bind('#resultForm','/api/natijalar');setupOrderWorkflowForm();setupOrderCurrencyForm();bind('#orderForm','/api/buyurtmalar');bind('#stockForm','/api/ombor-harakat');bind('#tripForm','/api/safarlar');bind('#paymentForm','/api/tolovlar');bind('#penaltyForm','/api/jarimalar');bind('#expenseForm','/api/xarajatlar');bind('#bonusForm','/api/bonuslar');bind('#statusForm','/api/ishchi-holatlari');bind('#finishedForm','/api/tayyor-mahsulot');setMonth();$('#finStart').value=$('#totalStart').value;$('#finEnd').value=$('#totalEnd').value;refresh();
</script>
<script>
(function(){
  const DAYS=["Yakshanba","Dushanba","Seshanba","Chorshanba","Payshanba","Juma","Shanba"];
  const MONTHS=["yanvar","fevral","mart","aprel","may","iyun","iyul","avgust","sentabr","oktabr","noyabr","dekabr"];
  function two(n){return String(n).padStart(2,"0")}
  function updatePharmClock(){
    const d=new Date(Date.now()+5*60*60*1000);
    const time=two(d.getUTCHours())+":"+two(d.getUTCMinutes())+":"+two(d.getUTCSeconds());
    const date=DAYS[d.getUTCDay()]+", "+d.getUTCDate()+" "+MONTHS[d.getUTCMonth()]+" "+d.getUTCFullYear();
    const clock=document.getElementById("pharmClock");
    const dateBox=document.getElementById("pharmDate");
    if(clock) clock.textContent=time;
    if(dateBox) dateBox.textContent=date+" · Toshkent";
  }
  updatePharmClock();
  setInterval(updatePharmClock,1000);
})();
</script>
</body></html>
"""

# Gunicorn modulni import qilganda ham bazani tayyorlash
init_db()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT",5000)), debug=False)
